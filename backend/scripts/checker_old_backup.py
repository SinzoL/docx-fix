#!/usr/bin/env python3
"""
Word 文档格式检查引擎

根据 YAML 规则配置文件，逐项检查文档的格式是否符合要求。
输出检查报告，包含通过/警告/错误三种级别。

用法：
    python checker.py <docx文件路径> [--rules <yaml规则文件>]
"""

import sys
import os
import re
import yaml
from docx import Document

NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# 中文字体名 ↔ 系统内部名 映射（等价关系）
FONT_ALIASES = {
    'SimHei': '黑体', '黑体': '黑体',
    'SimSun': '宋体', '宋体': '宋体',
    'STXinwei': '华文新魏', '华文新魏': '华文新魏',
    'STKaiti': '华文楷体', '华文楷体': '华文楷体',
    'STSong': '华文宋体', '华文宋体': '华文宋体',
    'STFangsong': '华文仿宋', '华文仿宋': '华文仿宋',
    'KaiTi': '楷体', '楷体': '楷体',
    'FangSong': '仿宋', '仿宋': '仿宋',
    'Microsoft YaHei': '微软雅黑', '微软雅黑': '微软雅黑',
}


def fonts_match(actual, expected):
    """比较两个字体名是否等价（考虑中英文别名）"""
    if actual == expected:
        return True
    actual_cn = FONT_ALIASES.get(actual, actual)
    expected_cn = FONT_ALIASES.get(expected, expected)
    return actual_cn == expected_cn


# ===== 颜色输出 =====
class Color:
    RED = '\033[91m'
    GREEN = '\033[92m'
    YELLOW = '\033[93m'
    BLUE = '\033[94m'
    BOLD = '\033[1m'
    END = '\033[0m'


class CheckResult:
    """单条检查结果"""
    PASS = "PASS"
    WARN = "WARN"
    FAIL = "FAIL"

    def __init__(self, category, item, status, message, location=None, fixable=False):
        self.category = category
        self.item = item
        self.status = status
        self.message = message
        self.location = location  # 如 "段落 35"
        self.fixable = fixable

    def __str__(self):
        loc = f" [{self.location}]" if self.location else ""
        fix = " (可自动修复)" if self.fixable else ""
        if self.status == self.PASS:
            icon = f"{Color.GREEN}✓{Color.END}"
        elif self.status == self.WARN:
            icon = f"{Color.YELLOW}⚠{Color.END}"
        else:
            icon = f"{Color.RED}✗{Color.END}"
        return f"  {icon} [{self.category}] {self.item}{loc}: {self.message}{fix}"


class DocxChecker:
    """文档格式检查器"""

    def __init__(self, filepath, rules_path=None):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.results = []

        # 加载规则
        if rules_path is None:
            rules_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "rules", "default.yaml")
        with open(rules_path, 'r', encoding='utf-8') as f:
            self.rules = yaml.safe_load(f)

    def add_result(self, category, item, status, message, location=None, fixable=False):
        r = CheckResult(category, item, status, message, location, fixable)
        self.results.append(r)
        return r

    # ========================================
    # 1. 页面设置检查
    # ========================================
    def check_page_setup(self):
        rules = self.rules.get('page_setup', {})
        if not rules:
            return

        for i, section in enumerate(self.doc.sections):
            loc = f"节{i+1}" if len(self.doc.sections) > 1 else None

            # 纸张大小
            w_cm = round((section.page_width or 0) / 360000, 1)
            h_cm = round((section.page_height or 0) / 360000, 1)
            exp_w = rules.get('width_cm', 21.0)
            exp_h = rules.get('height_cm', 29.7)
            if abs(w_cm - exp_w) > 0.2 or abs(h_cm - exp_h) > 0.2:
                self.add_result("页面设置", "纸张大小", CheckResult.FAIL,
                                f"当前 {w_cm}×{h_cm}cm, 要求 {exp_w}×{exp_h}cm (A4)", loc, fixable=True)
            else:
                self.add_result("页面设置", "纸张大小", CheckResult.PASS,
                                f"A4 ({w_cm}×{h_cm}cm)", loc)

            # 页边距
            margins = {
                '上边距': (section.top_margin, 'margin_top_cm'),
                '下边距': (section.bottom_margin, 'margin_bottom_cm'),
                '左边距': (section.left_margin, 'margin_left_cm'),
                '右边距': (section.right_margin, 'margin_right_cm'),
            }
            for name, (actual_emu, key) in margins.items():
                if actual_emu is None:
                    continue
                actual_cm = round(actual_emu / 360000, 2)
                expected_cm = rules.get(key)
                if expected_cm is not None:
                    if abs(actual_cm - expected_cm) > 0.1:
                        self.add_result("页面设置", name, CheckResult.FAIL,
                                        f"当前 {actual_cm}cm, 要求 {expected_cm}cm", loc, fixable=True)
                    else:
                        self.add_result("页面设置", name, CheckResult.PASS,
                                        f"{actual_cm}cm", loc)

            # 页眉页脚距离
            if section.header_distance is not None:
                hd_cm = round(section.header_distance / 360000, 2)
                exp_hd = rules.get('header_distance_cm')
                if exp_hd and abs(hd_cm - exp_hd) > 0.1:
                    self.add_result("页面设置", "页眉距顶端", CheckResult.FAIL,
                                    f"当前 {hd_cm}cm, 要求 {exp_hd}cm", loc, fixable=True)

            if section.footer_distance is not None:
                fd_cm = round(section.footer_distance / 360000, 2)
                exp_fd = rules.get('footer_distance_cm')
                if exp_fd and abs(fd_cm - exp_fd) > 0.1:
                    self.add_result("页面设置", "页脚距底端", CheckResult.FAIL,
                                    f"当前 {fd_cm}cm, 要求 {exp_fd}cm", loc, fixable=True)

    # ========================================
    # 2. 页眉页脚检查
    # ========================================
    def check_header_footer(self):
        hf_rules = self.rules.get('header_footer', {})
        header_rules = hf_rules.get('header', {})

        if header_rules:
            expected_text = header_rules.get('text', '')
            found = False
            for section in self.doc.sections:
                header = section.header
                if header and header.paragraphs:
                    for p in header.paragraphs:
                        if p.text.strip():
                            found = True
                            if expected_text and expected_text not in p.text.strip() and p.text.strip() not in expected_text:
                                self.add_result("页眉页脚", "页眉内容", CheckResult.WARN,
                                                f"当前: \"{p.text.strip()[:30]}\", 期望包含: \"{expected_text[:30]}\"")
                            else:
                                self.add_result("页眉页脚", "页眉内容", CheckResult.PASS,
                                                f"\"{p.text.strip()[:40]}\"")
                            break
                if found:
                    break
            if not found:
                self.add_result("页眉页脚", "页眉", CheckResult.FAIL, "未找到页眉内容")

    # ========================================
    # 3. 样式定义检查
    # ========================================
    def _get_style_xml_info(self, style, inherit=True):
        """从样式 XML 中提取格式信息。

        Args:
            style: python-docx 的 Style 对象
            inherit: 是否沿 basedOn 链向上查找缺失属性（默认 True）
        """
        info = {}
        elem = style.element
        pPr = elem.find('.//w:pPr', NSMAP)
        rPr = elem.find('.//w:rPr', NSMAP)

        if pPr is not None:
            # 对齐
            jc = pPr.find('w:jc', NSMAP)
            if jc is not None:
                info['alignment'] = jc.get(f'{{{W}}}val')

            # 行距
            spacing = pPr.find('w:spacing', NSMAP)
            if spacing is not None:
                for attr in ['before', 'after', 'line', 'lineRule', 'beforeLines', 'afterLines']:
                    v = spacing.get(f'{{{W}}}{attr}')
                    if v is not None:
                        info[f'spacing_{attr}'] = v

            # 缩进
            ind = pPr.find('w:ind', NSMAP)
            if ind is not None:
                for attr in ['firstLine', 'firstLineChars', 'left', 'right', 'hanging']:
                    v = ind.get(f'{{{W}}}{attr}')
                    if v is not None:
                        info[f'indent_{attr}'] = v

            # 大纲级别
            outlineLvl = pPr.find('w:outlineLvl', NSMAP)
            if outlineLvl is not None:
                info['outlineLevel'] = int(outlineLvl.get(f'{{{W}}}val'))

        if rPr is not None:
            rFonts = rPr.find('w:rFonts', NSMAP)
            if rFonts is not None:
                for attr in ['ascii', 'eastAsia', 'hAnsi']:
                    v = rFonts.get(f'{{{W}}}{attr}')
                    if v is not None:
                        info[f'font_{attr}'] = v

            sz = rPr.find('w:sz', NSMAP)
            if sz is not None:
                info['fontSize_half_pt'] = int(sz.get(f'{{{W}}}val'))

            b = rPr.find('w:b', NSMAP)
            if b is not None:
                val = b.get(f'{{{W}}}val', 'true')
                info['bold'] = val not in ('0', 'false')

        # 沿 basedOn 链继承缺失属性
        if inherit:
            parent_style = self._get_parent_style(style)
            if parent_style is not None:
                parent_info = self._get_style_xml_info(parent_style, inherit=True)
                for key, val in parent_info.items():
                    if key not in info:
                        info[key] = val

        return info

    def _get_parent_style(self, style):
        """获取样式的 basedOn 父样式"""
        based_on = style.element.find('w:basedOn', NSMAP)
        if based_on is None:
            return None
        parent_id = based_on.get(f'{{{W}}}val')
        if not parent_id:
            return None
        for s in self.doc.styles:
            if s.style_id == parent_id:
                return s
        return None

    def check_style_definitions(self):
        """检查文档中样式定义是否与规则匹配"""
        style_rules = self.rules.get('styles', {})

        for style_name, rules in style_rules.items():
            if rules.get('check_type') == 'content_match':
                continue  # 封面等特殊处理
            if rules.get('should_not_exist'):
                continue  # 说明文字等

            # 查找文档中的样式
            found_style = None
            for style in self.doc.styles:
                if style.name == style_name:
                    found_style = style
                    break

            if found_style is None:
                self.add_result("样式定义", style_name, CheckResult.WARN,
                                f"文档中未找到样式 \"{style_name}\"")
                continue

            actual = self._get_style_xml_info(found_style)

            # 检查段落格式
            para_rules = rules.get('paragraph', {})
            char_rules = rules.get('character', {})

            # 大纲级别
            if 'outline_level' in para_rules:
                expected_lvl = para_rules['outline_level']
                actual_lvl = actual.get('outlineLevel')
                if actual_lvl is None or actual_lvl != expected_lvl:
                    self.add_result("样式定义", f"{style_name} 大纲级别", CheckResult.FAIL,
                                    f"当前={actual_lvl}, 要求={expected_lvl}", fixable=True)
                else:
                    self.add_result("样式定义", f"{style_name} 大纲级别", CheckResult.PASS,
                                    f"级别 {expected_lvl + 1}")

            # 对齐方式
            if 'alignment' in para_rules:
                expected_align = para_rules['alignment']
                actual_align = actual.get('alignment', 'left')
                if actual_align != expected_align:
                    self.add_result("样式定义", f"{style_name} 对齐", CheckResult.FAIL,
                                    f"当前={actual_align}, 要求={expected_align}", fixable=True)

            # 行距
            if 'line_spacing' in para_rules:
                expected_ls = str(para_rules['line_spacing'])
                actual_ls = actual.get('spacing_line')
                if actual_ls and actual_ls != expected_ls:
                    self.add_result("样式定义", f"{style_name} 行距", CheckResult.WARN,
                                    f"当前={actual_ls}, 要求={expected_ls}", fixable=True)

            # 字体
            font_map = {
                'font_ascii': 'font_ascii',
                'font_east_asia': 'font_eastAsia',
                'font_hAnsi': 'font_hAnsi',
            }
            for rule_key, actual_key in font_map.items():
                if rule_key in char_rules:
                    expected_font = char_rules[rule_key]
                    actual_font = actual.get(actual_key)
                    if actual_font and not fonts_match(actual_font, expected_font):
                        self.add_result("样式定义", f"{style_name} 字体({rule_key})", CheckResult.FAIL,
                                        f"当前={actual_font}, 要求={expected_font}", fixable=True)

            # 字号
            if 'font_size_pt' in char_rules:
                expected_pt = char_rules['font_size_pt']
                expected_half = int(expected_pt * 2)
                actual_half = actual.get('fontSize_half_pt')
                if actual_half and actual_half != expected_half:
                    self.add_result("样式定义", f"{style_name} 字号", CheckResult.FAIL,
                                    f"当前={actual_half/2}磅, 要求={expected_pt}磅", fixable=True)

    # ========================================
    # 4. 段落样式正确性检查
    # ========================================
    def _get_para_outline_level(self, para):
        """获取段落的实际大纲级别"""
        # 段落直接设置
        pPr = para._element.find('w:pPr', NSMAP)
        if pPr is not None:
            ol = pPr.find('w:outlineLvl', NSMAP)
            if ol is not None:
                return int(ol.get(f'{{{W}}}val'))

        # 样式定义
        if para.style and para.style.element is not None:
            style_pPr = para.style.element.find('.//w:pPr', NSMAP)
            if style_pPr is not None:
                ol = style_pPr.find('w:outlineLvl', NSMAP)
                if ol is not None:
                    return int(ol.get(f'{{{W}}}val'))

        # Heading 推断
        if para.style and para.style.name:
            m = re.match(r'Heading (\d+)', para.style.name)
            if m:
                return int(m.group(1)) - 1

        return 9

    def check_heading_styles(self):
        """检查标题段落的样式是否正确"""
        structure = self.rules.get('structure', {})
        heading_map = structure.get('heading_style_mapping', {})

        expected_styles = {}
        if heading_map:
            expected_styles[0] = heading_map.get('level_1', 'Heading 1')
            expected_styles[1] = heading_map.get('level_2', '一级节标题2.3')
            expected_styles[2] = heading_map.get('level_3', '二级节标题2.3.1')

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else "None"
            outline_lvl = self._get_para_outline_level(para)

            # 跳过 TOC 条目
            if (style_name or "").lower().startswith('toc ') or style_name == '目录标题':
                continue

            # 检查标题段落的样式与大纲级别一致
            if outline_lvl < 3:  # 1-3级标题
                expected_style = expected_styles.get(outline_lvl)
                if expected_style and style_name != expected_style:
                    # 也接受 Heading N（兼容性）
                    if not (style_name == f"Heading {outline_lvl + 1}"):
                        self.add_result("标题样式", f"段落{i}", CheckResult.WARN,
                                        f"\"{text[:40]}\" 样式={style_name}, "
                                        f"建议使用 \"{expected_style}\"",
                                        f"段落{i}", fixable=True)

            # 检查看起来像标题但不是标题的段落
            if outline_lvl >= 9 and style_name not in ('说明文字', 'Normal', '论文正文-首行缩进',
                                                         '表题注', '续表题注', '图题', '子图题', '续图题',
                                                         '表正文', '公式题注', '参考文献', '关键词'):
                if re.match(r'^\d+\.\d+', text) and len(text) < 80:
                    if not (style_name or "").lower().startswith('toc'):
                        self.add_result("标题样式", f"段落{i}", CheckResult.WARN,
                                        f"\"{text[:40]}\" 看起来像标题，但样式为 \"{style_name}\"(正文级别)",
                                        f"段落{i}")

    # ========================================
    # 5. 文档结构检查
    # ========================================
    def check_document_structure(self):
        """检查必要的章节是否存在"""
        structure = self.rules.get('structure', {})
        required = structure.get('required_chapters', [])

        # 收集所有一级标题
        h1_texts = []
        for para in self.doc.paragraphs:
            if self._get_para_outline_level(para) == 0:
                h1_texts.append(para.text.strip())

        for req in required:
            pattern = req['pattern']
            found = any(pattern in t for t in h1_texts)
            if found:
                self.add_result("文档结构", f"章节 \"{pattern[:20]}\"", CheckResult.PASS, "已找到")
            else:
                self.add_result("文档结构", f"章节 \"{pattern[:20]}\"", CheckResult.FAIL,
                                f"未找到包含 \"{pattern}\" 的一级标题")

    # ========================================
    # 6. TOC 目录域检查
    # ========================================
    def check_toc(self):
        """检查目录域代码"""
        toc_rules = self.rules.get('structure', {}).get('toc', {})
        if not toc_rules:
            return

        # 查找 TOC 域
        toc_found = False
        for para in self.doc.paragraphs:
            for run in para._element.findall('.//w:r', NSMAP):
                instrText = run.find('w:instrText', NSMAP)
                if instrText is not None and instrText.text and 'TOC' in instrText.text:
                    toc_found = True
                    code = instrText.text.strip()

                    # 检查 \o 范围
                    expected_range = toc_rules.get('outline_range', '1-3')
                    m = re.search(r'\\o\s*"1-(\d+)"', code)
                    if m:
                        actual_max = m.group(1)
                        expected_max = expected_range.split('-')[-1]
                        if actual_max != expected_max:
                            self.add_result("目录", "TOC大纲范围", CheckResult.FAIL,
                                            f"当前 \\o \"1-{actual_max}\", 要求 \\o \"1-{expected_max}\"",
                                            fixable=True)
                        else:
                            self.add_result("目录", "TOC大纲范围", CheckResult.PASS,
                                            f"\\o \"1-{actual_max}\"")

                    # 检查 \t 自定义样式映射
                    custom_styles = toc_rules.get('custom_styles', {})
                    if custom_styles:
                        if '\\t' not in code and '\\t' not in code:
                            self.add_result("目录", "自定义样式映射", CheckResult.WARN,
                                            "TOC 中未找到 \\t 自定义样式映射")
                    break
            if toc_found:
                break

        if not toc_found:
            self.add_result("目录", "TOC域", CheckResult.WARN, "未找到目录域代码")

    # ========================================
    # 7. 段落格式一致性检查
    # ========================================
    def _is_cover_page_paragraph(self, para):
        """判断是否为封面页段落（封面使用特殊格式，不应按正文规则检查）"""
        text = para.text.strip()
        style_rules = self.rules.get('styles', {})
        cover_rules = style_rules.get('cover_title', {})
        if cover_rules.get('check_type') == 'content_match':
            patterns = cover_rules.get('content_patterns', [])
            for pattern in patterns:
                if pattern in text:
                    return True
        return False

    def check_paragraph_formatting(self):
        """检查段落级别的格式（字体、字号等）"""
        style_rules = self.rules.get('styles', {})

        # 统计各样式下的格式异常
        issues_by_style = {}

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else "None"

            # 跳过 TOC 条目、说明文字和封面段落
            if (style_name or "").lower().startswith('toc ') or style_name == '说明文字':
                continue
            if self._is_cover_page_paragraph(para):
                continue

            rules = style_rules.get(style_name)
            if not rules:
                continue

            char_rules = rules.get('character', {})
            if not char_rules:
                continue

            # 检查 Run 级别的字体和字号
            for _j, run in enumerate(para.runs):
                rt = run.text
                if not rt or not rt.strip():
                    continue

                # 字号检查
                if run.font.size is not None and 'font_size_pt' in char_rules:
                    actual_pt = run.font.size / 12700
                    expected_pt = char_rules['font_size_pt']
                    if abs(actual_pt - expected_pt) > 0.5:
                        key = f"{style_name}_字号"
                        if key not in issues_by_style:
                            issues_by_style[key] = []
                        issues_by_style[key].append(
                            (i, f"\"{rt[:20]}\" 字号={actual_pt}磅, 要求={expected_pt}磅"))

                # 英文字体检查
                if run.font.name is not None and 'font_ascii' in char_rules:
                    expected = char_rules['font_ascii']
                    if not fonts_match(run.font.name, expected):
                        key = f"{style_name}_英文字体"
                        if key not in issues_by_style:
                            issues_by_style[key] = []
                        issues_by_style[key].append(
                            (i, f"\"{rt[:20]}\" 字体={run.font.name}, 要求={expected}"))

                # 中文字体检查（eastAsia）
                rPr = run._element.find('.//w:rPr', NSMAP)
                if rPr is not None:
                    rFonts = rPr.find('w:rFonts', NSMAP)
                    if rFonts is not None:
                        ea = rFonts.get(f'{{{W}}}eastAsia')
                        expected_ea = char_rules.get('font_east_asia')
                        if ea and expected_ea and not fonts_match(ea, expected_ea):
                            key = f"{style_name}_中文字体"
                            if key not in issues_by_style:
                                issues_by_style[key] = []
                            issues_by_style[key].append(
                                (i, f"\"{rt[:20]}\" 中文字体={ea}, 要求={expected_ea}"))

        # 汇总输出
        for key, issues in issues_by_style.items():
            if len(issues) <= 3:
                for para_idx, msg in issues:
                    self.add_result("段落格式", key, CheckResult.FAIL, msg, f"段落{para_idx}", fixable=True)
            else:
                # 太多就汇总
                self.add_result("段落格式", key, CheckResult.FAIL,
                                f"共 {len(issues)} 处不一致 (首例: {issues[0][1]})", fixable=True)

    # ========================================
    # 7.5 标题编号检查（numId=0 导致编号被禁用）
    # ========================================
    def check_heading_numbering(self):
        """检查标题段落的编号是否被意外禁用（numId=0 覆盖样式编号）"""
        style_rules = self.rules.get('styles', {})

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else "None"
            if (style_name or "").lower().startswith('toc ') or style_name == '说明文字':
                continue

            rules = style_rules.get(style_name)
            if not rules:
                continue
            if rules.get('paragraph', {}).get('outline_level') is None:
                continue

            # 获取样式定义中的编号信息
            style_numId = None
            if para.style and para.style.element is not None:
                s_pPr = para.style.element.find('.//w:pPr', NSMAP)
                if s_pPr is not None:
                    s_numPr = s_pPr.find('w:numPr', NSMAP)
                    if s_numPr is not None:
                        s_numId_el = s_numPr.find('w:numId', NSMAP)
                        if s_numId_el is not None:
                            style_numId = s_numId_el.get(f'{{{W}}}val')

            # 只在样式定义有编号时检查
            if not style_numId or style_numId == '0':
                continue

            # 检查段落级是否覆盖为 numId=0（禁用编号）
            pPr = para._element.find('w:pPr', NSMAP)
            if pPr is not None:
                numPr = pPr.find('w:numPr', NSMAP)
                if numPr is not None:
                    numId_el = numPr.find('w:numId', NSMAP)
                    if numId_el is not None and numId_el.get(f'{{{W}}}val') == '0':
                        self.add_result("标题编号", f"段落{i} 编号被禁用", CheckResult.FAIL,
                                        f"\"{text[:30]}\" 样式 \"{style_name}\" 的编号被 numId=0 覆盖禁用",
                                        f"段落{i}", fixable=True)

    # ========================================
    # 7.6 标题编号 lvlText 格式检查
    # ========================================
    def _get_numbering_part(self):
        """安全获取 numbering part"""
        try:
            return self.doc.part.numbering_part._element
        except Exception:
            return None

    def _get_heading_abstract_num_id(self):
        """找到标题样式绑定的 numId 对应的 abstractNumId"""
        numbering_el = self._get_numbering_part()
        if numbering_el is None:
            return None, None, None

        heading_numId = None
        for style in self.doc.styles:
            if style.name == 'Heading 1' or (hasattr(style, 'element') and
                    style.element.get(f'{{{W}}}styleId') == '1'):
                s_pPr = style.element.find('.//w:pPr', NSMAP)
                if s_pPr is not None:
                    s_numPr = s_pPr.find('w:numPr', NSMAP)
                    if s_numPr is not None:
                        s_numId_el = s_numPr.find('w:numId', NSMAP)
                        if s_numId_el is not None:
                            heading_numId = s_numId_el.get(f'{{{W}}}val')
                break

        if not heading_numId:
            return numbering_el, None, None

        for num_el in numbering_el.findall('w:num', NSMAP):
            if num_el.get(f'{{{W}}}numId') == heading_numId:
                abs_ref = num_el.find('w:abstractNumId', NSMAP)
                if abs_ref is not None:
                    return numbering_el, heading_numId, abs_ref.get(f'{{{W}}}val')
        return numbering_el, heading_numId, None

    def check_heading_lvl_text(self):
        """检查标题编号的 lvlText 格式是否正确。

        python-docx 保存后可能破坏 lvlText（如 '%1.' 而非 '%1'），
        导致标题编号显示错误。
        """
        numbering_el, _heading_numId2, abs_id = self._get_heading_abstract_num_id()
        if numbering_el is None or abs_id is None:
            return

        expected_lvl_texts = {
            '0': '%1', '1': '%1.%2', '2': '%1.%2.%3',
        }

        for abs_num in numbering_el.findall('w:abstractNum', NSMAP):
            if abs_num.get(f'{{{W}}}abstractNumId') != abs_id:
                continue
            for lvl in abs_num.findall('w:lvl', NSMAP):
                ilvl = lvl.get(f'{{{W}}}ilvl')
                if ilvl not in expected_lvl_texts:
                    continue
                pstyle_el = lvl.find('w:pStyle', NSMAP)
                if pstyle_el is None:
                    continue
                text_el = lvl.find('w:lvlText', NSMAP)
                if text_el is None:
                    continue
                actual = text_el.get(f'{{{W}}}val')
                expected = expected_lvl_texts[ilvl]
                if actual != expected:
                    self.add_result("编号格式", f"abstractNum={abs_id} ilvl={ilvl} lvlText",
                                    CheckResult.FAIL,
                                    f"当前=\"{actual}\", 要求=\"{expected}\"",
                                    fixable=True)

    # ========================================
    # 7.7 标题段落级 numId 覆盖检查
    # ========================================
    def check_heading_numid_override(self):
        """检查标题段落是否有段落级 numId 覆盖导致脱离样式编号链。

        如 Heading 1 段落级 numId=11 覆盖了样式的 numId=4，
        使得 Heading 1 与子标题不在同一编号链中。
        """
        style_rules = self.rules.get('styles', {})

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else "None"
            if (style_name or "").lower().startswith('toc ') or style_name == '说明文字':
                continue

            rules = style_rules.get(style_name)
            if not rules:
                continue
            if rules.get('paragraph', {}).get('outline_level') is None:
                continue

            style_numId = None
            if para.style and para.style.element is not None:
                s_pPr = para.style.element.find('.//w:pPr', NSMAP)
                if s_pPr is not None:
                    s_numPr = s_pPr.find('w:numPr', NSMAP)
                    if s_numPr is not None:
                        s_numId_el = s_numPr.find('w:numId', NSMAP)
                        if s_numId_el is not None:
                            style_numId = s_numId_el.get(f'{{{W}}}val')

            if not style_numId or style_numId == '0':
                continue

            pPr = para._element.find('w:pPr', NSMAP)
            if pPr is not None:
                numPr = pPr.find('w:numPr', NSMAP)
                if numPr is not None:
                    numId_el = numPr.find('w:numId', NSMAP)
                    if numId_el is not None:
                        para_numId = numId_el.get(f'{{{W}}}val')
                        if para_numId != '0' and para_numId != style_numId:
                            self.add_result("编号统一", f"段落{i} numId 覆盖",
                                            CheckResult.FAIL,
                                            f"\"{text[:30]}\" 段落级 numId={para_numId} "
                                            f"覆盖了样式 numId={style_numId}，脱离编号链",
                                            f"段落{i}", fixable=True)

    # ========================================
    # 7.8 共享 abstractNum 计数器污染检查
    # ========================================
    def check_shared_abstract_num(self):
        """检查是否有非标题列表共享了标题编号的 abstractNum。

        多个 numId 引用同一 abstractNumId 时会共享计数器，
        普通列表段落会推进标题的编号计数器，导致 Heading 编号跳跃。
        """
        numbering_el, _heading_numId, heading_abs_id = self._get_heading_abstract_num_id()
        if numbering_el is None or heading_abs_id is None:
            return

        # 收集标题样式使用的 numId
        heading_numIds = set()
        style_rules = self.rules.get('styles', {})
        for style in self.doc.styles:
            rules = style_rules.get(style.name)
            if not rules:
                continue
            if rules.get('paragraph', {}).get('outline_level') is None:
                continue
            s_pPr = style.element.find('.//w:pPr', NSMAP)
            if s_pPr is not None:
                s_numPr = s_pPr.find('w:numPr', NSMAP)
                if s_numPr is not None:
                    s_numId_el = s_numPr.find('w:numId', NSMAP)
                    if s_numId_el is not None:
                        heading_numIds.add(s_numId_el.get(f'{{{W}}}val'))

        # 找引用同一 abstractNum 但不是标题的 numId
        polluting_numIds = set()
        for num_el in numbering_el.findall('w:num', NSMAP):
            num_id = num_el.get(f'{{{W}}}numId')
            abs_ref = num_el.find('w:abstractNumId', NSMAP)
            if abs_ref is None:
                continue
            abs_id = abs_ref.get(f'{{{W}}}val')
            if abs_id == heading_abs_id and num_id not in heading_numIds:
                polluting_numIds.add(num_id)

        if not polluting_numIds:
            return

        # 检查这些 numId 是否实际被使用
        body = self.doc.element.body
        used_count = 0
        for p_el in body.findall('w:p', NSMAP):
            pPr = p_el.find('w:pPr', NSMAP)
            if pPr is not None:
                numPr = pPr.find('w:numPr', NSMAP)
                if numPr is not None:
                    ne = numPr.find('w:numId', NSMAP)
                    if ne is not None and ne.get(f'{{{W}}}val') in polluting_numIds:
                        used_count += 1

        if used_count > 0:
            self.add_result("编号隔离", "abstractNum 共享污染", CheckResult.FAIL,
                            f"{len(polluting_numIds)} 个非标题列表编号共享了标题的 "
                            f"abstractNum={heading_abs_id}，影响 {used_count} 个段落，"
                            f"会导致标题编号跳跃",
                            fixable=True)

    # ========================================
    # 7.9 标题样式错误 & 手动编号检查
    # ========================================
    def check_heading_style_and_manual_numbering(self):
        """检查标题段落是否使用了错误的样式（如 Heading 2 而非 一级节标题2.3），
        以及是否包含手动输入的编号前缀。

        这类问题通常发生在：Heading 2/3 样式没有 numPr 配置，
        用户只能手动输入 "1.1"、"1.2.1" 等编号，不参与自动编号。
        """
        fix_rules = self.rules.get('heading_style_fix', {})
        if not fix_rules.get('enabled'):
            return

        style_map = fix_rules.get('style_replacement', {})
        num_patterns = fix_rules.get('manual_numbering_patterns', {})

        if not style_map:
            return

        wrong_style_count = 0
        manual_num_count = 0

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else "None"
            if (style_name or "").lower().startswith('toc ') or style_name == '说明文字':
                continue

            if style_name not in style_map:
                continue

            new_style_name = style_map[style_name]
            wrong_style_count += 1

            # 检查手动编号
            pattern = num_patterns.get(new_style_name)
            has_manual_num = False
            if pattern:
                m = re.match(pattern, text)
                if m:
                    has_manual_num = True
                    manual_num_count += 1

            msg = (f"\"{text[:40]}\" 使用了 \"{style_name}\"，"
                   f"应使用 \"{new_style_name}\"")
            if has_manual_num and m:
                msg += f"，且包含手动编号 \"{m.group(0).strip()}\""

            self.add_result("标题样式", f"段落{i} 样式错误", CheckResult.FAIL,
                            msg, f"段落{i}", fixable=True)

        if wrong_style_count == 0:
            # 额外检查：即使样式正确，也检查是否有手动编号残留
            for i, para in enumerate(self.doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                style_name = para.style.name if para.style else "None"
                if (style_name or "").lower().startswith('toc ') or style_name == '说明文字':
                    continue

                pattern = num_patterns.get(style_name)
                if pattern:
                    m = re.match(pattern, text)
                    if m:
                        self.add_result("标题样式", f"段落{i} 手动编号残留", CheckResult.WARN,
                                        f"\"{text[:40]}\" 样式 \"{style_name}\" 可能包含手动编号 "
                                        f"\"{m.group(0).strip()}\"",
                                        f"段落{i}", fixable=True)

    # ========================================
    # 8. 特殊检查
    # ========================================
    def check_template_instructions(self):
        """检查是否残留模板说明文字。

        区分两种情况：
        1. 真正的模板说明文字（红色字体）→ 警告：未删除模板说明
        2. 被错误设置为"说明文字"样式的正文段落（非红色）→ 错误：样式错误，可自动修复
        """
        special = self.rules.get('special_checks', {})
        check = special.get('template_instructions_check', {})
        if not check.get('enabled'):
            return

        style_name = check.get('style_name', '说明文字')
        template_paras = []  # 真正的模板说明（红色）
        wrong_style_paras = []  # 错误样式的正文段落

        for i, para in enumerate(self.doc.paragraphs):
            if not (para.style and para.style.name == style_name and para.text.strip()):
                continue

            # 判断是否为红色字体
            is_red = False
            for run in para.runs:
                if run.font.color and run.font.color.rgb:
                    rgb = str(run.font.color.rgb)
                    if rgb.upper() in ('FF0000', 'CC0000', 'FF3333'):
                        is_red = True
                        break

            if is_red:
                template_paras.append((i, para.text.strip()[:50]))
            else:
                wrong_style_paras.append((i, para.text.strip()[:50]))

        if template_paras:
            self.add_result("特殊检查", "模板说明文字", CheckResult.WARN,
                            f"发现 {len(template_paras)} 段未删除的模板说明文字(红色)")
            for idx, text in template_paras[:5]:
                self.add_result("特殊检查", f"  说明文字-段落{idx}", CheckResult.WARN,
                                f"\"{text}\"", f"段落{idx}")

        if wrong_style_paras:
            self.add_result("特殊检查", "说明文字样式误用", CheckResult.FAIL,
                            f"发现 {len(wrong_style_paras)} 段正文被错误设置为\"说明文字\"样式",
                            fixable=True)
            for idx, text in wrong_style_paras[:5]:
                self.add_result("特殊检查", f"  误用-段落{idx}", CheckResult.FAIL,
                                f"\"{text}\" 应改为正文样式", f"段落{idx}", fixable=True)

        if not template_paras and not wrong_style_paras:
            self.add_result("特殊检查", "模板说明文字", CheckResult.PASS, "无残留模板说明")

    def check_font_consistency(self):
        """检查正文和标题字体一致性"""
        special = self.rules.get('special_checks', {})

        for check_key in ('body_font_consistency', 'heading_font_consistency'):
            check = special.get(check_key, {})
            if not check.get('enabled'):
                continue

            target_styles = check.get('target_styles', [])
            exp_cn = check.get('expected_chinese_font')
            exp_en = check.get('expected_english_font')
            desc = check.get('description', check_key)

            cn_issues = 0
            en_issues = 0

            for para in self.doc.paragraphs:
                if not para.style or para.style.name not in target_styles:
                    continue
                if not para.text.strip():
                    continue
                # 跳过封面段落
                if self._is_cover_page_paragraph(para):
                    continue

                for run in para.runs:
                    if not run.text.strip():
                        continue

                    # 英文字体
                    if exp_en and run.font.name and not fonts_match(run.font.name, exp_en):
                        en_issues += 1

                    # 中文字体
                    if exp_cn:
                        rPr = run._element.find('.//w:rPr', NSMAP)
                        if rPr is not None:
                            rFonts = rPr.find('w:rFonts', NSMAP)
                            if rFonts is not None:
                                ea = rFonts.get(f'{{{W}}}eastAsia')
                                if ea and not fonts_match(ea, exp_cn):
                                    cn_issues += 1

            if cn_issues == 0 and en_issues == 0:
                self.add_result("字体一致性", desc, CheckResult.PASS,
                                f"中文={exp_cn or '-'}, 英文={exp_en or '-'}")
            else:
                msgs = []
                if cn_issues:
                    msgs.append(f"中文字体不一致 {cn_issues} 处(要求{exp_cn})")
                if en_issues:
                    msgs.append(f"英文字体不一致 {en_issues} 处(要求{exp_en})")
                self.add_result("字体一致性", desc, CheckResult.FAIL,
                                "; ".join(msgs), fixable=True)

    # ========================================
    # 9. 图表标题检查
    # ========================================
    def check_figure_table_captions(self):
        """检查图注和表注的样式以及引用完整性。

        检查内容：
        1. 以"图X-Y"开头的段落是否使用了"图题"样式
        2. 以"表X-Y"开头的段落是否使用了"表题注"样式
        3. 每张图/表是否在正文中被引用
        4. 正文中引用的图/表是否有对应的图注/表注
        """
        import re
        special = self.rules.get('special_checks', {})
        ft_check = special.get('figure_table_caption', {})
        if not ft_check.get('enabled'):
            return

        fig_style_name = ft_check.get('figure_caption_style', '图题')
        tbl_style_name = ft_check.get('table_caption_style', '表题注')

        fig_pattern = re.compile(r'^图\s*(\d+-\d+)')
        tbl_pattern = re.compile(r'^表\s*(\d+-\d+)')
        fig_ref_pattern = re.compile(r'图\s*(\d+-\d+)')
        tbl_ref_pattern = re.compile(r'表\s*(\d+-\d+)')

        # 收集图注/表注及其样式
        fig_captions = {}  # {编号: (段落idx, 样式, 文本)}
        tbl_captions = {}
        fig_wrong_style = []
        tbl_wrong_style = []

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else 'None'

            m = fig_pattern.match(text)
            if m:
                fig_num = m.group(1)
                fig_captions[fig_num] = (i, style_name, text[:50])
                if style_name != fig_style_name:
                    fig_wrong_style.append((i, fig_num, style_name, text[:50]))
                continue

            m = tbl_pattern.match(text)
            if m:
                tbl_num = m.group(1)
                tbl_captions[tbl_num] = (i, style_name, text[:50])
                if style_name != tbl_style_name:
                    tbl_wrong_style.append((i, tbl_num, style_name, text[:50]))

        # 收集正文中的引用（排除图注/表注本身）
        fig_refs = set()
        tbl_refs = set()
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            # 跳过图注/表注本身和目录
            if fig_pattern.match(text) or tbl_pattern.match(text):
                continue
            style_name = para.style.name if para.style else 'None'
            if (style_name or '').lower().startswith('toc'):
                continue

            for m in fig_ref_pattern.finditer(text):
                fig_refs.add(m.group(1))
            for m in tbl_ref_pattern.finditer(text):
                tbl_refs.add(m.group(1))

        # 报告图注样式问题
        if fig_wrong_style:
            self.add_result("图表标题", "图注样式", CheckResult.FAIL,
                            f"{len(fig_wrong_style)} 个图注未使用\"{fig_style_name}\"样式",
                            fixable=True)
            for idx, num, style, text in fig_wrong_style[:5]:
                self.add_result("图表标题", f"  图{num}-段落{idx}", CheckResult.FAIL,
                                f"\"{text}\" 样式=\"{style}\"(应为\"{fig_style_name}\")",
                                f"段落{idx}", fixable=True)
        elif fig_captions:
            self.add_result("图表标题", "图注样式", CheckResult.PASS,
                            f"全部 {len(fig_captions)} 个图注样式正确")

        # 报告表注样式问题
        if tbl_wrong_style:
            self.add_result("图表标题", "表注样式", CheckResult.FAIL,
                            f"{len(tbl_wrong_style)} 个表注未使用\"{tbl_style_name}\"样式",
                            fixable=True)
            for idx, num, style, text in tbl_wrong_style[:5]:
                self.add_result("图表标题", f"  表{num}-段落{idx}", CheckResult.FAIL,
                                f"\"{text}\" 样式=\"{style}\"(应为\"{tbl_style_name}\")",
                                f"段落{idx}", fixable=True)
        elif tbl_captions:
            self.add_result("图表标题", "表注样式", CheckResult.PASS,
                            f"全部 {len(tbl_captions)} 个表注样式正确")

        # 图引用检查
        fig_nums = set(fig_captions.keys())
        unreferenced_figs = fig_nums - fig_refs
        phantom_fig_refs = fig_refs - fig_nums

        if unreferenced_figs:
            self.add_result("图表引用", "未被引用的图", CheckResult.WARN,
                            f"{len(unreferenced_figs)} 张图未在正文中被引用: "
                            f"{', '.join('图' + n for n in sorted(unreferenced_figs))}")
        if phantom_fig_refs:
            self.add_result("图表引用", "悬空图引用", CheckResult.WARN,
                            f"{len(phantom_fig_refs)} 个图引用找不到对应图注: "
                            f"{', '.join('图' + n for n in sorted(phantom_fig_refs))}")
        if not unreferenced_figs and not phantom_fig_refs and fig_captions:
            self.add_result("图表引用", "图引用完整性", CheckResult.PASS,
                            f"全部 {len(fig_captions)} 张图均被正确引用")

        # 表引用检查
        tbl_nums = set(tbl_captions.keys())
        unreferenced_tbls = tbl_nums - tbl_refs
        phantom_tbl_refs = tbl_refs - tbl_nums

        if unreferenced_tbls:
            self.add_result("图表引用", "未被引用的表", CheckResult.WARN,
                            f"{len(unreferenced_tbls)} 个表未在正文中被引用: "
                            f"{', '.join('表' + n for n in sorted(unreferenced_tbls))}")
        if phantom_tbl_refs:
            self.add_result("图表引用", "悬空表引用", CheckResult.WARN,
                            f"{len(phantom_tbl_refs)} 个表引用找不到对应表注: "
                            f"{', '.join('表' + n for n in sorted(phantom_tbl_refs))}")
        if not unreferenced_tbls and not phantom_tbl_refs and tbl_captions:
            self.add_result("图表引用", "表引用完整性", CheckResult.PASS,
                            f"全部 {len(tbl_captions)} 个表均被正确引用")

        # 无图无表
        if not fig_captions and not tbl_captions:
            self.add_result("图表标题", "图表标题", CheckResult.PASS, "文档中无图表标题")

    # ========================================
    # 10. 标题编号缩进检查
    # ========================================
    def check_heading_numbering_indent(self):
        """检查标题编号在 numbering.xml 中的缩进、suff、tabs 是否正确。

        模板要求：
        - ilvl 0: suff=tab, tab_pos=420, left=0, firstLine=0
        - ilvl 1/2: suff=space, left=0, firstLine=0
        python-docx 可能将 suff 重置为 tab、tabs 设为错误值。
        """
        numbering_el = None
        try:
            numbering_el = self.doc.part.numbering_part._element
        except Exception:
            return

        # 从 YAML 获取编号规则
        num_rules = self.rules.get('numbering', {}).get('heading_numbering', {})
        levels = num_rules.get('levels', {})

        # 找标题 abstractNum
        heading_numId = None
        for style in self.doc.styles:
            if style.name == 'Heading 1':
                s_pPr = style.element.find('.//w:pPr', NSMAP)
                if s_pPr is not None:
                    s_numPr = s_pPr.find('w:numPr', NSMAP)
                    if s_numPr is not None:
                        s_numId_el = s_numPr.find('w:numId', NSMAP)
                        if s_numId_el is not None:
                            heading_numId = s_numId_el.get(f'{{{W}}}val')
                break

        if not heading_numId:
            return

        abs_id = None
        for num_el in numbering_el.findall('w:num', NSMAP):
            if num_el.get(f'{{{W}}}numId') == heading_numId:
                abs_ref = num_el.find('w:abstractNumId', NSMAP)
                if abs_ref is not None:
                    abs_id = abs_ref.get(f'{{{W}}}val')
                break

        if not abs_id:
            return

        issues = []
        for abs_num in numbering_el.findall('w:abstractNum', NSMAP):
            if abs_num.get(f'{{{W}}}abstractNumId') != abs_id:
                continue
            for lvl in abs_num.findall('w:lvl', NSMAP):
                ilvl = lvl.get(f'{{{W}}}ilvl')
                ilvl_int = int(ilvl) if ilvl else -1
                lvl_rule = levels.get(ilvl_int, {})
                if not lvl_rule:
                    continue

                # 检查 suff
                expected_suff = lvl_rule.get('suff')
                if expected_suff:
                    suff_el = lvl.find('w:suff', NSMAP)
                    actual_suff = suff_el.get(f'{{{W}}}val') if suff_el is not None else 'tab'
                    if actual_suff != expected_suff:
                        issues.append(f"ilvl={ilvl} suff={actual_suff}(应为{expected_suff})")

                # 检查 tabs
                expected_tab_pos = lvl_rule.get('tab_pos')
                pPr = lvl.find('w:pPr', NSMAP)
                if expected_suff == 'space' and pPr is not None:
                    tabs = pPr.find('w:tabs', NSMAP)
                    if tabs is not None and len(tabs.findall('w:tab', NSMAP)) > 0:
                        issues.append(f"ilvl={ilvl} suff=space时不应有tabs")
                elif expected_tab_pos is not None and pPr is not None:
                    tabs = pPr.find('w:tabs', NSMAP)
                    if tabs is not None:
                        tab_els = tabs.findall('w:tab', NSMAP)
                        if len(tab_els) == 1:
                            pos = tab_els[0].get(f'{{{W}}}pos', '0')
                            if pos != str(expected_tab_pos):
                                issues.append(f"ilvl={ilvl} tab_pos={pos}(应为{expected_tab_pos})")
                    else:
                        issues.append(f"ilvl={ilvl} 缺少tab_pos={expected_tab_pos}")

                # 检查 indent
                expected_left = lvl_rule.get('ind_left')
                if expected_left is not None and pPr is not None:
                    ind = pPr.find('w:ind', NSMAP)
                    if ind is not None:
                        left = ind.get(f'{{{W}}}left', '0')
                        hanging = ind.get(f'{{{W}}}hanging')
                        if left != str(expected_left) or hanging:
                            issues.append(f"ilvl={ilvl} left={left} hanging={hanging or 'N/A'}")

        if issues:
            self.add_result("编号缩进", "标题编号属性", CheckResult.FAIL,
                            f"标题编号级别配置不正确: {'; '.join(issues)}",
                            fixable=True)
        else:
            self.add_result("编号缩进", "标题编号属性", CheckResult.PASS,
                            "标题编号级别属性正确 (suff/tabs/indent)")

    # ========================================
    # 运行所有检查
    # ========================================
    def run_all_checks(self):
        """执行所有检查项"""
        self.results = []

        self.check_page_setup()
        self.check_header_footer()
        self.check_style_definitions()
        self.check_heading_styles()
        self.check_document_structure()
        self.check_toc()
        self.check_heading_numbering()
        self.check_heading_lvl_text()
        self.check_heading_numid_override()
        self.check_shared_abstract_num()
        self.check_heading_style_and_manual_numbering()
        self.check_paragraph_formatting()
        self.check_template_instructions()
        self.check_font_consistency()
        self.check_figure_table_captions()
        self.check_heading_numbering_indent()

        return self.results

    def print_report(self):
        """打印检查报告"""
        print(f"\n{'=' * 70}")
        print(f"  文档格式检查报告: {os.path.basename(self.filepath)}")
        print(f"  规则: {self.rules.get('meta', {}).get('name', 'N/A')}")
        print(f"{'=' * 70}\n")

        # 按类别分组
        categories = {}
        for r in self.results:
            if r.category not in categories:
                categories[r.category] = []
            categories[r.category].append(r)

        pass_count = sum(1 for r in self.results if r.status == CheckResult.PASS)
        warn_count = sum(1 for r in self.results if r.status == CheckResult.WARN)
        fail_count = sum(1 for r in self.results if r.status == CheckResult.FAIL)
        fixable_count = sum(1 for r in self.results if r.fixable)

        for cat, items in categories.items():
            print(f"{Color.BOLD}▸ {cat}{Color.END}")
            for item in items:
                print(str(item))
            print()

        # 汇总
        print(f"{'─' * 70}")
        print(f"  汇总: {Color.GREEN}通过 {pass_count}{Color.END} / "
              f"{Color.YELLOW}警告 {warn_count}{Color.END} / "
              f"{Color.RED}错误 {fail_count}{Color.END}")
        if fixable_count:
            print(f"  其中 {Color.BLUE}{fixable_count} 项可自动修复{Color.END}，"
                  f"运行 fix-format 命令进行修复")
        print()

        return fail_count == 0


def main():
    if len(sys.argv) < 2:
        print("用法: python checker.py <docx文件路径> [--rules <yaml规则文件>]")
        sys.exit(1)

    filepath = sys.argv[1]
    rules_path = None
    if "--rules" in sys.argv:
        idx = sys.argv.index("--rules")
        if idx + 1 < len(sys.argv):
            rules_path = sys.argv[idx + 1]

    checker = DocxChecker(filepath, rules_path)
    checker.run_all_checks()
    checker.print_report()


if __name__ == "__main__":
    main()
