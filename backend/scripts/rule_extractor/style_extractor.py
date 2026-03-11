"""
规则提取器 — 样式提取模块

负责从 Word 文档的 styles.xml 中提取样式定义，
包括段落格式（对齐、间距、缩进等）和字符格式（字体、字号、粗斜体等）。
"""

import re
from collections import OrderedDict

from .constants import NSMAP, W


class StyleExtractorMixin:
    """样式提取功能的 Mixin，注入到 RuleExtractor 中使用。

    依赖宿主类提供：self.doc, self.rules
    """

    def _extract_style_info(self, style, inherit=False):
        """从样式 XML 中提取完整的格式信息"""
        info = OrderedDict()
        para_info = OrderedDict()
        char_info = OrderedDict()

        elem = style.element
        pPr = elem.find('.//w:pPr', NSMAP)
        rPr = elem.find('.//w:rPr', NSMAP)

        # ---- 段落属性 ----
        if pPr is not None:
            self._extract_paragraph_props(pPr, para_info)

        # ---- 字符属性 ----
        if rPr is not None:
            self._extract_character_props(rPr, char_info)

        # 沿 basedOn 链继承
        if inherit:
            parent = self._get_parent_style(style)
            if parent:
                parent_info = self._extract_style_info(parent, inherit=True)
                p_para = parent_info.get('paragraph', {})
                p_char = parent_info.get('character', {})
                for k, v in p_para.items():
                    if k not in para_info:
                        para_info[k] = v
                for k, v in p_char.items():
                    if k not in char_info:
                        char_info[k] = v

        if para_info:
            info['paragraph'] = para_info
        if char_info:
            info['character'] = char_info

        return info

    def _extract_paragraph_props(self, pPr, para_info):
        """从 pPr 元素中提取段落属性"""
        # 对齐方式
        jc = pPr.find('w:jc', NSMAP)
        if jc is not None:
            para_info['alignment'] = jc.get(f'{{{W}}}val')

        # 间距
        spacing = pPr.find('w:spacing', NSMAP)
        if spacing is not None:
            before_lines = spacing.get(f'{{{W}}}beforeLines')
            if before_lines:
                para_info['spacing_before_lines'] = int(before_lines)
            after_lines = spacing.get(f'{{{W}}}afterLines')
            if after_lines:
                para_info['spacing_after_lines'] = int(after_lines)
            before = spacing.get(f'{{{W}}}before')
            if before and not before_lines:
                para_info['spacing_before_twips'] = int(before)
            after = spacing.get(f'{{{W}}}after')
            if after and not after_lines:
                para_info['spacing_after_twips'] = int(after)
            line = spacing.get(f'{{{W}}}line')
            if line:
                para_info['line_spacing'] = int(line)
            line_rule = spacing.get(f'{{{W}}}lineRule')
            if line_rule:
                para_info['line_spacing_rule'] = line_rule

        # 缩进
        ind = pPr.find('w:ind', NSMAP)
        if ind is not None:
            first_line_chars = ind.get(f'{{{W}}}firstLineChars')
            if first_line_chars:
                para_info['first_line_indent_chars'] = int(first_line_chars)
            first_line = ind.get(f'{{{W}}}firstLine')
            if first_line and not first_line_chars:
                para_info['first_line_indent_twips'] = int(first_line)
            left = ind.get(f'{{{W}}}left')
            if left and left != '0':
                para_info['indent_left_twips'] = int(left)
            right = ind.get(f'{{{W}}}right')
            if right and right != '0':
                para_info['indent_right_twips'] = int(right)
            hanging = ind.get(f'{{{W}}}hanging')
            if hanging:
                para_info['indent_hanging_twips'] = int(hanging)

        # 大纲级别
        outline_lvl = pPr.find('w:outlineLvl', NSMAP)
        if outline_lvl is not None:
            para_info['outline_level'] = int(outline_lvl.get(f'{{{W}}}val'))

        # 与下段同页
        keep_next = pPr.find('w:keepNext', NSMAP)
        if keep_next is not None:
            val = keep_next.get(f'{{{W}}}val', 'true')
            if val not in ('0', 'false'):
                para_info['keep_next'] = True

        # 段中不分页
        keep_lines = pPr.find('w:keepLines', NSMAP)
        if keep_lines is not None:
            val = keep_lines.get(f'{{{W}}}val', 'true')
            if val not in ('0', 'false'):
                para_info['keep_lines'] = True

    def _extract_character_props(self, rPr, char_info):
        """从 rPr 元素中提取字符属性"""
        # 字体
        rFonts = rPr.find('w:rFonts', NSMAP)
        if rFonts is not None:
            for attr, rule_key in [('ascii', 'font_ascii'),
                                    ('eastAsia', 'font_east_asia'),
                                    ('hAnsi', 'font_hAnsi'),
                                    ('cs', 'font_cs')]:
                v = rFonts.get(f'{{{W}}}{attr}')
                if v:
                    char_info[rule_key] = v

        # 字号
        sz = rPr.find('w:sz', NSMAP)
        if sz is not None:
            half_pt = int(sz.get(f'{{{W}}}val'))
            char_info['font_size_pt'] = half_pt / 2

        # 字号（cs）
        szCs = rPr.find('w:szCs', NSMAP)
        if szCs is not None:
            half_pt_cs = int(szCs.get(f'{{{W}}}val'))
            char_info['font_size_cs_pt'] = half_pt_cs / 2

        # 粗体
        b = rPr.find('w:b', NSMAP)
        if b is not None:
            val = b.get(f'{{{W}}}val', 'true')
            char_info['bold'] = val not in ('0', 'false')

        # 斜体
        i_el = rPr.find('w:i', NSMAP)
        if i_el is not None:
            val = i_el.get(f'{{{W}}}val', 'true')
            char_info['italic'] = val not in ('0', 'false')

        # 字间距
        spacing_el = rPr.find('w:spacing', NSMAP)
        if spacing_el is not None:
            val = spacing_el.get(f'{{{W}}}val')
            if val:
                char_info['char_spacing'] = int(val)

        # 字距调整（kern）
        kern = rPr.find('w:kern', NSMAP)
        if kern is not None:
            val = kern.get(f'{{{W}}}val')
            if val:
                char_info['kern'] = int(val)

        # 颜色
        color = rPr.find('w:color', NSMAP)
        if color is not None:
            val = color.get(f'{{{W}}}val')
            if val and val.upper() != 'auto' and val.upper() != '000000':
                char_info['color'] = val.upper()

    def _get_parent_style(self, style):
        """获取样式的 basedOn 父样式"""
        based_on = style.element.find('w:basedOn', NSMAP)
        if based_on is None:
            return None
        parent_id = based_on.get(f'{{{W}}}val')
        if not parent_id:
            return None
        for s in self.doc.styles:
            if s.style_id == parent_id:
                return s
        return None

    def _get_style_based_on(self, style):
        """获取 basedOn 样式名"""
        based_on = style.element.find('w:basedOn', NSMAP)
        if based_on is None:
            return None
        parent_id = based_on.get(f'{{{W}}}val')
        if not parent_id:
            return None
        for s in self.doc.styles:
            if s.style_id == parent_id:
                return s.name
        return parent_id

    def extract_styles(self):
        """提取文档中所有有意义的样式定义"""
        styles = OrderedDict()

        # 收集文档中实际使用的样式名
        used_styles = set()
        for para in self.doc.paragraphs:
            if para.style:
                used_styles.add(para.style.name)
        # 也检查页眉页脚
        for section in self.doc.sections:
            for p in section.header.paragraphs:
                if p.style:
                    used_styles.add(p.style.name)
            for p in section.footer.paragraphs:
                if p.style:
                    used_styles.add(p.style.name)

        # 检查每个样式
        for style in self.doc.styles:
            # 只处理段落样式
            if style.type is not None and str(style.type) != 'PARAGRAPH (1)' and style.type != 1:
                try:
                    from docx.enum.style import WD_STYLE_TYPE
                    if style.type != WD_STYLE_TYPE.PARAGRAPH:
                        continue
                except Exception:
                    if 'PARAGRAPH' not in str(style.type):
                        continue

            name = style.name
            if not name:
                continue

            # 跳过无内容的内置样式
            # 但保留常用内置样式（Normal, Heading 1-9 等）和文档中使用的样式
            is_used = name in used_styles
            is_important = name in ('Normal', 'Header', 'Footer', '目录标题') or \
                           name.startswith('Heading ') or \
                           name.lower().startswith('toc ')

            if not is_used and not is_important:
                continue

            info = self._extract_style_info(style, inherit=False)

            # 如果没有任何有效的格式信息，跳过
            if not info.get('paragraph') and not info.get('character'):
                if not is_used:
                    continue

            style_entry = OrderedDict()

            # 添加描述（基于样式名推断）
            desc = self._generate_style_description(name, style)
            if desc:
                style_entry['description'] = desc

            # basedOn
            based_on = self._get_style_based_on(style)
            if based_on and based_on != name:
                style_entry['based_on'] = based_on

            # 检查是否应标记为 should_not_exist（如"说明文字"类红色样式）
            if self._is_instruction_style(style):
                style_entry['should_not_exist'] = True

            # 段落格式
            if info.get('paragraph'):
                style_entry['paragraph'] = info['paragraph']

            # 字符格式
            if info.get('character'):
                style_entry['character'] = info['character']

            if style_entry:
                styles[name] = style_entry

        self.rules['styles'] = styles

    def _is_instruction_style(self, style):
        """判断是否是模板说明文字样式（红色字体）"""
        rPr = style.element.find('.//w:rPr', NSMAP)
        if rPr is not None:
            color = rPr.find('w:color', NSMAP)
            if color is not None:
                val = (color.get(f'{{{W}}}val') or '').upper()
                if val in ('FF0000', 'CC0000', 'FF3333'):
                    return True
        return False

    def _generate_style_description(self, name, style):
        """根据样式名自动生成描述"""
        desc_map = {
            'Normal': '基础正文样式',
            'Header': '页眉样式',
            'Footer': '页脚样式',
            '目录标题': '目录页标题',
        }
        if name in desc_map:
            return desc_map[name]

        # Heading 系列
        m = re.match(r'Heading (\d+)', name)
        if m:
            lvl = int(m.group(1))
            return f'{lvl}级标题'

        # TOC 系列
        m = re.match(r'toc (\d+)', name, re.IGNORECASE)
        if m:
            return f'目录中的{m.group(1)}级条目'

        # 中文样式名通常自解释
        if any('\u4e00' <= c <= '\u9fff' for c in name):
            return f'自定义样式 "{name}"'

        return None
