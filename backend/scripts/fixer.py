#!/usr/bin/env python3
"""
Word 文档格式自动修复引擎

根据 YAML 规则配置文件，自动修复文档的格式问题。
仅修复格式，不改动文档内容。

用法：
    python fixer.py <docx文件路径> [--rules <yaml规则文件>] [--dry-run]
"""

import sys
import os
import re
import shutil
import yaml
from docx import Document
from typing import cast
from docx.shared import Pt, Cm
from docx.styles.style import ParagraphStyle
from lxml import etree  # type: ignore[attr-defined]

NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# 中文字体名 ↔ 系统内部名 映射
FONT_ALIASES = {
    'SimHei': '黑体', '黑体': '黑体',
    'SimSun': '宋体', '宋体': '宋体',
    'STXinwei': '华文新魏', '华文新魏': '华文新魏',
    'STKaiti': '华文楷体', '华文楷体': '华文楷体',
    'STSong': '华文宋体', '华文宋体': '华文宋体',
    'STFangsong': '华文仿宋', '华文仿宋': '华文仿宋',
    'KaiTi': '楷体', '楷体': '楷体',
    'FangSong': '仿宋', '仿宋': '仿宋',
    'Microsoft YaHei': '微软雅黑', '微软雅黑': '微软雅黑',
}


def fonts_match(actual, expected):
    if actual == expected:
        return True
    return FONT_ALIASES.get(actual, actual) == FONT_ALIASES.get(expected, expected)


class DocxFixer:
    """文档格式修复器"""

    def __init__(self, filepath, rules_path=None):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.fixes = []
        self.dry_run = False

        if rules_path is None:
            rules_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "rules", "default.yaml")
        with open(rules_path, 'r', encoding='utf-8') as f:
            self.rules = yaml.safe_load(f)

    def log_fix(self, category, description):
        self.fixes.append((category, description))
        prefix = "[DRY RUN] " if self.dry_run else ""
        print(f"  {prefix}🔧 [{category}] {description}")

    # ========================================
    # 1. 修复页面设置
    # ========================================
    def fix_page_setup(self):
        rules = self.rules.get('page_setup', {})
        if not rules:
            return

        for section in self.doc.sections:
            # 纸张大小
            exp_w = rules.get('width_cm')
            exp_h = rules.get('height_cm')
            if exp_w and exp_h:
                if section.page_width is None or section.page_height is None:
                    continue
                actual_w = round(section.page_width / 360000, 1)
                actual_h = round(section.page_height / 360000, 1)
                if abs(actual_w - exp_w) > 0.2 or abs(actual_h - exp_h) > 0.2:
                    self.log_fix("页面设置", f"纸张大小 {actual_w}×{actual_h} → {exp_w}×{exp_h}cm")
                    if not self.dry_run:
                        section.page_width = Cm(exp_w)
                        section.page_height = Cm(exp_h)

            # 页边距
            margin_map = {
                'margin_top_cm': 'top_margin',
                'margin_bottom_cm': 'bottom_margin',
                'margin_left_cm': 'left_margin',
                'margin_right_cm': 'right_margin',
            }
            for rule_key, attr in margin_map.items():
                expected = rules.get(rule_key)
                if expected is None:
                    continue
                actual_emu = getattr(section, attr)
                if actual_emu is None:
                    continue
                actual_cm = round(actual_emu / 360000, 2)
                if abs(actual_cm - expected) > 0.1:
                    self.log_fix("页面设置", f"{attr} {actual_cm}cm → {expected}cm")
                    if not self.dry_run:
                        setattr(section, attr, Cm(expected))

            # 页眉距离
            exp_hd = rules.get('header_distance_cm')
            if exp_hd and section.header_distance is not None:
                actual = round(section.header_distance / 360000, 2)
                if abs(actual - exp_hd) > 0.1:
                    self.log_fix("页面设置", f"页眉距顶端 {actual}cm → {exp_hd}cm")
                    if not self.dry_run:
                        section.header_distance = Cm(exp_hd)

            # 页脚距离
            exp_fd = rules.get('footer_distance_cm')
            if exp_fd and section.footer_distance is not None:
                actual = round(section.footer_distance / 360000, 2)
                if abs(actual - exp_fd) > 0.1:
                    self.log_fix("页面设置", f"页脚距底端 {actual}cm → {exp_fd}cm")
                    if not self.dry_run:
                        section.footer_distance = Cm(exp_fd)

    # ========================================
    # 2. 修复样式定义
    # ========================================
    def _set_style_outline_level(self, style_elem, level):
        pPr = style_elem.find('.//w:pPr', NSMAP)
        if pPr is None:
            pPr = etree.SubElement(style_elem, f'{{{W}}}pPr')
        ol = pPr.find('w:outlineLvl', NSMAP)
        if ol is None:
            ol = etree.SubElement(pPr, f'{{{W}}}outlineLvl')
        ol.set(f'{{{W}}}val', str(level))

    def _set_style_alignment(self, style_elem, alignment):
        pPr = style_elem.find('.//w:pPr', NSMAP)
        if pPr is None:
            pPr = etree.SubElement(style_elem, f'{{{W}}}pPr')
        jc = pPr.find('w:jc', NSMAP)
        if jc is None:
            jc = etree.SubElement(pPr, f'{{{W}}}jc')
        jc.set(f'{{{W}}}val', alignment)

    def _set_style_spacing(self, style_elem, **kwargs):
        pPr = style_elem.find('.//w:pPr', NSMAP)
        if pPr is None:
            pPr = etree.SubElement(style_elem, f'{{{W}}}pPr')
        spacing = pPr.find('w:spacing', NSMAP)
        if spacing is None:
            spacing = etree.SubElement(pPr, f'{{{W}}}spacing')
        for k, v in kwargs.items():
            spacing.set(f'{{{W}}}{k}', str(v))

    def _set_style_font(self, style_elem, **kwargs):
        rPr = style_elem.find('.//w:rPr', NSMAP)
        if rPr is None:
            rPr = etree.SubElement(style_elem, f'{{{W}}}rPr')
        rFonts = rPr.find('w:rFonts', NSMAP)
        if rFonts is None:
            rFonts = etree.SubElement(rPr, f'{{{W}}}rFonts')
        for k, v in kwargs.items():
            rFonts.set(f'{{{W}}}{k}', v)

    def _set_style_font_size(self, style_elem, half_pt):
        rPr = style_elem.find('.//w:rPr', NSMAP)
        if rPr is None:
            rPr = etree.SubElement(style_elem, f'{{{W}}}rPr')
        sz = rPr.find('w:sz', NSMAP)
        if sz is None:
            sz = etree.SubElement(rPr, f'{{{W}}}sz')
        sz.set(f'{{{W}}}val', str(half_pt))
        szCs = rPr.find('w:szCs', NSMAP)
        if szCs is None:
            szCs = etree.SubElement(rPr, f'{{{W}}}szCs')
        szCs.set(f'{{{W}}}val', str(half_pt))

    def _get_style_outline_level(self, style_elem):
        pPr = style_elem.find('.//w:pPr', NSMAP)
        if pPr is not None:
            ol = pPr.find('w:outlineLvl', NSMAP)
            if ol is not None:
                return int(ol.get(f'{{{W}}}val'))
        return None

    def _get_style_font_info(self, style_elem):
        info = {}
        rPr = style_elem.find('.//w:rPr', NSMAP)
        if rPr is not None:
            rFonts = rPr.find('w:rFonts', NSMAP)
            if rFonts is not None:
                for attr in ['ascii', 'eastAsia', 'hAnsi']:
                    v = rFonts.get(f'{{{W}}}{attr}')
                    if v:
                        info[attr] = v
            sz = rPr.find('w:sz', NSMAP)
            if sz is not None:
                info['fontSize_half_pt'] = int(sz.get(f'{{{W}}}val'))
        return info

    def _get_style_spacing(self, style_elem):
        info = {}
        pPr = style_elem.find('.//w:pPr', NSMAP)
        if pPr is not None:
            spacing = pPr.find('w:spacing', NSMAP)
            if spacing is not None:
                for attr in ['line', 'lineRule', 'before', 'after', 'beforeLines', 'afterLines']:
                    v = spacing.get(f'{{{W}}}{attr}')
                    if v:
                        info[attr] = v
        return info

    def _get_style_alignment(self, style_elem):
        pPr = style_elem.find('.//w:pPr', NSMAP)
        if pPr is not None:
            jc = pPr.find('w:jc', NSMAP)
            if jc is not None:
                return jc.get(f'{{{W}}}val')
        return None

    def fix_style_definitions(self):
        """修复样式定义中的格式"""
        style_rules = self.rules.get('styles', {})

        for style_name, rules in style_rules.items():
            if rules.get('check_type') == 'content_match':
                continue
            if rules.get('should_not_exist'):
                continue

            # 查找样式
            found_style = None
            for style in self.doc.styles:
                if style.name == style_name:
                    found_style = style
                    break

            if found_style is None:
                continue

            para_rules = rules.get('paragraph', {})
            char_rules = rules.get('character', {})
            elem = found_style.element

            # 大纲级别
            if 'outline_level' in para_rules:
                expected = para_rules['outline_level']
                actual = self._get_style_outline_level(elem)
                if actual is None or actual != expected:
                    self.log_fix("样式定义", f"{style_name} 大纲级别 {actual} → {expected}")
                    if not self.dry_run:
                        self._set_style_outline_level(elem, expected)

            # 对齐
            if 'alignment' in para_rules:
                expected = para_rules['alignment']
                actual = self._get_style_alignment(elem)
                if actual != expected:
                    self.log_fix("样式定义", f"{style_name} 对齐 {actual} → {expected}")
                    if not self.dry_run:
                        self._set_style_alignment(elem, expected)

            # 行距
            actual_spacing = self._get_style_spacing(elem)
            spacing_updates = {}
            if 'line_spacing' in para_rules:
                expected_ls = str(para_rules['line_spacing'])
                if actual_spacing.get('line') != expected_ls:
                    spacing_updates['line'] = expected_ls
            if 'line_spacing_rule' in para_rules:
                expected_lr = para_rules['line_spacing_rule']
                if actual_spacing.get('lineRule') != expected_lr:
                    spacing_updates['lineRule'] = expected_lr
            if 'spacing_before_lines' in para_rules:
                expected = str(para_rules['spacing_before_lines'])
                if actual_spacing.get('beforeLines') != expected:
                    spacing_updates['beforeLines'] = expected
            if 'spacing_after_lines' in para_rules:
                expected = str(para_rules['spacing_after_lines'])
                if actual_spacing.get('afterLines') != expected:
                    spacing_updates['afterLines'] = expected

            if spacing_updates:
                self.log_fix("样式定义", f"{style_name} 间距 {spacing_updates}")
                if not self.dry_run:
                    self._set_style_spacing(elem, **spacing_updates)

            # 字体
            actual_font = self._get_style_font_info(elem)
            font_updates = {}
            font_key_map = {'font_ascii': 'ascii', 'font_east_asia': 'eastAsia', 'font_hAnsi': 'hAnsi'}
            for rule_key, xml_key in font_key_map.items():
                if rule_key in char_rules:
                    expected = char_rules[rule_key]
                    actual_val = actual_font.get(xml_key)
                    if actual_val and not fonts_match(actual_val, expected):
                        font_updates[xml_key] = expected

            if font_updates:
                self.log_fix("样式定义", f"{style_name} 字体 {font_updates}")
                if not self.dry_run:
                    self._set_style_font(elem, **font_updates)

            # 字号
            if 'font_size_pt' in char_rules:
                expected_pt = char_rules['font_size_pt']
                expected_half = int(expected_pt * 2)
                actual_half = actual_font.get('fontSize_half_pt')
                if actual_half is None or actual_half != expected_half:
                    self.log_fix("样式定义", f"{style_name} 字号 {(actual_half or 0)/2}磅 → {expected_pt}磅")
                    if not self.dry_run:
                        self._set_style_font_size(elem, expected_half)

    # ========================================
    # 3. 修复段落大纲级别
    # ========================================
    def fix_paragraph_outline_levels(self):
        """确保标题段落有正确的大纲级别"""
        style_rules = self.rules.get('styles', {})

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            style_name = (para.style.name or "None") if para.style else "None"
            if style_name.lower().startswith('toc ') or style_name == '说明文字':
                continue

            rules = style_rules.get(style_name)
            if not rules:
                continue

            para_rules = rules.get('paragraph', {})
            if 'outline_level' not in para_rules:
                continue

            expected_lvl = para_rules['outline_level']

            # 检查段落直接设置的大纲级别
            pPr = para._element.find('w:pPr', NSMAP)
            current_lvl = None
            if pPr is not None:
                ol = pPr.find('w:outlineLvl', NSMAP)
                if ol is not None:
                    current_lvl = int(ol.get(f'{{{W}}}val'))

            if current_lvl is not None and current_lvl != expected_lvl:
                self.log_fix("段落大纲", f"段落{i} \"{text[:30]}\" 大纲级别 {current_lvl} → {expected_lvl}")
                if not self.dry_run:
                    ol.set(f'{{{W}}}val', str(expected_lvl))

    # ========================================
    # 4. 修复 TOC 域代码
    # ========================================
    def fix_toc(self):
        """修复目录域代码"""
        toc_rules = self.rules.get('structure', {}).get('toc', {})
        if not toc_rules:
            return

        expected_range = toc_rules.get('outline_range', '1-3')
        expected_max = expected_range.split('-')[-1]

        for para in self.doc.paragraphs:
            for run in para._element.findall('.//w:r', NSMAP):
                instrText = run.find('w:instrText', NSMAP)
                if instrText is not None and instrText.text and 'TOC' in instrText.text:
                    old_text = instrText.text
                    m = re.search(r'\\o\s*"1-(\d+)"', old_text)
                    if m and m.group(1) != expected_max:
                        new_text = re.sub(r'\\o\s*"1-\d+"', f'\\\\o "1-{expected_max}"', old_text)
                        self.log_fix("TOC", f"大纲范围 \\o \"1-{m.group(1)}\" → \\o \"1-{expected_max}\"")
                        if not self.dry_run:
                            instrText.text = new_text

    # ========================================
    # 5. 修复段落字体（Run 级别）
    # ========================================
    def _is_cover_page_paragraph(self, para):
        """判断是否为封面页段落"""
        text = para.text.strip()
        style_rules = self.rules.get('styles', {})
        cover_rules = style_rules.get('cover_title', {})
        if cover_rules.get('check_type') == 'content_match':
            for pattern in cover_rules.get('content_patterns', []):
                if pattern in text:
                    return True
        return False

    def fix_run_fonts(self):
        """修复段落中 Run 级别的字体和字号"""
        style_rules = self.rules.get('styles', {})
        fix_count = 0

        for _i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            style_name = (para.style.name or "None") if para.style else "None"
            if style_name.lower().startswith('toc ') or style_name == '说明文字':
                continue
            # 跳过封面段落
            if self._is_cover_page_paragraph(para):
                continue

            rules = style_rules.get(style_name)
            if not rules:
                continue

            char_rules = rules.get('character', {})
            if not char_rules:
                continue

            for run in para.runs:
                if not run.text or not run.text.strip():
                    continue

                fixed_this_run = False

                # 字号修复
                if run.font.size is not None and 'font_size_pt' in char_rules:
                    expected_pt = char_rules['font_size_pt']
                    actual_pt = run.font.size / 12700
                    if abs(actual_pt - expected_pt) > 0.5:
                        if not self.dry_run:
                            run.font.size = Pt(expected_pt)
                        fixed_this_run = True

                # 英文字体修复
                if run.font.name is not None and 'font_ascii' in char_rules:
                    expected = char_rules['font_ascii']
                    if not fonts_match(run.font.name, expected):
                        if not self.dry_run:
                            run.font.name = expected
                        fixed_this_run = True

                # 中文字体修复（eastAsia）
                exp_ea = char_rules.get('font_east_asia')
                if exp_ea:
                    rPr = run._element.find('.//w:rPr', NSMAP)
                    if rPr is not None:
                        rFonts = rPr.find('w:rFonts', NSMAP)
                        if rFonts is not None:
                            ea = rFonts.get(f'{{{W}}}eastAsia')
                            if ea and not fonts_match(ea, exp_ea):
                                if not self.dry_run:
                                    rFonts.set(f'{{{W}}}eastAsia', exp_ea)
                                fixed_this_run = True

                if fixed_this_run:
                    fix_count += 1

        if fix_count > 0:
            self.log_fix("段落字体", f"共修复 {fix_count} 个 Run 的字体/字号")

    # ========================================
    # 6. 修复被禁用的标题编号（numId=0 覆盖）
    # ========================================
    def fix_disabled_heading_numbering(self):
        """修复标题段落中 numId=0 导致编号被禁用的问题。
        
        当段落级 numPr 中 numId=0 时，它会覆盖样式定义中的编号，
        导致标题不显示自动编号（如 3.1、4.1 等）。
        修复方式：移除段落级的 numPr，让编号继承自样式定义。
        """
        style_rules = self.rules.get('styles', {})
        fix_count = 0

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name or "None") if para.style else "None"
            if style_name.lower().startswith('toc ') or style_name == '说明文字':
                continue

            rules = style_rules.get(style_name)
            if not rules:
                continue
            if rules.get('paragraph', {}).get('outline_level') is None:
                continue

            # 获取样式定义中的编号信息
            style_numId = None
            if para.style and para.style.element is not None:
                s_pPr = para.style.element.find('.//w:pPr', NSMAP)
                if s_pPr is not None:
                    s_numPr = s_pPr.find('w:numPr', NSMAP)
                    if s_numPr is not None:
                        s_numId_el = s_numPr.find('w:numId', NSMAP)
                        if s_numId_el is not None:
                            style_numId = s_numId_el.get(f'{{{W}}}val')

            if not style_numId or style_numId == '0':
                continue

            # 检查段落级 numId=0
            pPr = para._element.find('w:pPr', NSMAP)
            if pPr is not None:
                numPr = pPr.find('w:numPr', NSMAP)
                if numPr is not None:
                    numId_el = numPr.find('w:numId', NSMAP)
                    if numId_el is not None and numId_el.get(f'{{{W}}}val') == '0':
                        self.log_fix("编号恢复",
                                     f"段落{i} \"{text[:30]}\" 移除 numId=0 覆盖，恢复样式编号")
                        if not self.dry_run:
                            pPr.remove(numPr)
                        fix_count += 1

        if fix_count > 0:
            self.log_fix("编号恢复", f"共恢复 {fix_count} 个标题段落的自动编号")

    # ========================================
    # 7. 修复 numbering.xml 中标题编号的 lvlText 格式
    # ========================================
    def _get_numbering_part(self):
        """安全获取 numbering part，不存在则返回 None"""
        try:
            return self.doc.part.numbering_part._element
        except Exception:
            return None

    def _get_heading_abstract_num_id(self):
        """找到标题样式绑定的 numId 对应的 abstractNumId"""
        numbering_el = self._get_numbering_part()
        if numbering_el is None:
            return None, None, None

        # 从样式中获取 Heading 1 的 numId
        heading_numId = None
        for style in self.doc.styles:
            if style.name == 'Heading 1' or (hasattr(style, 'element') and
                    style.element.get(f'{{{W}}}styleId') == '1'):
                s_pPr = style.element.find('.//w:pPr', NSMAP)
                if s_pPr is not None:
                    s_numPr = s_pPr.find('w:numPr', NSMAP)
                    if s_numPr is not None:
                        s_numId_el = s_numPr.find('w:numId', NSMAP)
                        if s_numId_el is not None:
                            heading_numId = s_numId_el.get(f'{{{W}}}val')
                break

        if not heading_numId:
            return numbering_el, None, None

        # 找 numId -> abstractNumId
        for num_el in numbering_el.findall('w:num', NSMAP):
            if num_el.get(f'{{{W}}}numId') == heading_numId:
                abs_ref = num_el.find('w:abstractNumId', NSMAP)
                if abs_ref is not None:
                    return numbering_el, heading_numId, abs_ref.get(f'{{{W}}}val')
        return numbering_el, heading_numId, None

    def fix_heading_lvl_text(self):
        """修复 numbering.xml 中标题编号的 lvlText 格式。

        python-docx 保存文档后，abstractNum 的 lvlText 可能变成错误格式
        如 '%1.'、'%2.'、'%3.' 而非正确的多级编号格式 '%1'、'%1.%2'、'%1.%2.%3'。
        这会导致标题编号显示为 "1." 而不是 "1"，子标题显示 "1." 而不是 "1.1"。
        """
        numbering_el, _heading_numId, abs_id = self._get_heading_abstract_num_id()
        if numbering_el is None or abs_id is None:
            return

        # 期望的 lvlText 格式（仅标题编号级别 ilvl 0-2）
        # 注意：ilvl=3 是图编号(如"图%1-%4")，ilvl=4 是表编号(如"表%1-%5")，
        # 这些不是标题编号，不应被修改
        expected_lvl_texts = {
            '0': '%1',
            '1': '%1.%2',
            '2': '%1.%2.%3',
        }

        for abs_num in numbering_el.findall('w:abstractNum', NSMAP):
            if abs_num.get(f'{{{W}}}abstractNumId') != abs_id:
                continue
            for lvl in abs_num.findall('w:lvl', NSMAP):
                ilvl = lvl.get(f'{{{W}}}ilvl')
                if ilvl not in expected_lvl_texts:
                    continue
                # 检查是否有 pStyle 绑定（说明这是标题编号级别）
                pstyle_el = lvl.find('w:pStyle', NSMAP)
                if pstyle_el is None:
                    continue

                text_el = lvl.find('w:lvlText', NSMAP)
                if text_el is None:
                    continue
                actual = text_el.get(f'{{{W}}}val')
                expected = expected_lvl_texts[ilvl]
                if actual != expected:
                    self.log_fix("编号格式",
                                 f"abstractNum={abs_id} ilvl={ilvl} lvlText \"{actual}\" → \"{expected}\"")
                    if not self.dry_run:
                        text_el.set(f'{{{W}}}val', expected)

    # ========================================
    # 8. 修复标题段落级 numId 覆盖（统一编号链）
    # ========================================
    def fix_heading_numid_override(self):
        """修复标题段落级别的 numId 覆盖问题。

        当 Heading 1 段落级有独立的 numId 覆盖（如 numId=11 → abstractNumId=3），
        而样式定义中的 numId（如 numId=4 → abstractNumId=4）是与子标题共享的编号链，
        段落级覆盖会导致 Heading 1 和子标题（一级节标题2.3、二级节标题2.3.1）
        不在同一编号链中，造成编号不连贯。

        修复方式：移除段落级 numPr，让标题继承样式定义的编号。
        """
        style_rules = self.rules.get('styles', {})
        fix_count = 0

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name or "None") if para.style else "None"
            if style_name.lower().startswith('toc ') or style_name == '说明文字':
                continue

            rules = style_rules.get(style_name)
            if not rules:
                continue
            if rules.get('paragraph', {}).get('outline_level') is None:
                continue

            # 获取样式定义中的编号信息
            style_numId = None
            if para.style and para.style.element is not None:
                s_pPr = para.style.element.find('.//w:pPr', NSMAP)
                if s_pPr is not None:
                    s_numPr = s_pPr.find('w:numPr', NSMAP)
                    if s_numPr is not None:
                        s_numId_el = s_numPr.find('w:numId', NSMAP)
                        if s_numId_el is not None:
                            style_numId = s_numId_el.get(f'{{{W}}}val')

            if not style_numId or style_numId == '0':
                continue

            # 检查段落级 numId 是否与样式定义不同（且不是 numId=0）
            pPr = para._element.find('w:pPr', NSMAP)
            if pPr is not None:
                numPr = pPr.find('w:numPr', NSMAP)
                if numPr is not None:
                    numId_el = numPr.find('w:numId', NSMAP)
                    if numId_el is not None:
                        para_numId = numId_el.get(f'{{{W}}}val')
                        if para_numId != '0' and para_numId != style_numId:
                            self.log_fix("编号统一",
                                         f"段落{i} \"{text[:30]}\" 移除段落级 numId={para_numId} 覆盖"
                                         f"（样式 numId={style_numId}）")
                            if not self.dry_run:
                                pPr.remove(numPr)
                            fix_count += 1

        if fix_count > 0:
            self.log_fix("编号统一", f"共统一 {fix_count} 个标题段落到样式编号链")

    # ========================================
    # 9. 修复共享 abstractNum 导致编号计数器污染
    # ========================================
    def fix_shared_abstract_num(self):
        """修复多个 numId 共享标题 abstractNum 导致编号计数器污染的问题。

        当文档中的普通列表（项目符号、有序列表等）的 numId 也引用了
        标题编号系统的 abstractNumId 时，这些列表段落会推进标题的编号计数器，
        导致 Heading 1 编号跳跃（如 1, 5, 6, 7 而非 1, 2, 3, 4）。

        修复方式：为这些非标题列表创建独立的 abstractNum，
        使它们不再与标题编号系统共享计数器。
        """
        numbering_el, _heading_numId, heading_abs_id = self._get_heading_abstract_num_id()
        if numbering_el is None or heading_abs_id is None:
            return

        # 收集所有标题样式使用的 numId（通过样式定义）
        heading_numIds = set()
        style_rules = self.rules.get('styles', {})
        for style in self.doc.styles:
            rules = style_rules.get(style.name)
            if not rules:
                continue
            if rules.get('paragraph', {}).get('outline_level') is None:
                continue
            s_pPr = style.element.find('.//w:pPr', NSMAP)
            if s_pPr is not None:
                s_numPr = s_pPr.find('w:numPr', NSMAP)
                if s_numPr is not None:
                    s_numId_el = s_numPr.find('w:numId', NSMAP)
                    if s_numId_el is not None:
                        heading_numIds.add(s_numId_el.get(f'{{{W}}}val'))

        # 找所有引用同一 abstractNumId 但不是标题使用的 numId
        polluting_numIds = []
        for num_el in numbering_el.findall('w:num', NSMAP):
            num_id = num_el.get(f'{{{W}}}numId')
            abs_ref = num_el.find('w:abstractNumId', NSMAP)
            if abs_ref is None:
                continue
            abs_id = abs_ref.get(f'{{{W}}}val')
            if abs_id == heading_abs_id and num_id not in heading_numIds:
                polluting_numIds.append((num_id, num_el))

        if not polluting_numIds:
            return

        # 检查这些 numId 是否实际被文档段落使用
        body = self.doc.element.body
        used_polluting = set()
        for p_el in body.findall('w:p', NSMAP):
            pPr = p_el.find('w:pPr', NSMAP)
            if pPr is not None:
                numPr = pPr.find('w:numPr', NSMAP)
                if numPr is not None:
                    ne = numPr.find('w:numId', NSMAP)
                    if ne is not None:
                        nid = ne.get(f'{{{W}}}val')
                        if nid in {pid for pid, _ in polluting_numIds}:
                            used_polluting.add(nid)

        if not used_polluting:
            return

        # 创建独立的简单列表 abstractNum
        max_abs_id = 0
        for abs_num in numbering_el.findall('w:abstractNum', NSMAP):
            aid = int(abs_num.get(f'{{{W}}}abstractNumId'))
            if aid > max_abs_id:
                max_abs_id = aid

        new_abs_id = max_abs_id + 1

        new_abs_xml = (
            f'<w:abstractNum w:abstractNumId="{new_abs_id}" '
            f'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f'<w:multiLevelType w:val="hybridMultilevel"/>'
        )
        lvl_formats = [
            ('decimal', '%1.'), ('lowerLetter', '%2.'), ('lowerRoman', '%3.'),
            ('decimal', '%4.'), ('lowerLetter', '%5.'), ('lowerRoman', '%6.'),
            ('decimal', '%7.'), ('lowerLetter', '%8.'), ('lowerRoman', '%9.'),
        ]
        for i, (fmt, txt) in enumerate(lvl_formats):
            new_abs_xml += (
                f'<w:lvl w:ilvl="{i}">'
                f'<w:start w:val="1"/>'
                f'<w:numFmt w:val="{fmt}"/>'
                f'<w:lvlText w:val="{txt}"/>'
                f'<w:lvlJc w:val="left"/>'
                f'</w:lvl>'
            )
        new_abs_xml += '</w:abstractNum>'

        new_abs_el = etree.fromstring(new_abs_xml)

        # 插入到 numbering.xml 中
        last_abs = numbering_el.findall('w:abstractNum', NSMAP)[-1]
        last_abs.addnext(new_abs_el)

        # 修改污染的 numId 指向新 abstractNum
        for num_id, num_el in polluting_numIds:
            if num_id in used_polluting:
                abs_ref = num_el.find('w:abstractNumId', NSMAP)
                abs_ref.set(f'{{{W}}}val', str(new_abs_id))
                self.log_fix("编号隔离",
                             f"numId={num_id} 从共享 abstractNum={heading_abs_id} "
                             f"改为独立 abstractNum={new_abs_id}")

        self.log_fix("编号隔离",
                     f"共隔离 {len(used_polluting)} 个非标题列表编号，"
                     f"防止计数器污染标题编号")

    # ========================================
    # 10. 修复标题样式 & 去除手动编号
    # ========================================
    def fix_heading_style_and_manual_numbering(self):
        """修复使用错误样式的标题段落，并去除手动编号文本。

        当 Heading 2/3 等样式没有配置 numPr（不参与自动编号），
        用户往往手动输入 "1.1"、"1.2.1" 等编号前缀。
        正确做法是将这些段落改为有 numPr 配置的标题样式
        （如 一级节标题2.3、二级节标题2.3.1），然后去除手动编号文本。

        修复逻辑：
        1. 根据 heading_style_fix.style_replacement 映射替换样式
        2. 根据 heading_style_fix.manual_numbering_patterns 去除手动编号
        3. 保留段落中每个 Run 的字体/格式信息
        """
        fix_rules = self.rules.get('heading_style_fix', {})
        if not fix_rules.get('enabled'):
            return

        style_map = fix_rules.get('style_replacement', {})
        num_patterns = fix_rules.get('manual_numbering_patterns', {})

        if not style_map:
            return

        # 预编译目标样式对象
        target_styles = {}
        for _old_name, new_name in style_map.items():
            for style in self.doc.styles:
                if style.name == new_name:
                    target_styles[new_name] = style
                    break

        fix_count = 0

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            style_name = (para.style.name or "None") if para.style else "None"

            # 跳过 TOC 条目
            if style_name.lower().startswith('toc ') or style_name == '说明文字':
                continue

            # 检查是否需要样式替换
            if style_name not in style_map:
                continue

            new_style_name = style_map[style_name]
            new_style = target_styles.get(new_style_name)
            if new_style is None:
                continue

            # 替换样式
            old_style = style_name
            if not self.dry_run:
                para.style = cast(ParagraphStyle, new_style)

            # 去除手动编号前缀
            pattern = num_patterns.get(new_style_name)
            removed_prefix = ""
            if pattern:
                m = re.match(pattern, text)
                if m:
                    removed_prefix = m.group(0)
                    if not self.dry_run:
                        self._remove_text_prefix(para, len(removed_prefix))

            desc = f"段落{i} \"{text[:40]}\" 样式 {old_style} → {new_style_name}"
            if removed_prefix:
                desc += f"，去除手动编号 \"{removed_prefix.strip()}\""
            self.log_fix("标题样式", desc)
            fix_count += 1

        if fix_count > 0:
            self.log_fix("标题样式", f"共修复 {fix_count} 个标题段落的样式和手动编号")

    def _remove_text_prefix(self, para, char_count):
        """从段落开头移除指定数量的字符，保留 Run 的格式信息。

        逐个 Run 删除字符，当一个 Run 的文本被完全消耗后移至下一个 Run。
        """
        remaining = char_count
        runs_to_remove = []

        for run in para.runs:
            if remaining <= 0:
                break
            rt = run.text
            if len(rt) <= remaining:
                remaining -= len(rt)
                runs_to_remove.append(run)
            else:
                run.text = rt[remaining:]
                remaining = 0

        # 删除已完全消耗的 Run
        for run in runs_to_remove:
            run._element.getparent().remove(run._element)

    # ========================================
    # 11. 修复"说明文字"样式误用
    # ========================================
    def fix_wrong_caption_style(self):
        """将被错误设置为"说明文字"样式的正文段落改为正确的正文样式。

        区分真正的模板说明（红色字体）和被误用的正文段落：
        - 红色字体 → 跳过（属于模板说明，应手动删除）
        - 非红色字体 → 替换为正文样式
        """
        fix_rules = self.rules.get('heading_style_fix', {})
        if not fix_rules.get('enabled'):
            return

        caption_fix = fix_rules.get('caption_style_fix', {})
        if not caption_fix:
            return

        source_style = caption_fix.get('source', '说明文字')
        target_style_name = caption_fix.get('target', '论文正文-首行缩进')

        # 查找目标样式对象
        target_style = None
        for style in self.doc.styles:
            if style.name == target_style_name:
                target_style = style
                break

        if target_style is None:
            return

        fix_count = 0
        for i, para in enumerate(self.doc.paragraphs):
            if not (para.style and para.style.name == source_style and para.text.strip()):
                continue

            # 跳过红色字体的真模板说明
            is_red = False
            for run in para.runs:
                if run.font.color and run.font.color.rgb:
                    rgb = str(run.font.color.rgb)
                    if rgb.upper() in ('FF0000', 'CC0000', 'FF3333'):
                        is_red = True
                        break
            if is_red:
                continue

            text = para.text.strip()
            if not self.dry_run:
                para.style = cast(ParagraphStyle, target_style)
            self.log_fix("说明文字修复",
                         f"段落{i} \"{text[:40]}\" 样式 {source_style} → {target_style_name}")
            fix_count += 1

        if fix_count > 0:
            self.log_fix("说明文字修复", f"共修复 {fix_count} 个误用说明文字样式的段落")

    # ========================================
    # 12. 修复图注样式（Normal/正文 → 图题）
    # ========================================
    def fix_figure_caption_style(self):
        """将以"图X-Y"开头的正文段落改为"图题"样式。

        用户手动输入了"图2-1 xxx"之类的图注文本，但样式用的是 Normal 或
        论文正文-首行缩进。需要改为"图题"样式。由于图注文本是手动编写的，
        改为图题样式后需要在段落级设置 numId=0 禁用自动编号，避免编号重复。
        """
        import re
        fig_pattern = re.compile(r'^图\s*\d+-\d+')

        # 找到图题样式
        fig_style = None
        for style in self.doc.styles:
            if style.name == '图题':
                fig_style = style
                break
        if fig_style is None:
            return

        # 可以被修复的源样式列表
        fixable_sources = {'Normal', '论文正文-首行缩进'}

        fix_count = 0
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else 'None'
            if style_name not in fixable_sources:
                continue
            if not fig_pattern.match(text):
                continue

            self.log_fix("图注样式",
                         f"段落{i} \"{text[:40]}\" 样式 {style_name} → 图题")
            if not self.dry_run:
                para.style = cast(ParagraphStyle, fig_style)
                # 段落级设置 numId=0 禁用自动编号（避免与手动文本重复）
                pPr = para._element.find('w:pPr', NSMAP)
                if pPr is None:
                    pPr = etree.SubElement(para._element, f'{{{W}}}pPr')
                numPr = pPr.find('w:numPr', NSMAP)
                if numPr is None:
                    numPr = etree.SubElement(pPr, f'{{{W}}}numPr')
                numId_el = numPr.find('w:numId', NSMAP)
                if numId_el is None:
                    numId_el = etree.SubElement(numPr, f'{{{W}}}numId')
                numId_el.set(f'{{{W}}}val', '0')
            fix_count += 1

        if fix_count > 0:
            self.log_fix("图注样式", f"共修复 {fix_count} 个图注段落的样式")

    # ========================================
    # 13. 修复标题编号缩进（numbering.xml）
    # ========================================
    def fix_heading_numbering_indent(self):
        """修复 numbering.xml 中标题编号级别的缩进、suff 和 tabs。

        python-docx 保存文档后，标题编号级别的属性可能被重置为错误值：
        - 缩进：left=720, hanging=720 → 应为 left=0, firstLine=0
        - suff：tab（默认）→ 应根据级别使用 tab 或 space
        - tabs：val=num pos=720/1440 → 应根据级别设置正确的 tab stop
        同时修复图/表编号级别的 lvlText 和 suff。
        """
        numbering_el, _heading_numId, abs_id = self._get_heading_abstract_num_id()
        if numbering_el is None or abs_id is None:
            return

        # 从 YAML 规则获取编号定义
        num_rules = self.rules.get('numbering', {}).get('heading_numbering', {})
        levels = num_rules.get('levels', {})

        for abs_num in numbering_el.findall('w:abstractNum', NSMAP):
            if abs_num.get(f'{{{W}}}abstractNumId') != abs_id:
                continue
            for lvl in abs_num.findall('w:lvl', NSMAP):
                ilvl = lvl.get(f'{{{W}}}ilvl')
                ilvl_int = int(ilvl) if ilvl else -1

                lvl_rule = levels.get(ilvl_int, {})

                # --- 修复 suff（编号后缀类型：tab / space / nothing）---
                expected_suff = lvl_rule.get('suff')
                if expected_suff:
                    suff_el = lvl.find('w:suff', NSMAP)
                    actual_suff = suff_el.get(f'{{{W}}}val') if suff_el is not None else 'tab'
                    if actual_suff != expected_suff:
                        self.log_fix("编号后缀",
                                     f"abstractNum={abs_id} ilvl={ilvl} "
                                     f"suff \"{actual_suff}\" → \"{expected_suff}\"")
                        if not self.dry_run:
                            if suff_el is None:
                                suff_el = etree.SubElement(lvl, f'{{{W}}}suff')
                            suff_el.set(f'{{{W}}}val', expected_suff)

                # --- 修复 tabs（tab stop 位置）---
                expected_tab_pos = lvl_rule.get('tab_pos')
                pPr = lvl.find('w:pPr', NSMAP)
                if pPr is None:
                    pPr = etree.SubElement(lvl, f'{{{W}}}pPr')

                if expected_suff == 'space':
                    # suff=space 时，清除 tabs（tab stop 不再需要）
                    tabs = pPr.find('w:tabs', NSMAP)
                    if tabs is not None:
                        old_tabs = []
                        for tab in tabs.findall('w:tab', NSMAP):
                            old_tabs.append(f"val={tab.get(f'{{{W}}}val')} pos={tab.get(f'{{{W}}}pos')}")
                        if old_tabs:
                            self.log_fix("编号制表位",
                                         f"abstractNum={abs_id} ilvl={ilvl} "
                                         f"移除 tabs [{'; '.join(old_tabs)}]（suff=space 不需要制表位）")
                        if not self.dry_run:
                            pPr.remove(tabs)
                elif expected_tab_pos is not None:
                    # suff=tab 时，确保 tab stop 位置正确
                    tabs = pPr.find('w:tabs', NSMAP)
                    expected_pos_str = str(expected_tab_pos)

                    needs_tab_fix = False
                    old_tab_info = 'none'
                    if tabs is not None:
                        tab_els = tabs.findall('w:tab', NSMAP)
                        if len(tab_els) == 1:
                            t = tab_els[0]
                            old_val = t.get(f'{{{W}}}val', '')
                            old_pos = t.get(f'{{{W}}}pos', '')
                            old_tab_info = f"val={old_val} pos={old_pos}"
                            if old_pos != expected_pos_str or old_val == 'num':
                                needs_tab_fix = True
                        else:
                            old_tab_info = f"{len(tab_els)} tabs"
                            needs_tab_fix = True
                    else:
                        needs_tab_fix = True

                    if needs_tab_fix:
                        self.log_fix("编号制表位",
                                     f"abstractNum={abs_id} ilvl={ilvl} "
                                     f"tabs [{old_tab_info}] → [val=left pos={expected_pos_str}]")
                        if not self.dry_run:
                            if tabs is not None:
                                pPr.remove(tabs)
                            tabs = etree.SubElement(pPr, f'{{{W}}}tabs')
                            tab = etree.SubElement(tabs, f'{{{W}}}tab')
                            tab.set(f'{{{W}}}val', 'left')
                            tab.set(f'{{{W}}}pos', expected_pos_str)

                # --- 修复 indent ---
                expected_left = lvl_rule.get('ind_left')
                expected_fl = lvl_rule.get('ind_firstLine')
                if expected_left is not None:
                    ind = pPr.find('w:ind', NSMAP)
                    if ind is None:
                        ind = etree.SubElement(pPr, f'{{{W}}}ind')

                    old_left = ind.get(f'{{{W}}}left', '0')
                    old_hanging = ind.get(f'{{{W}}}hanging', 'N/A')
                    old_fl = ind.get(f'{{{W}}}firstLine', 'N/A')
                    exp_left_str = str(expected_left)
                    exp_fl_str = str(expected_fl) if expected_fl is not None else '0'

                    needs_fix = (old_left != exp_left_str or old_hanging != 'N/A' or old_fl != exp_fl_str)
                    if needs_fix:
                        self.log_fix("编号缩进",
                                     f"abstractNum={abs_id} ilvl={ilvl} "
                                     f"left={old_left},hanging={old_hanging},firstLine={old_fl} "
                                     f"→ left={exp_left_str},firstLine={exp_fl_str}")
                        if not self.dry_run:
                            ind.set(f'{{{W}}}left', exp_left_str)
                            ind.set(f'{{{W}}}firstLine', exp_fl_str)
                            # 移除 hanging（与 firstLine 互斥）
                            if f'{{{W}}}hanging' in ind.attrib:
                                del ind.attrib[f'{{{W}}}hanging']

                # --- 修复 lvlText ---
                expected_text = lvl_rule.get('lvlText')
                if expected_text:
                    text_el = lvl.find('w:lvlText', NSMAP)
                    if text_el is not None:
                        actual = text_el.get(f'{{{W}}}val')
                        if actual != expected_text:
                            self.log_fix("编号格式",
                                         f"abstractNum={abs_id} ilvl={ilvl} "
                                         f"lvlText \"{actual}\" → \"{expected_text}\"")
                            if not self.dry_run:
                                text_el.set(f'{{{W}}}val', expected_text)

    # ========================================
    # 14. 修复标题段落级多余缩进
    # ========================================
    def fix_heading_paragraph_indent(self):
        """移除标题段落级别多余的缩进属性。

        有些标题段落可能被手动设置了段落级 ind（如 left=720），
        导致标题不靠左。对于应该靠左的标题样式，移除段落级 ind。
        """
        style_rules = self.rules.get('styles', {})
        heading_styles = set()
        for name, rules in style_rules.items():
            if rules.get('paragraph', {}).get('outline_level') is not None:
                heading_styles.add(name)

        # 不处理 Heading 1 和非章节标题（它们居中显示，不受 ind 影响）
        # 主要处理一级节标题和二级节标题
        fix_count = 0
        for i, para in enumerate(self.doc.paragraphs):
            style_name = para.style.name if para.style else 'None'
            if style_name not in heading_styles:
                continue

            pPr = para._element.find('w:pPr', NSMAP)
            if pPr is None:
                continue
            ind = pPr.find('w:ind', NSMAP)
            if ind is None:
                continue

            # 检查是否有 left 缩进
            left_val = ind.get(f'{{{W}}}left')
            if left_val and left_val != '0':
                text = para.text.strip()
                self.log_fix("标题缩进",
                             f"段落{i} \"{text[:30]}\" 移除段落级 ind left={left_val}")
                if not self.dry_run:
                    # 移除 left 属性或整个 ind（如果只有 left）
                    if f'{{{W}}}left' in ind.attrib:
                        del ind.attrib[f'{{{W}}}left']
                    # 如果 ind 没有其他属性了，整个移除
                    if len(ind.attrib) == 0:
                        pPr.remove(ind)
                fix_count += 1

        if fix_count > 0:
            self.log_fix("标题缩进", f"共修复 {fix_count} 个标题段落的多余缩进")

    # ========================================
    # 15. 修复异常编号
    # ========================================
    def fix_abnormal_numbering(self):
        """修复标题段落的异常编号"""
        style_rules = self.rules.get('styles', {})
        heading_styles = set()
        for name, rules in style_rules.items():
            if rules.get('paragraph', {}).get('outline_level') is not None:
                heading_styles.add(name)

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name or "None") if para.style else "None"

            if style_name not in heading_styles:
                continue

            # 检查是否有编号
            pPr = para._element.find('w:pPr', NSMAP)
            if pPr is None:
                continue
            numPr = pPr.find('w:numPr', NSMAP)
            if numPr is None:
                continue

            numId_elem = numPr.find('w:numId', NSMAP)
            ilvl_elem = numPr.find('w:ilvl', NSMAP)
            numId = numId_elem.get(f'{{{W}}}val') if numId_elem is not None else None
            ilvl = ilvl_elem.get(f'{{{W}}}val') if ilvl_elem is not None else None

            # 如果 numId 不是 0 且 ilvl 不是 0，可能是异常编号
            if numId and numId != '0' and ilvl and ilvl != '0':
                # 检查同样式的其他段落
                same_style_non_zero = 0
                same_style_total = 0
                for p in self.doc.paragraphs:
                    if p.style and p.style.name == style_name and p.text.strip():
                        same_style_total += 1
                        pp = p._element.find('w:pPr', NSMAP)
                        if pp is not None:
                            np = pp.find('w:numPr', NSMAP)
                            if np is not None:
                                ni = np.find('w:numId', NSMAP)
                                il = np.find('w:ilvl', NSMAP)
                                if ni is not None and il is not None:
                                    if ni.get(f'{{{W}}}val') == numId and il.get(f'{{{W}}}val') == ilvl:
                                        same_style_non_zero += 1

                # 如果少于一半的同样式段落有这个编号配置，认为是异常
                if same_style_non_zero <= same_style_total // 2:
                    self.log_fix("编号修复", f"段落{i} \"{text[:30]}\" 移除异常编号 numId={numId} ilvl={ilvl}")
                    if not self.dry_run:
                        pPr.remove(numPr)

    # ========================================
    # 16. 设置更新域标志（最后执行）
    # ========================================
    def set_update_fields(self):
        """设置文档打开时自动更新域"""
        from docx.oxml.ns import qn
        settings = self.doc.settings.element
        updateFields = settings.find(qn('w:updateFields'))
        if updateFields is None:
            updateFields = etree.SubElement(settings, qn('w:updateFields'))
        updateFields.set(qn('w:val'), 'true')
        self.log_fix("域更新", "设置文档打开时自动更新域")

    # ========================================
    # 运行所有修复
    # ========================================
    def run_all_fixes(self, dry_run=False):
        """执行所有修复"""
        self.dry_run = dry_run
        self.fixes = []

        print(f"\n{'=' * 60}")
        print(f"  格式修复: {os.path.basename(self.filepath)}")
        print(f"  模式: {'DRY RUN (仅检测)' if dry_run else '实际修复'}")
        print(f"{'=' * 60}\n")

        self.fix_heading_style_and_manual_numbering()
        self.fix_wrong_caption_style()
        self.fix_figure_caption_style()
        self.fix_page_setup()
        self.fix_style_definitions()
        self.fix_paragraph_outline_levels()
        self.fix_toc()
        self.fix_run_fonts()
        self.fix_disabled_heading_numbering()
        self.fix_heading_lvl_text()
        self.fix_heading_numid_override()
        self.fix_shared_abstract_num()
        self.fix_heading_numbering_indent()
        self.fix_heading_paragraph_indent()
        self.fix_abnormal_numbering()

        if not dry_run and self.fixes:
            self.set_update_fields()

        # 保存
        if not dry_run and self.fixes:
            bak_path = self.filepath + ".bak"
            if not os.path.exists(bak_path):
                shutil.copy2(self.filepath, bak_path)
                print(f"\n  已备份原文件: {bak_path}")

            self.doc.save(self.filepath)
            print(f"  已保存: {self.filepath}")

        print(f"\n  {'─' * 50}")
        print(f"  共 {len(self.fixes)} 项{'需要' if dry_run else '已'}修复")
        if dry_run and self.fixes:
            print("  去掉 --dry-run 参数执行实际修复")
        print()

        return self.fixes


def main():
    if len(sys.argv) < 2:
        print("用法: python fixer.py <docx文件路径> [--rules <yaml规则文件>] [--dry-run]")
        sys.exit(1)

    filepath = sys.argv[1]
    dry_run = "--dry-run" in sys.argv

    rules_path = None
    if "--rules" in sys.argv:
        idx = sys.argv.index("--rules")
        if idx + 1 < len(sys.argv):
            rules_path = sys.argv[idx + 1]

    fixer = DocxFixer(filepath, rules_path)
    fixer.run_all_fixes(dry_run=dry_run)


if __name__ == "__main__":
    main()
