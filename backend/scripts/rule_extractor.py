#!/usr/bin/env python3
"""
从 Word 模板文档自动提取格式规则，生成 YAML 规则配置文件。

通过解析 .docx 文件的 XML 结构（styles.xml、numbering.xml、document.xml），
自动提取所有格式信息并生成与 checker.py / fixer.py 兼容的 YAML 规则文件。

用法：
    python rule_extractor.py <模板docx文件> [--output <输出yaml路径>] [--name <规则名称>]
"""

import sys
import os
import re
import yaml
from collections import OrderedDict
from docx import Document

NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# 中文字体名 ↔ 系统内部名 映射
FONT_ALIASES = {
    'SimHei': '黑体', '黑体': '黑体',
    'SimSun': '宋体', '宋体': '宋体',
    'STXinwei': '华文新魏', '华文新魏': '华文新魏',
    'STKaiti': '华文楷体', '华文楷体': '华文楷体',
    'STSong': '华文宋体', '华文宋体': '华文宋体',
    'STFangsong': '华文仿宋', '华文仿宋': '华文仿宋',
    'KaiTi': '楷体', '楷体': '楷体',
    'FangSong': '仿宋', '仿宋': '仿宋',
    'Microsoft YaHei': '微软雅黑', '微软雅黑': '微软雅黑',
}

# 半磅值 → 常见中文字号名映射
HALF_PT_TO_CN_SIZE = {
    84: '初号 (42pt)',
    72: '小初 (36pt)',
    52: '一号 (26pt)',
    48: '小一 (24pt)',
    44: '二号 (22pt)',
    36: '小二 (18pt)',
    32: '三号 (16pt)',
    30: '小三 (15pt)',
    28: '四号 (14pt)',
    24: '小四 (12pt)',
    21: '五号 (10.5pt)',
    18: '小五 (9pt)',
    15: '六号 (7.5pt)',
    12: '小六 (6.5pt)',
}

# Word 对齐方式映射
ALIGNMENT_MAP = {
    'left': '左对齐',
    'center': '居中',
    'right': '右对齐',
    'both': '两端对齐',
    'distribute': '分散对齐',
}


# ===== YAML 有序输出支持 =====
class OrderedDumper(yaml.Dumper):
    """保持字典键顺序的 YAML Dumper"""
    pass


def _dict_representer(dumper, data):
    return dumper.represent_mapping(
        yaml.resolver.BaseResolver.DEFAULT_MAPPING_TAG,
        data.items()
    )


OrderedDumper.add_representer(OrderedDict, _dict_representer)


# ===== 颜色输出 =====
class Color:
    RED = '\033[91m'
    GREEN = '\033[92m'
    YELLOW = '\033[93m'
    BLUE = '\033[94m'
    CYAN = '\033[96m'
    BOLD = '\033[1m'
    END = '\033[0m'


class RuleExtractor:
    """从 Word 模板文档提取格式规则"""

    def __init__(self, filepath):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.rules = OrderedDict()

    # ========================================
    # 1. 提取元信息
    # ========================================
    def extract_meta(self, name=None, description=None):
        """生成规则文件的元信息"""
        basename = os.path.splitext(os.path.basename(self.filepath))[0]
        self.rules['meta'] = OrderedDict([
            ('name', name or f'从 {basename} 提取的格式规则'),
            ('version', '1.0'),
            ('description', description or f'从模板文件 "{os.path.basename(self.filepath)}" 自动提取'),
            ('source_file', os.path.basename(self.filepath)),
        ])

    # ========================================
    # 2. 提取页面设置
    # ========================================
    def extract_page_setup(self):
        """从第一个节（section）提取页面设置"""
        if not self.doc.sections:
            return

        section = self.doc.sections[0]
        page_setup = OrderedDict()

        # 纸张大小
        if section.page_width and section.page_height:
            w_cm = round(section.page_width / 360000, 1)
            h_cm = round(section.page_height / 360000, 1)
            if abs(w_cm - 21.0) < 0.3 and abs(h_cm - 29.7) < 0.3:
                page_setup['paper_size'] = 'A4'
            else:
                page_setup['paper_size'] = 'custom'
            page_setup['width_cm'] = w_cm
            page_setup['height_cm'] = h_cm

        # 页边距
        margin_attrs = [
            ('margin_top_cm', 'top_margin'),
            ('margin_bottom_cm', 'bottom_margin'),
            ('margin_left_cm', 'left_margin'),
            ('margin_right_cm', 'right_margin'),
        ]
        for rule_key, attr in margin_attrs:
            val = getattr(section, attr, None)
            if val is not None:
                page_setup[rule_key] = round(val / 360000, 2)

        # 页眉页脚距离
        if section.header_distance is not None:
            page_setup['header_distance_cm'] = round(section.header_distance / 360000, 2)
        if section.footer_distance is not None:
            page_setup['footer_distance_cm'] = round(section.footer_distance / 360000, 2)

        # 装订线
        sect_pr = section._sectPr
        if sect_pr is not None:
            pgMar = sect_pr.find('w:pgMar', NSMAP)
            if pgMar is not None:
                gutter = pgMar.get(f'{{{W}}}gutter')
                if gutter:
                    page_setup['gutter_cm'] = round(int(gutter) / 567, 2)

        # 文档网格
        if sect_pr is not None:
            docGrid = sect_pr.find('w:docGrid', NSMAP)
            if docGrid is not None:
                grid_info = OrderedDict()
                grid_type = docGrid.get(f'{{{W}}}type')
                if grid_type:
                    grid_info['type'] = grid_type
                line_pitch = docGrid.get(f'{{{W}}}linePitch')
                if line_pitch:
                    grid_info['line_pitch'] = int(line_pitch)
                char_space = docGrid.get(f'{{{W}}}charSpace')
                if char_space:
                    grid_info['char_space'] = int(char_space)
                if grid_info:
                    page_setup['doc_grid'] = grid_info

        self.rules['page_setup'] = page_setup

    # ========================================
    # 3. 提取页眉页脚
    # ========================================
    def extract_header_footer(self):
        """提取页眉页脚内容和格式"""
        hf = OrderedDict()

        # 页眉
        for section in self.doc.sections:
            header = section.header
            if header and header.paragraphs:
                for p in header.paragraphs:
                    text = p.text.strip()
                    if text:
                        header_info = OrderedDict()
                        header_info['text'] = text
                        style_name = p.style.name if p.style else None
                        if style_name:
                            header_info['style'] = style_name

                        # 对齐
                        pPr = p._element.find('w:pPr', NSMAP)
                        if pPr is not None:
                            jc = pPr.find('w:jc', NSMAP)
                            if jc is not None:
                                header_info['alignment'] = jc.get(f'{{{W}}}val')

                        # 字号
                        for run in p.runs:
                            if run.font.size:
                                header_info['font_size_pt'] = round(run.font.size / 12700, 1)
                                break

                        hf['header'] = header_info
                        break
                if 'header' in hf:
                    break

        # 页脚
        for section in self.doc.sections:
            footer = section.footer
            if footer and footer.paragraphs:
                for p in footer.paragraphs:
                    footer_info = OrderedDict()
                    style_name = p.style.name if p.style else None
                    if style_name:
                        footer_info['style'] = style_name

                    pPr = p._element.find('w:pPr', NSMAP)
                    if pPr is not None:
                        jc = pPr.find('w:jc', NSMAP)
                        if jc is not None:
                            footer_info['alignment'] = jc.get(f'{{{W}}}val')

                    for run in p.runs:
                        if run.font.size:
                            footer_info['font_size_pt'] = round(run.font.size / 12700, 1)
                            break

                    if footer_info:
                        hf['footer'] = footer_info
                    break
                if 'footer' in hf:
                    break

        if hf:
            self.rules['header_footer'] = hf

    # ========================================
    # 4. 提取样式定义
    # ========================================
    def _extract_style_info(self, style, inherit=False):
        """从样式 XML 中提取完整的格式信息"""
        info = OrderedDict()
        para_info = OrderedDict()
        char_info = OrderedDict()

        elem = style.element
        pPr = elem.find('.//w:pPr', NSMAP)
        rPr = elem.find('.//w:rPr', NSMAP)

        # ---- 段落属性 ----
        if pPr is not None:
            # 对齐方式
            jc = pPr.find('w:jc', NSMAP)
            if jc is not None:
                val = jc.get(f'{{{W}}}val')
                para_info['alignment'] = val

            # 间距
            spacing = pPr.find('w:spacing', NSMAP)
            if spacing is not None:
                before_lines = spacing.get(f'{{{W}}}beforeLines')
                if before_lines:
                    para_info['spacing_before_lines'] = int(before_lines)
                after_lines = spacing.get(f'{{{W}}}afterLines')
                if after_lines:
                    para_info['spacing_after_lines'] = int(after_lines)
                before = spacing.get(f'{{{W}}}before')
                if before and not before_lines:
                    para_info['spacing_before_twips'] = int(before)
                after = spacing.get(f'{{{W}}}after')
                if after and not after_lines:
                    para_info['spacing_after_twips'] = int(after)
                line = spacing.get(f'{{{W}}}line')
                if line:
                    para_info['line_spacing'] = int(line)
                line_rule = spacing.get(f'{{{W}}}lineRule')
                if line_rule:
                    para_info['line_spacing_rule'] = line_rule

            # 缩进
            ind = pPr.find('w:ind', NSMAP)
            if ind is not None:
                first_line_chars = ind.get(f'{{{W}}}firstLineChars')
                if first_line_chars:
                    para_info['first_line_indent_chars'] = int(first_line_chars)
                first_line = ind.get(f'{{{W}}}firstLine')
                if first_line and not first_line_chars:
                    para_info['first_line_indent_twips'] = int(first_line)
                left = ind.get(f'{{{W}}}left')
                if left and left != '0':
                    para_info['indent_left_twips'] = int(left)
                right = ind.get(f'{{{W}}}right')
                if right and right != '0':
                    para_info['indent_right_twips'] = int(right)
                hanging = ind.get(f'{{{W}}}hanging')
                if hanging:
                    para_info['indent_hanging_twips'] = int(hanging)

            # 大纲级别
            outline_lvl = pPr.find('w:outlineLvl', NSMAP)
            if outline_lvl is not None:
                para_info['outline_level'] = int(outline_lvl.get(f'{{{W}}}val'))

            # 与下段同页
            keep_next = pPr.find('w:keepNext', NSMAP)
            if keep_next is not None:
                val = keep_next.get(f'{{{W}}}val', 'true')
                if val not in ('0', 'false'):
                    para_info['keep_next'] = True

            # 段中不分页
            keep_lines = pPr.find('w:keepLines', NSMAP)
            if keep_lines is not None:
                val = keep_lines.get(f'{{{W}}}val', 'true')
                if val not in ('0', 'false'):
                    para_info['keep_lines'] = True

        # ---- 字符属性 ----
        if rPr is not None:
            # 字体
            rFonts = rPr.find('w:rFonts', NSMAP)
            if rFonts is not None:
                for attr, rule_key in [('ascii', 'font_ascii'),
                                        ('eastAsia', 'font_east_asia'),
                                        ('hAnsi', 'font_hAnsi'),
                                        ('cs', 'font_cs')]:
                    v = rFonts.get(f'{{{W}}}{attr}')
                    if v:
                        char_info[rule_key] = v

            # 字号
            sz = rPr.find('w:sz', NSMAP)
            if sz is not None:
                half_pt = int(sz.get(f'{{{W}}}val'))
                char_info['font_size_pt'] = half_pt / 2

            # 字号（cs）
            szCs = rPr.find('w:szCs', NSMAP)
            if szCs is not None:
                half_pt_cs = int(szCs.get(f'{{{W}}}val'))
                char_info['font_size_cs_pt'] = half_pt_cs / 2

            # 粗体
            b = rPr.find('w:b', NSMAP)
            if b is not None:
                val = b.get(f'{{{W}}}val', 'true')
                char_info['bold'] = val not in ('0', 'false')

            # 斜体
            i_el = rPr.find('w:i', NSMAP)
            if i_el is not None:
                val = i_el.get(f'{{{W}}}val', 'true')
                char_info['italic'] = val not in ('0', 'false')

            # 字间距
            spacing_el = rPr.find('w:spacing', NSMAP)
            if spacing_el is not None:
                val = spacing_el.get(f'{{{W}}}val')
                if val:
                    char_info['char_spacing'] = int(val)

            # 字距调整（kern）
            kern = rPr.find('w:kern', NSMAP)
            if kern is not None:
                val = kern.get(f'{{{W}}}val')
                if val:
                    char_info['kern'] = int(val)

            # 颜色
            color = rPr.find('w:color', NSMAP)
            if color is not None:
                val = color.get(f'{{{W}}}val')
                if val and val.upper() != 'auto' and val.upper() != '000000':
                    char_info['color'] = val.upper()

        # 沿 basedOn 链继承
        if inherit:
            parent = self._get_parent_style(style)
            if parent:
                parent_info = self._extract_style_info(parent, inherit=True)
                p_para = parent_info.get('paragraph', {})
                p_char = parent_info.get('character', {})
                for k, v in p_para.items():
                    if k not in para_info:
                        para_info[k] = v
                for k, v in p_char.items():
                    if k not in char_info:
                        char_info[k] = v

        if para_info:
            info['paragraph'] = para_info
        if char_info:
            info['character'] = char_info

        return info

    def _get_parent_style(self, style):
        """获取样式的 basedOn 父样式"""
        based_on = style.element.find('w:basedOn', NSMAP)
        if based_on is None:
            return None
        parent_id = based_on.get(f'{{{W}}}val')
        if not parent_id:
            return None
        for s in self.doc.styles:
            if s.style_id == parent_id:
                return s
        return None

    def _get_style_based_on(self, style):
        """获取 basedOn 样式名"""
        based_on = style.element.find('w:basedOn', NSMAP)
        if based_on is None:
            return None
        parent_id = based_on.get(f'{{{W}}}val')
        if not parent_id:
            return None
        for s in self.doc.styles:
            if s.style_id == parent_id:
                return s.name
        return parent_id

    def extract_styles(self):
        """提取文档中所有有意义的样式定义"""
        styles = OrderedDict()

        # 收集文档中实际使用的样式名
        used_styles = set()
        for para in self.doc.paragraphs:
            if para.style:
                used_styles.add(para.style.name)
        # 也检查页眉页脚
        for section in self.doc.sections:
            for p in section.header.paragraphs:
                if p.style:
                    used_styles.add(p.style.name)
            for p in section.footer.paragraphs:
                if p.style:
                    used_styles.add(p.style.name)

        # 检查每个样式
        for style in self.doc.styles:
            # 只处理段落样式
            if style.type is not None and str(style.type) != 'PARAGRAPH (1)' and style.type != 1:
                try:
                    from docx.enum.style import WD_STYLE_TYPE
                    if style.type != WD_STYLE_TYPE.PARAGRAPH:
                        continue
                except Exception:
                    if 'PARAGRAPH' not in str(style.type):
                        continue

            name = style.name
            if not name:
                continue

            # 跳过无内容的内置样式
            # 但保留常用内置样式（Normal, Heading 1-9 等）和文档中使用的样式
            is_used = name in used_styles
            is_important = name in ('Normal', 'Header', 'Footer', '目录标题') or \
                           name.startswith('Heading ') or \
                           name.lower().startswith('toc ')

            if not is_used and not is_important:
                continue

            info = self._extract_style_info(style, inherit=False)

            # 如果没有任何有效的格式信息，跳过
            if not info.get('paragraph') and not info.get('character'):
                if not is_used:
                    continue

            style_entry = OrderedDict()

            # 添加描述（基于样式名推断）
            desc = self._generate_style_description(name, style)
            if desc:
                style_entry['description'] = desc

            # basedOn
            based_on = self._get_style_based_on(style)
            if based_on and based_on != name:
                style_entry['based_on'] = based_on

            # 检查是否应标记为 should_not_exist（如"说明文字"类红色样式）
            if self._is_instruction_style(style):
                style_entry['should_not_exist'] = True

            # 段落格式
            if info.get('paragraph'):
                style_entry['paragraph'] = info['paragraph']

            # 字符格式
            if info.get('character'):
                style_entry['character'] = info['character']

            if style_entry:
                styles[name] = style_entry

        self.rules['styles'] = styles

    def _is_instruction_style(self, style):
        """判断是否是模板说明文字样式（红色字体）"""
        rPr = style.element.find('.//w:rPr', NSMAP)
        if rPr is not None:
            color = rPr.find('w:color', NSMAP)
            if color is not None:
                val = (color.get(f'{{{W}}}val') or '').upper()
                if val in ('FF0000', 'CC0000', 'FF3333'):
                    return True
        return False

    def _generate_style_description(self, name, style):
        """根据样式名自动生成描述"""
        desc_map = {
            'Normal': '基础正文样式',
            'Header': '页眉样式',
            'Footer': '页脚样式',
            '目录标题': '目录页标题',
        }
        if name in desc_map:
            return desc_map[name]

        # Heading 系列
        m = re.match(r'Heading (\d+)', name)
        if m:
            lvl = int(m.group(1))
            return f'{lvl}级标题'

        # TOC 系列
        m = re.match(r'toc (\d+)', name, re.IGNORECASE)
        if m:
            return f'目录中的{m.group(1)}级条目'

        # 中文样式名通常自解释
        if any('\u4e00' <= c <= '\u9fff' for c in name):
            return f'自定义样式 "{name}"'

        return None

    # ========================================
    # 5. 提取文档结构
    # ========================================
    def extract_structure(self):
        """提取文档结构信息（章节、标题映射、TOC 配置）"""
        structure = OrderedDict()

        # 收集一级标题
        chapters = []
        heading_style_names = {}  # outline_level -> style_name

        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            outline_lvl = self._get_para_outline_level(para)

            if outline_lvl < 9:
                style_name = para.style.name if para.style else 'Normal'
                if outline_lvl not in heading_style_names:
                    heading_style_names[outline_lvl] = style_name

                if outline_lvl == 0:
                    chapters.append(OrderedDict([
                        ('pattern', text[:30]),
                        ('style', style_name),
                    ]))

        # 必要章节
        if chapters:
            structure['required_chapters'] = chapters

        # 标题样式映射
        if heading_style_names:
            mapping = OrderedDict()
            for lvl in sorted(heading_style_names.keys()):
                if lvl <= 5:
                    mapping[f'level_{lvl + 1}'] = heading_style_names[lvl]
            if mapping:
                structure['heading_style_mapping'] = mapping

        # TOC 配置
        toc_info = self._extract_toc_info()
        if toc_info:
            structure['toc'] = toc_info

        if structure:
            self.rules['structure'] = structure

    def _get_para_outline_level(self, para):
        """获取段落的实际大纲级别"""
        # 段落直接设置
        pPr = para._element.find('w:pPr', NSMAP)
        if pPr is not None:
            ol = pPr.find('w:outlineLvl', NSMAP)
            if ol is not None:
                return int(ol.get(f'{{{W}}}val'))

        # 样式定义
        if para.style and para.style.element is not None:
            style_pPr = para.style.element.find('.//w:pPr', NSMAP)
            if style_pPr is not None:
                ol = style_pPr.find('w:outlineLvl', NSMAP)
                if ol is not None:
                    return int(ol.get(f'{{{W}}}val'))

        # Heading 推断
        if para.style and para.style.name:
            m = re.match(r'Heading (\d+)', para.style.name)
            if m:
                return int(m.group(1)) - 1

        return 9

    def _extract_toc_info(self):
        """提取 TOC 域代码配置"""
        toc_info = OrderedDict()

        for para in self.doc.paragraphs:
            for run in para._element.findall('.//w:r', NSMAP):
                instrText = run.find('w:instrText', NSMAP)
                if instrText is not None and instrText.text and 'TOC' in instrText.text:
                    code = instrText.text.strip()

                    # TOC 标题样式
                    # 检查段落前是否有目录标题
                    toc_info['toc_title_style'] = '目录标题'

                    # \o 大纲范围
                    m = re.search(r'\\o\s*"1-(\d+)"', code)
                    if m:
                        toc_info['outline_range'] = f'1-{m.group(1)}'

                    # \t 自定义样式映射
                    m = re.search(r'\\t\s*"([^"]+)"', code)
                    if m:
                        custom_styles = OrderedDict()
                        parts = m.group(1).split(',')
                        for i in range(0, len(parts) - 1, 2):
                            style_name = parts[i].strip()
                            try:
                                level = int(parts[i + 1].strip())
                                custom_styles[style_name] = level
                            except (ValueError, IndexError):
                                pass
                        if custom_styles:
                            toc_info['custom_styles'] = custom_styles

                    return toc_info

        return toc_info

    # ========================================
    # 6. 提取编号定义
    # ========================================
    def extract_numbering(self):
        """提取 numbering.xml 中的编号定义"""
        try:
            numbering_el = self.doc.part.numbering_part._element
        except Exception:
            return

        numbering = OrderedDict()

        # 收集每个样式关联的 numId 和 ilvl
        style_num_map = {}  # style_name -> (numId, ilvl)
        for style in self.doc.styles:
            s_pPr = style.element.find('.//w:pPr', NSMAP)
            if s_pPr is not None:
                numPr = s_pPr.find('w:numPr', NSMAP)
                if numPr is not None:
                    numId_el = numPr.find('w:numId', NSMAP)
                    ilvl_el = numPr.find('w:ilvl', NSMAP)
                    numId = numId_el.get(f'{{{W}}}val') if numId_el is not None else None
                    ilvl = ilvl_el.get(f'{{{W}}}val') if ilvl_el is not None else '0'
                    if numId and numId != '0':
                        style_num_map[style.name] = (numId, ilvl)

        # numId → abstractNumId 映射
        numId_to_absId = {}
        for num_el in numbering_el.findall('w:num', NSMAP):
            nid = num_el.get(f'{{{W}}}numId')
            abs_ref = num_el.find('w:abstractNumId', NSMAP)
            if abs_ref is not None:
                numId_to_absId[nid] = abs_ref.get(f'{{{W}}}val')

        # 按 abstractNumId 分组
        absId_groups = {}  # absId -> [(style_name, numId, ilvl)]
        for sname, (numId, ilvl) in style_num_map.items():
            absId = numId_to_absId.get(numId)
            if absId:
                if absId not in absId_groups:
                    absId_groups[absId] = []
                absId_groups[absId].append((sname, numId, ilvl))

        # 提取每个 abstractNum 的详细信息
        for abs_num in numbering_el.findall('w:abstractNum', NSMAP):
            abs_id = abs_num.get(f'{{{W}}}abstractNumId')

            # 获取多级类型
            multi_type = abs_num.find('w:multiLevelType', NSMAP)
            if multi_type is not None:
                multi_val = multi_type.get(f'{{{W}}}val')
            else:
                multi_val = 'singleLevel'

            # 判断是否是标题编号（包含 Heading 样式绑定的）
            associated_styles = absId_groups.get(abs_id, [])
            has_heading = any('Heading' in s[0] for s in associated_styles)

            # 提取每个级别
            levels_info = OrderedDict()
            for lvl in abs_num.findall('w:lvl', NSMAP):
                ilvl = lvl.get(f'{{{W}}}ilvl')
                ilvl_int = int(ilvl) if ilvl else 0

                lvl_info = OrderedDict()

                # numFmt
                numFmt = lvl.find('w:numFmt', NSMAP)
                if numFmt is not None:
                    lvl_info['numFmt'] = numFmt.get(f'{{{W}}}val')

                # lvlText
                lvlText = lvl.find('w:lvlText', NSMAP)
                if lvlText is not None:
                    lvl_info['lvlText'] = lvlText.get(f'{{{W}}}val')

                # pStyle
                pStyle = lvl.find('w:pStyle', NSMAP)
                if pStyle is not None:
                    psval = pStyle.get(f'{{{W}}}val')
                    # 尝试从 styleId 找回样式名
                    for s in self.doc.styles:
                        if s.style_id == psval:
                            lvl_info['pStyle'] = s.name
                            break
                    else:
                        lvl_info['pStyle'] = psval

                # suff
                suff = lvl.find('w:suff', NSMAP)
                if suff is not None:
                    lvl_info['suff'] = suff.get(f'{{{W}}}val')
                else:
                    lvl_info['suff'] = 'tab'  # 默认值

                # start
                start = lvl.find('w:start', NSMAP)
                if start is not None:
                    start_val = int(start.get(f'{{{W}}}val'))
                    if start_val != 1:
                        lvl_info['start'] = start_val

                # tabs 和 indent
                pPr = lvl.find('w:pPr', NSMAP)
                if pPr is not None:
                    tabs = pPr.find('w:tabs', NSMAP)
                    if tabs is not None:
                        tab_els = tabs.findall('w:tab', NSMAP)
                        if len(tab_els) == 1:
                            pos = tab_els[0].get(f'{{{W}}}pos')
                            if pos:
                                lvl_info['tab_pos'] = int(pos)

                    ind = pPr.find('w:ind', NSMAP)
                    if ind is not None:
                        left = ind.get(f'{{{W}}}left')
                        if left:
                            lvl_info['ind_left'] = int(left)
                        fl = ind.get(f'{{{W}}}firstLine')
                        if fl:
                            lvl_info['ind_firstLine'] = int(fl)
                        hanging = ind.get(f'{{{W}}}hanging')
                        if hanging:
                            lvl_info['ind_hanging'] = int(hanging)

                levels_info[ilvl_int] = lvl_info

            # 只保留有实际内容的编号定义
            if not levels_info:
                continue

            # 生成编号定义名称
            if has_heading:
                num_key = 'heading_numbering'
                desc = '标题使用多级编号'
                if any('图' in (lv.get('lvlText', '') or '') for lv in levels_info.values()):
                    desc += '，章节标题和图表题注共享同一编号链'
            else:
                # 根据关联样式命名
                style_names = [s[0] for s in associated_styles]
                if style_names:
                    num_key = f'{style_names[0]}_numbering'.replace(' ', '_')
                    desc = f'样式 {", ".join(style_names)} 的编号'
                else:
                    # 尝试从 pStyle 获取
                    p_styles = [lv.get('pStyle', '') for lv in levels_info.values() if lv.get('pStyle')]
                    if p_styles:
                        num_key = f'{p_styles[0]}_numbering'.replace(' ', '_')
                        desc = f'样式 {", ".join(p_styles)} 的编号'
                    else:
                        num_key = f'numbering_{abs_id}'
                        desc = f'编号定义 abstractNumId={abs_id}'

            num_entry = OrderedDict([
                ('description', desc),
                ('type', 'multilevel' if multi_val in ('multilevel', 'hybridMultilevel') else 'singleLevel'),
                ('levels', levels_info),
            ])

            numbering[num_key] = num_entry

        if numbering:
            self.rules['numbering'] = numbering

    # ========================================
    # 7. 提取特殊检查规则
    # ========================================
    def extract_special_checks(self):
        """生成特殊检查规则（基于文档内容自动推断）"""
        checks = OrderedDict()

        # 检查是否有"说明文字"样式 → 启用模板说明检查
        has_instruction_style = False
        instruction_style_name = None
        for style in self.doc.styles:
            if self._is_instruction_style(style):
                has_instruction_style = True
                instruction_style_name = style.name
                break

        checks['template_instructions_check'] = OrderedDict([
            ('enabled', has_instruction_style),
            ('description', '检查是否存在未删除的模板说明文字'),
        ])
        if has_instruction_style:
            checks['template_instructions_check']['style_name'] = instruction_style_name
            checks['template_instructions_check']['color'] = 'FF0000'

        # 正文字体一致性
        body_styles = []
        body_cn_font = None
        body_en_font = None
        styles_section = self.rules.get('styles', {})
        for sname in ('论文正文-首行缩进', 'Normal'):
            if sname in styles_section:
                body_styles.append(sname)
                char = styles_section[sname].get('character', {})
                if not body_cn_font:
                    body_cn_font = char.get('font_east_asia')
                if not body_en_font:
                    body_en_font = char.get('font_ascii')

        checks['body_font_consistency'] = OrderedDict([
            ('enabled', bool(body_styles and (body_cn_font or body_en_font))),
            ('description', f'检查正文中文使用{body_cn_font or "未指定"}、英文使用{body_en_font or "未指定"}'),
        ])
        if body_styles:
            checks['body_font_consistency']['target_styles'] = body_styles
            if body_cn_font:
                checks['body_font_consistency']['expected_chinese_font'] = body_cn_font
            if body_en_font:
                checks['body_font_consistency']['expected_english_font'] = body_en_font

        # 标题字体一致性
        heading_styles = []
        heading_cn_font = None
        heading_en_font = None
        for sname, sinfo in styles_section.items():
            para = sinfo.get('paragraph', {})
            if para.get('outline_level') is not None:
                heading_styles.append(sname)
                char = sinfo.get('character', {})
                if not heading_cn_font:
                    heading_cn_font = char.get('font_east_asia')
                if not heading_en_font:
                    heading_en_font = char.get('font_ascii')

        checks['heading_font_consistency'] = OrderedDict([
            ('enabled', bool(heading_styles and (heading_cn_font or heading_en_font))),
            ('description', f'检查标题中文使用{heading_cn_font or "未指定"}、英文使用{heading_en_font or "未指定"}'),
        ])
        if heading_styles:
            checks['heading_font_consistency']['target_styles'] = heading_styles
            if heading_cn_font:
                checks['heading_font_consistency']['expected_chinese_font'] = heading_cn_font
            if heading_en_font:
                checks['heading_font_consistency']['expected_english_font'] = heading_en_font

        # 图表标题检查
        fig_style = '图题' if '图题' in styles_section else None
        tbl_style = '表题注' if '表题注' in styles_section else None
        checks['figure_table_caption'] = OrderedDict([
            ('enabled', bool(fig_style or tbl_style)),
            ('description', '检查图片是否有对应的图题'),
        ])
        if fig_style:
            checks['figure_table_caption']['figure_caption_style'] = fig_style
        if tbl_style:
            checks['figure_table_caption']['table_caption_style'] = tbl_style

        self.rules['special_checks'] = checks

    # ========================================
    # 8. 提取标题样式修复规则
    # ========================================
    def extract_heading_style_fix(self):
        """生成标题样式修复规则（自动推断错误样式 → 正确样式的映射）"""
        styles_section = self.rules.get('styles', {})
        structure = self.rules.get('structure', {})
        heading_map = structure.get('heading_style_mapping', {})

        if not heading_map:
            self.rules['heading_style_fix'] = OrderedDict([('enabled', False)])
            return

        fix = OrderedDict()
        fix['enabled'] = True
        fix['description'] = '修复使用错误样式的标题段落，替换为正确样式并去除手动编号'

        # 样式替换映射
        style_replacement = OrderedDict()
        for level_key, correct_style in heading_map.items():
            m = re.match(r'level_(\d+)', level_key)
            if m:
                lvl = int(m.group(1))
                builtin_name = f'Heading {lvl}'
                # 如果正确样式不是 Heading N，则 Heading N 是错误的
                if correct_style != builtin_name:
                    style_replacement[builtin_name] = correct_style

        if style_replacement:
            fix['style_replacement'] = style_replacement

        # 说明文字修复
        if '说明文字' in styles_section:
            # 找一个合适的正文样式作为目标
            target = '论文正文-首行缩进' if '论文正文-首行缩进' in styles_section else 'Normal'
            fix['caption_style_fix'] = OrderedDict([
                ('source', '说明文字'),
                ('target', target),
            ])

        # 手动编号模式
        manual_patterns = OrderedDict()
        for level_key, correct_style in heading_map.items():
            m = re.match(r'level_(\d+)', level_key)
            if not m:
                continue
            lvl = int(m.group(1))
            if lvl == 1:
                continue  # 一级标题通常不需要匹配
            # 生成编号模式
            pattern_parts = r'\d+' + (r'\.\d+' * (lvl - 1))
            manual_patterns[correct_style] = f'^{pattern_parts}\\s*'

        if manual_patterns:
            fix['manual_numbering_patterns'] = manual_patterns

        self.rules['heading_style_fix'] = fix

    # ========================================
    # 汇总提取
    # ========================================
    def extract_all(self, name=None, description=None):
        """执行所有提取步骤"""
        print(f"\n{Color.BOLD}{'=' * 60}{Color.END}")
        print(f"  {Color.CYAN}从模板提取格式规则{Color.END}")
        print(f"  源文件: {os.path.basename(self.filepath)}")
        print(f"{Color.BOLD}{'=' * 60}{Color.END}\n")

        steps = [
            ('元信息', lambda: self.extract_meta(name, description)),
            ('页面设置', self.extract_page_setup),
            ('页眉页脚', self.extract_header_footer),
            ('样式定义', self.extract_styles),
            ('文档结构', self.extract_structure),
            ('编号定义', self.extract_numbering),
            ('特殊检查规则', self.extract_special_checks),
            ('标题修复规则', self.extract_heading_style_fix),
        ]

        for step_name, func in steps:
            try:
                func()
                print(f"  {Color.GREEN}✓{Color.END} {step_name}")
            except Exception as e:
                print(f"  {Color.YELLOW}⚠{Color.END} {step_name}: {e}")

        return self.rules

    def save_yaml(self, output_path):
        """将提取的规则保存为 YAML 文件"""
        # 自定义注释和分节
        sections = [
            ('meta', '元信息'),
            ('page_setup', '一、页面设置'),
            ('header_footer', '二、页眉页脚'),
            ('styles', '三、样式定义规则\n# 每个样式包含：段落格式 + 字符格式'),
            ('structure', '四、文档结构规则'),
            ('numbering', '五、编号定义规则'),
            ('special_checks', '六、特殊检查规则'),
            ('heading_style_fix', '七、标题样式自动修复规则'),
        ]

        lines = []
        lines.append(f'# {"=" * 60}')
        lines.append(f'# {self.rules.get("meta", {}).get("name", "格式规则")}')
        lines.append(f'# {self.rules.get("meta", {}).get("description", "")}')
        lines.append(f'# {"=" * 60}')
        lines.append('')

        for key, comment in sections:
            if key not in self.rules:
                continue

            lines.append(f'# {"=" * 28}')
            lines.append(f'# {comment}')
            lines.append(f'# {"=" * 28}')

            # 使用 YAML dump 该段
            section_data = OrderedDict([(key, self.rules[key])])
            yaml_str = yaml.dump(
                dict(section_data),
                Dumper=OrderedDumper,
                default_flow_style=False,
                allow_unicode=True,
                width=120,
                sort_keys=False,
            )
            lines.append(yaml_str)

        content = '\n'.join(lines)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)

        return output_path

    def print_summary(self):
        """打印提取结果摘要"""
        print(f"\n{Color.BOLD}{'─' * 60}{Color.END}")
        print(f"  {Color.CYAN}提取结果摘要{Color.END}")
        print(f"{Color.BOLD}{'─' * 60}{Color.END}")

        ps = self.rules.get('page_setup', {})
        if ps:
            print(f"\n  {Color.BOLD}页面设置:{Color.END}")
            print(f"    纸张: {ps.get('paper_size', '?')} ({ps.get('width_cm', '?')}×{ps.get('height_cm', '?')}cm)")
            print(f"    页边距: 上{ps.get('margin_top_cm', '?')} 下{ps.get('margin_bottom_cm', '?')} "
                  f"左{ps.get('margin_left_cm', '?')} 右{ps.get('margin_right_cm', '?')}cm")

        styles = self.rules.get('styles', {})
        if styles:
            heading_styles = []
            body_styles = []
            other_styles = []
            for name, info in styles.items():
                para = info.get('paragraph', {})
                if para.get('outline_level') is not None:
                    heading_styles.append(name)
                elif name in ('Normal', '论文正文-首行缩进', '英文正文'):
                    body_styles.append(name)
                else:
                    other_styles.append(name)

            print(f"\n  {Color.BOLD}样式定义:{Color.END} 共 {len(styles)} 个")
            if heading_styles:
                print(f"    标题样式: {', '.join(heading_styles)}")
            if body_styles:
                print(f"    正文样式: {', '.join(body_styles)}")
            if other_styles:
                print(f"    其他样式: {', '.join(other_styles[:10])}"
                      f"{'...' if len(other_styles) > 10 else ''}")

        structure = self.rules.get('structure', {})
        chapters = structure.get('required_chapters', [])
        if chapters:
            print(f"\n  {Color.BOLD}文档结构:{Color.END}")
            print(f"    章节: {len(chapters)} 个一级标题")
            for ch in chapters[:5]:
                print(f"      · {ch.get('pattern', '?')}")

        numbering = self.rules.get('numbering', {})
        if numbering:
            print(f"\n  {Color.BOLD}编号定义:{Color.END} {len(numbering)} 组")
            for nk, nv in numbering.items():
                lvl_count = len(nv.get('levels', {}))
                print(f"    {nk}: {nv.get('type', '?')} ({lvl_count} 级)")

        print()


def main():
    if len(sys.argv) < 2:
        print("用法: python rule_extractor.py <模板docx文件> [--output <yaml路径>] [--name <规则名称>]")
        print()
        print("示例:")
        print("  python rule_extractor.py template.docx")
        print("  python rule_extractor.py template.docx --output rules/my_rules.yaml")
        print("  python rule_extractor.py template.docx --name '我的学校论文格式规则'")
        sys.exit(1)

    filepath = sys.argv[1]

    if not os.path.exists(filepath):
        print(f"❌ 文件不存在: {filepath}")
        sys.exit(1)

    # 解析参数
    output_path = None
    if '--output' in sys.argv:
        idx = sys.argv.index('--output')
        if idx + 1 < len(sys.argv):
            output_path = sys.argv[idx + 1]

    name = None
    if '--name' in sys.argv:
        idx = sys.argv.index('--name')
        if idx + 1 < len(sys.argv):
            name = sys.argv[idx + 1]

    # 默认输出路径
    if output_path is None:
        basename = os.path.splitext(os.path.basename(filepath))[0]
        # 清理文件名
        safe_name = re.sub(r'[^\w\u4e00-\u9fff-]', '_', basename)
        output_path = os.path.join(
            os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
            'rules',
            f'{safe_name}.yaml'
        )

    # 提取规则
    extractor = RuleExtractor(filepath)
    extractor.extract_all(name=name)
    extractor.print_summary()

    # 保存
    extractor.save_yaml(output_path)
    print(f"  {Color.GREEN}✓ 规则已保存到:{Color.END} {output_path}")
    print("\n  使用方法:")
    print(f"    python docx_fixer.py check <文档.docx> --rules {output_path}")
    print(f"    python docx_fixer.py fix-format <文档.docx> --rules {output_path}")
    print()


if __name__ == "__main__":
    main()
