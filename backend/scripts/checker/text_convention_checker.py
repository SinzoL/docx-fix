"""
通用文本排版习惯检查器

检查文档文本内容层面的排版规范问题：
- 确定性检查（高置信度，可直接报告）：
  1. 括号不对称
  2. 引号不匹配
  3. 连续标点
  4. 中文之间多余空格
  5. 连续多个空格
  6. 行首/行尾空格
  7. 全角空格混入

- 争议候选（需 LLM 二次审查）：
  8. 中英文间距不一致
  9. 全半角标点混用
  10. 句末标点缺失

架构说明：
  本模块作为 checker 子包中的独立模块，在 run_all_checks() 中注册调用。
  检查结果统一使用 CheckResult / add_result() 基础设施。
  所有检查规则由 YAML text_conventions section 驱动。
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from typing import Generator

from docx.document import Document

from scripts.checker.base import CheckResult

# ============================================================
# 常量
# ============================================================

# CJK 字符正则（覆盖主要中日韩统一表意文字区段）
_CJK_RE = re.compile(
    r'[\u4e00-\u9fff'       # CJK Unified Ideographs
    r'\u3400-\u4dbf'        # CJK Unified Ideographs Extension A
    r'\u2e80-\u2eff'        # CJK Radicals Supplement
    r'\u3000-\u303f'        # CJK Symbols and Punctuation
    r'\uf900-\ufaff'        # CJK Compatibility Ideographs
    r'\ufe30-\ufe4f'        # CJK Compatibility Forms
    r'\U00020000-\U0002a6df'  # CJK Unified Ideographs Extension B
    r']'
)

# CJK 基本表意文字（用于占比判断，不含符号）
_CJK_IDEO_RE = re.compile(
    r'[\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff\U00020000-\U0002a6df]'
)

# 中文标点
_CN_PUNCT = set('，。、；：！？""''（）【】《》〈〉『』「」〔〕…—～·')

# 英文标点（在中文语境中可能是混用）
_EN_PUNCT_IN_CN = set(',.;:!?()')

# 括号配对
_BRACKET_PAIRS = {
    '（': '）', '）': '（',
    '(': ')', ')': '(',
    '【': '】', '】': '【',
    '[': ']', ']': '[',
    '｛': '｝', '｝': '｛',
    '{': '}', '}': '{',
    '《': '》', '》': '《',
    '〈': '〉', '〉': '〈',
}

_OPEN_BRACKETS = {'（', '(', '【', '[', '｛', '{', '《', '〈'}
_CLOSE_BRACKETS = {'）', ')', '】', ']', '｝', '}', '》', '〉'}

# 引号配对
_OPEN_QUOTES = {'\u201c', '\u2018'}   # " '
_CLOSE_QUOTES = {'\u201d', '\u2019'}  # " '

# 需要报告的连续标点（排除 ！！、？？、……）
_DUPLICATE_PUNCT_RE = re.compile(
    r'([。，、；：])\1+'  # 中文标点连续
    r'|'
    r'(?<!\.)\.\.(?!\.)'  # 英文两个连续点（非省略号上下文）
)

# 中文之间多余空格：CJK + 空格 + CJK
_CJK_SPACE_CJK_RE = re.compile(
    r'([\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff])'
    r'(\s+)'
    r'([\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff])'
)

# 连续多个空格（2个以上连续半角空格）
_MULTI_SPACE_RE = re.compile(r'[ ]{2,}')

# 全角空格
_FULLWIDTH_SPACE = '\u3000'

# URL / 邮箱跳过模式
_URL_RE = re.compile(r'https?://\S+|www\.\S+|\S+@\S+\.\S+', re.IGNORECASE)

# 代码样式名关键字
_CODE_STYLE_KEYWORDS = {'code', '代码'}

NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


# ============================================================
# 数据类
# ============================================================

@dataclass
class ParagraphInfo:
    """段落遍历信息"""
    paragraph: object  # docx Paragraph 对象
    index: int         # 段落序号（全局从 0 开始）
    source: str        # "body" | "table" | "footnote" | "endnote"
    text: str          # 所有 Run 文本拼接
    style_name: str    # 样式名
    has_xml_indent: bool = False  # 是否有 XML 首行缩进


@dataclass
class TextIssue:
    """文本检查发现的问题"""
    rule: str               # 规则名（对应 YAML text_conventions 的 key）
    category: str           # 分类："通用·标点" / "通用·空格" / "通用·全半角"
    item: str               # 检查项名称
    status: str             # "FAIL" / "WARN"
    message: str            # 描述信息
    paragraph_index: int    # 段落序号
    paragraph_source: str   # 来源（body/table/footnote/endnote）
    char_offset: int = 0    # 字符偏移
    context: str = ""       # 上下文片段
    fixable: bool = False   # 是否可自动修复
    is_disputed: bool = False  # 是否为争议项（需 LLM 审查）


@dataclass
class DocumentStats:
    """文档级统计数据（用于 LLM 审查上下文）"""
    total_paragraphs: int = 0
    cjk_spaced_count: int = 0    # 中英交界有空格的数量
    cjk_unspaced_count: int = 0  # 中英交界无空格的数量


# ============================================================
# 段落遍历器
# ============================================================

def iter_all_paragraphs(doc: Document) -> Generator[ParagraphInfo, None, None]:
    """遍历文档所有段落（主体 + 表格 + 脚注 + 尾注），附加来源标记。

    遍历顺序：
    1. doc.paragraphs（主体）
    2. doc.tables → cell.paragraphs（表格）
    3. 脚注/尾注段落

    跳过：
    - 页眉/页脚段落
    - 空段落
    """
    idx = 0

    # 1. 主体段落
    for para in doc.paragraphs:
        text = para.text
        if not text or not text.strip():
            idx += 1
            continue
        style_name = (para.style.name or "Normal") if para.style else "Normal"
        has_indent = _has_xml_first_line_indent(para)
        yield ParagraphInfo(
            paragraph=para, index=idx, source="body",
            text=text, style_name=style_name, has_xml_indent=has_indent,
        )
        idx += 1

    # 2. 表格段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    if not text or not text.strip():
                        idx += 1
                        continue
                    style_name = (para.style.name or "Normal") if para.style else "Normal"
                    yield ParagraphInfo(
                        paragraph=para, index=idx, source="table",
                        text=text, style_name=style_name,
                    )
                    idx += 1

    # 3. 脚注
    try:
        footnotes_part = doc.part.footnotes_part
        if footnotes_part is not None:
            fn_el = footnotes_part._element
            for footnote in fn_el.findall(f'{{{W}}}footnote'):
                fn_type = footnote.get(f'{{{W}}}type')
                if fn_type in ('separator', 'continuationSeparator'):
                    continue
                for p_el in footnote.findall(f'{{{W}}}p'):
                    text = _extract_text_from_element(p_el)
                    if not text or not text.strip():
                        idx += 1
                        continue
                    yield ParagraphInfo(
                        paragraph=None, index=idx, source="footnote",
                        text=text, style_name="Footnote",
                    )
                    idx += 1
    except (AttributeError, Exception):
        pass

    # 4. 尾注
    try:
        endnotes_part = doc.part.endnotes_part
        if endnotes_part is not None:
            en_el = endnotes_part._element
            for endnote in en_el.findall(f'{{{W}}}endnote'):
                en_type = endnote.get(f'{{{W}}}type')
                if en_type in ('separator', 'continuationSeparator'):
                    continue
                for p_el in endnote.findall(f'{{{W}}}p'):
                    text = _extract_text_from_element(p_el)
                    if not text or not text.strip():
                        idx += 1
                        continue
                    yield ParagraphInfo(
                        paragraph=None, index=idx, source="endnote",
                        text=text, style_name="Endnote",
                    )
                    idx += 1
    except (AttributeError, Exception):
        pass


def _extract_text_from_element(p_el) -> str:
    """从 XML 段落元素中提取文本"""
    texts = []
    for r in p_el.findall(f'{{{W}}}r'):
        for t in r.findall(f'{{{W}}}t'):
            if t.text:
                texts.append(t.text)
    return ''.join(texts)


def _has_xml_first_line_indent(para) -> bool:
    """检查段落是否有 XML 首行缩进属性"""
    try:
        pPr = para._element.find('w:pPr', NSMAP)
        if pPr is not None:
            ind = pPr.find('w:ind', NSMAP)
            if ind is not None:
                fi = ind.get(f'{{{W}}}firstLine')
                fic = ind.get(f'{{{W}}}firstLineChars')
                if (fi and fi != '0') or (fic and fic != '0'):
                    return True
    except Exception:
        pass
    return False


# ============================================================
# 辅助函数
# ============================================================

def _cjk_ratio(text: str) -> float:
    """计算段落中 CJK 表意文字的占比"""
    if not text:
        return 0.0
    total = len(text.strip())
    if total == 0:
        return 0.0
    cjk_count = len(_CJK_IDEO_RE.findall(text))
    return cjk_count / total


def _is_code_style(style_name: str) -> bool:
    """判断是否为代码样式"""
    name_lower = style_name.lower()
    return any(kw in name_lower for kw in _CODE_STYLE_KEYWORDS)


def _context_snippet(text: str, pos: int, width: int = 10) -> str:
    """提取上下文片段"""
    start = max(0, pos - width)
    end = min(len(text), pos + width)
    prefix = "..." if start > 0 else ""
    suffix = "..." if end < len(text) else ""
    return f"{prefix}{text[start:end]}{suffix}"


def _mask_urls(text: str) -> str:
    """将 URL 和邮箱替换为占位符，避免误判"""
    return _URL_RE.sub(lambda m: ' ' * len(m.group()), text)


def _source_label(source: str) -> str:
    """来源标记翻译"""
    labels = {"body": "主体", "table": "表格", "footnote": "脚注", "endnote": "尾注"}
    return labels.get(source, source)


def _location_str(para_info: ParagraphInfo, char_offset: int = -1) -> str:
    """构建位置字符串"""
    loc = f"段落{para_info.index + 1} [{_source_label(para_info.source)}]"
    if char_offset >= 0:
        loc += f", 第{char_offset + 1}字符"
    return loc


# ============================================================
# 核心检查函数
# ============================================================

def run_text_convention_checks(
    checker,
    doc: Document,
    rules: dict,
) -> tuple[list[TextIssue], DocumentStats]:
    """执行所有文本排版习惯检查。

    Args:
        checker: DocxChecker 实例（用于 add_result）
        doc: python-docx Document 对象
        rules: YAML 规则字典（完整的 rules）

    Returns:
        (issues, stats) — 所有检查问题列表 + 文档统计
    """
    tc_rules = rules.get('text_conventions', {})
    if not tc_rules:
        # 没有 text_conventions 配置，使用默认（全部启用）
        tc_rules = _default_text_conventions()

    all_issues: list[TextIssue] = []
    stats = DocumentStats()

    # 收集所有段落
    paragraphs = list(iter_all_paragraphs(doc))
    stats.total_paragraphs = len(paragraphs)

    # 遍历段落执行检查
    for i, para_info in enumerate(paragraphs):
        text = para_info.text
        style_name = para_info.style_name
        is_code = _is_code_style(style_name)
        cjk_r = _cjk_ratio(text)
        is_cjk_para = cjk_r >= 0.1  # 段落级 CJK 占比 >= 10%

        # 遮罩 URL/邮箱
        masked_text = _mask_urls(text)

        # 检查是否有 OMath 元素（数学公式段落跳过）
        if para_info.paragraph is not None:
            try:
                omath = para_info.paragraph._element.findall(
                    './/{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath'
                )
                if omath:
                    continue  # 跳过公式段落
            except Exception:
                pass

        # === 确定性检查 ===

        # 1. 括号不对称
        if tc_rules.get('bracket_mismatch', {}).get('enabled', True):
            next_text = paragraphs[i + 1].text[:5] if i + 1 < len(paragraphs) else ""
            issues = _check_bracket_mismatch(para_info, masked_text, next_text)
            all_issues.extend(issues)

        # 2. 引号不匹配
        if tc_rules.get('quote_mismatch', {}).get('enabled', True):
            issues = _check_quote_mismatch(para_info, masked_text)
            all_issues.extend(issues)

        # 3. 连续标点
        if tc_rules.get('duplicate_punctuation', {}).get('enabled', True):
            issues = _check_duplicate_punctuation(para_info, masked_text)
            all_issues.extend(issues)

        # 以下检查跳过代码样式段落
        if not is_code:
            # 4. 中文之间多余空格（仅 CJK 段落）
            if is_cjk_para and tc_rules.get('extra_spaces_in_chinese', {}).get('enabled', True):
                issues = _check_extra_spaces_in_chinese(para_info, masked_text)
                all_issues.extend(issues)

            # 5. 连续多个空格
            if tc_rules.get('consecutive_spaces', {}).get('enabled', True):
                issues = _check_consecutive_spaces(para_info, masked_text)
                all_issues.extend(issues)

            # 6. 行首/行尾空格
            if tc_rules.get('leading_trailing_spaces', {}).get('enabled', True):
                issues = _check_leading_trailing_spaces(para_info, text)
                all_issues.extend(issues)

            # 7. 全角空格混入
            if tc_rules.get('fullwidth_space', {}).get('enabled', True):
                issues = _check_fullwidth_space(para_info, text)
                all_issues.extend(issues)

        # === 争议候选检查 ===

        if not is_code and is_cjk_para:
            # 8. 中英文间距统计
            if tc_rules.get('cjk_english_spacing', {}).get('enabled', True):
                s, u = _count_cjk_english_spacing(masked_text)
                stats.cjk_spaced_count += s
                stats.cjk_unspaced_count += u

            # 9. 全半角标点混用
            if tc_rules.get('fullwidth_halfwidth_punctuation', {}).get('enabled', True):
                issues = _check_halfwidth_punctuation_in_chinese(para_info, masked_text)
                all_issues.extend(issues)

            # 10. 句末标点缺失
            if tc_rules.get('sentence_ending_punctuation', {}).get('enabled', True):
                issues = _check_sentence_ending(para_info, text)
                all_issues.extend(issues)

    # 争议候选：中英文间距不一致（文档级判断）
    if tc_rules.get('cjk_english_spacing', {}).get('enabled', True):
        total = stats.cjk_spaced_count + stats.cjk_unspaced_count
        if total > 0:
            # 如果有不一致（两种风格都存在且少数派占比 > 5%），生成争议项
            minority = min(stats.cjk_spaced_count, stats.cjk_unspaced_count)
            if minority > 0 and minority / total > 0.05:
                # 为每个少数派位置生成争议项（由前端在 AI 审查时处理）
                # 这里只标记统计数据，具体位置在第二遍扫描中标记
                _mark_cjk_spacing_disputes(paragraphs, all_issues, stats, tc_rules)

    # 将确定性问题注册到 checker
    for issue in all_issues:
        if not issue.is_disputed:
            checker.add_result(
                category=issue.category,
                item=issue.item,
                status=issue.status,
                message=issue.message,
                location=_location_str(
                    ParagraphInfo(None, issue.paragraph_index, issue.paragraph_source,
                                  "", "", False),
                    issue.char_offset if issue.char_offset > 0 else -1
                ),
                fixable=issue.fixable,
            )

    return all_issues, stats


# ============================================================
# 个别检查实现
# ============================================================

def _check_bracket_mismatch(
    para_info: ParagraphInfo,
    text: str,
    next_para_start: str,
) -> list[TextIssue]:
    """检查括号不对称（段落级 + 相邻段落宽松匹配）"""
    issues = []
    stack: list[tuple[str, int]] = []  # (bracket_char, position)

    for i, ch in enumerate(text):
        if ch in _OPEN_BRACKETS:
            stack.append((ch, i))
        elif ch in _CLOSE_BRACKETS:
            expected_open = _BRACKET_PAIRS.get(ch)
            if stack and stack[-1][0] == expected_open:
                stack.pop()
            else:
                # 多余的右括号
                issues.append(TextIssue(
                    rule="bracket_mismatch",
                    category="通用·标点",
                    item="括号不对称",
                    status=CheckResult.FAIL,
                    message=f"多余的右括号 '{ch}'",
                    paragraph_index=para_info.index,
                    paragraph_source=para_info.source,
                    char_offset=i,
                    context=_context_snippet(text, i),
                    fixable=False,
                ))

    # 检查未闭合的左括号
    for bracket, pos in stack:
        expected_close = _BRACKET_PAIRS.get(bracket, '')
        # 相邻段落宽松匹配
        if expected_close and expected_close in next_para_start:
            continue  # 下一段开头找到了对应右括号，不报告
        issues.append(TextIssue(
            rule="bracket_mismatch",
            category="通用·标点",
            item="括号不对称",
            status=CheckResult.FAIL,
            message=f"未闭合的左括号 '{bracket}'",
            paragraph_index=para_info.index,
            paragraph_source=para_info.source,
            char_offset=pos,
            context=_context_snippet(text, pos),
            fixable=False,
        ))

    return issues


def _check_quote_mismatch(
    para_info: ParagraphInfo,
    text: str,
) -> list[TextIssue]:
    """检查引号不匹配"""
    issues = []

    # 中文双引号
    open_double = text.count('\u201c')  # "
    close_double = text.count('\u201d')  # "
    if open_double != close_double:
        issues.append(TextIssue(
            rule="quote_mismatch",
            category="通用·标点",
            item="引号不匹配",
            status=CheckResult.WARN,
            message=f"中文双引号不匹配：左引号 {open_double} 个，右引号 {close_double} 个",
            paragraph_index=para_info.index,
            paragraph_source=para_info.source,
            fixable=False,
        ))

    # 中文单引号
    open_single = text.count('\u2018')  # '
    close_single = text.count('\u2019')  # '
    if open_single != close_single:
        issues.append(TextIssue(
            rule="quote_mismatch",
            category="通用·标点",
            item="引号不匹配",
            status=CheckResult.WARN,
            message=f"中文单引号不匹配：左引号 {open_single} 个，右引号 {close_single} 个",
            paragraph_index=para_info.index,
            paragraph_source=para_info.source,
            fixable=False,
        ))

    return issues


def _check_duplicate_punctuation(
    para_info: ParagraphInfo,
    text: str,
) -> list[TextIssue]:
    """检查连续标点。

    报告：。。 ，， 、、 ；； ：： .. (非省略号)
    不报告：！！ ？？ …… 单个…  ...
    """
    issues = []

    for m in _DUPLICATE_PUNCT_RE.finditer(text):
        matched = m.group()
        pos = m.start()
        issues.append(TextIssue(
            rule="duplicate_punctuation",
            category="通用·标点",
            item="连续标点",
            status=CheckResult.FAIL,
            message=f"连续标点 '{matched}'",
            paragraph_index=para_info.index,
            paragraph_source=para_info.source,
            char_offset=pos,
            context=_context_snippet(text, pos),
            fixable=True,
        ))

    return issues


def _check_extra_spaces_in_chinese(
    para_info: ParagraphInfo,
    text: str,
) -> list[TextIssue]:
    """检查中文之间多余空格"""
    issues = []

    for m in _CJK_SPACE_CJK_RE.finditer(text):
        pos = m.start(2)  # 空格的位置
        issues.append(TextIssue(
            rule="extra_spaces_in_chinese",
            category="通用·空格",
            item="中文之间多余空格",
            status=CheckResult.FAIL,
            message=f"中文字符之间不应有空格：'{m.group()}'",
            paragraph_index=para_info.index,
            paragraph_source=para_info.source,
            char_offset=pos,
            context=_context_snippet(text, pos),
            fixable=True,
        ))

    return issues


def _check_consecutive_spaces(
    para_info: ParagraphInfo,
    text: str,
) -> list[TextIssue]:
    """检查连续多个空格"""
    issues = []

    for m in _MULTI_SPACE_RE.finditer(text):
        pos = m.start()
        count = len(m.group())
        issues.append(TextIssue(
            rule="consecutive_spaces",
            category="通用·空格",
            item="连续多个空格",
            status=CheckResult.WARN,
            message=f"连续 {count} 个空格",
            paragraph_index=para_info.index,
            paragraph_source=para_info.source,
            char_offset=pos,
            context=_context_snippet(text, pos),
            fixable=True,
        ))

    return issues


def _check_leading_trailing_spaces(
    para_info: ParagraphInfo,
    text: str,
) -> list[TextIssue]:
    """检查行首/行尾空格（排除 XML 首行缩进控制的段落）"""
    issues = []

    # 行首空格（排除 XML 首行缩进）
    if text != text.lstrip(' ') and not para_info.has_xml_indent:
        leading = len(text) - len(text.lstrip(' '))
        issues.append(TextIssue(
            rule="leading_trailing_spaces",
            category="通用·空格",
            item="行首空格",
            status=CheckResult.WARN,
            message=f"段落开头有 {leading} 个空格",
            paragraph_index=para_info.index,
            paragraph_source=para_info.source,
            char_offset=0,
            fixable=True,
        ))

    # 行尾空格
    if text != text.rstrip(' '):
        trailing = len(text) - len(text.rstrip(' '))
        issues.append(TextIssue(
            rule="leading_trailing_spaces",
            category="通用·空格",
            item="行尾空格",
            status=CheckResult.WARN,
            message=f"段落末尾有 {trailing} 个空格",
            paragraph_index=para_info.index,
            paragraph_source=para_info.source,
            char_offset=len(text) - trailing,
            fixable=True,
        ))

    return issues


def _check_fullwidth_space(
    para_info: ParagraphInfo,
    text: str,
) -> list[TextIssue]:
    """检查全角空格混入"""
    issues = []

    for i, ch in enumerate(text):
        if ch == _FULLWIDTH_SPACE:
            issues.append(TextIssue(
                rule="fullwidth_space",
                category="通用·空格",
                item="全角空格",
                status=CheckResult.WARN,
                message="正文中出现全角空格（U+3000）",
                paragraph_index=para_info.index,
                paragraph_source=para_info.source,
                char_offset=i,
                context=_context_snippet(text, i),
                fixable=True,
            ))

    return issues


def _check_halfwidth_punctuation_in_chinese(
    para_info: ParagraphInfo,
    text: str,
) -> list[TextIssue]:
    """检查中文段落中的半角标点（争议项，需 LLM 审查）"""
    issues = []

    for i, ch in enumerate(text):
        if ch not in _EN_PUNCT_IN_CN:
            continue
        # 检查前后字符是否为 CJK
        prev_ch = text[i - 1] if i > 0 else ''
        next_ch = text[i + 1] if i + 1 < len(text) else ''
        if _CJK_IDEO_RE.match(prev_ch) or _CJK_IDEO_RE.match(next_ch):
            issues.append(TextIssue(
                rule="fullwidth_halfwidth_punctuation",
                category="通用·全半角",
                item="全半角标点混用",
                status=CheckResult.WARN,
                message=f"中文语境中出现半角标点 '{ch}'",
                paragraph_index=para_info.index,
                paragraph_source=para_info.source,
                char_offset=i,
                context=_context_snippet(text, i),
                fixable=True,
                is_disputed=True,
            ))

    return issues


def _check_sentence_ending(
    para_info: ParagraphInfo,
    text: str,
) -> list[TextIssue]:
    """检查句末标点缺失（争议项，需 LLM 审查）"""
    issues = []
    stripped = text.strip()
    if not stripped:
        return issues

    last_char = stripped[-1]
    # 如果段落以标点、数字、英文字母、括号结尾，不报告
    if last_char in '。！？…—）】》\u201d\u2019.!?)]}':
        return issues
    # 如果是标题样式，不报告
    style_lower = para_info.style_name.lower()
    if 'heading' in style_lower or '标题' in para_info.style_name or '目录' in para_info.style_name:
        return issues
    # 如果是纯英文段落，不检查
    if _cjk_ratio(stripped) < 0.1:
        return issues
    # 如果段落很短（可能是表格内容、公式等），不报告
    if len(stripped) < 10:
        return issues

    issues.append(TextIssue(
        rule="sentence_ending_punctuation",
        category="通用·标点",
        item="句末标点缺失",
        status=CheckResult.WARN,
        message=f"段落末尾缺少句号，末字符 '{last_char}'",
        paragraph_index=para_info.index,
        paragraph_source=para_info.source,
        char_offset=len(stripped) - 1,
        context=_context_snippet(stripped, len(stripped) - 1),
        fixable=False,
        is_disputed=True,
    ))

    return issues


def _count_cjk_english_spacing(text: str) -> tuple[int, int]:
    """统计中英文交界处有空格和无空格的数量"""
    spaced = 0
    unspaced = 0

    for i in range(1, len(text)):
        prev_ch = text[i - 1]
        curr_ch = text[i]

        # CJK → Latin 或 Latin → CJK
        prev_cjk = bool(_CJK_IDEO_RE.match(prev_ch))
        curr_cjk = bool(_CJK_IDEO_RE.match(curr_ch))
        prev_latin = prev_ch.isascii() and prev_ch.isalpha()
        curr_latin = curr_ch.isascii() and curr_ch.isalpha()

        if (prev_cjk and curr_latin) or (prev_latin and curr_cjk):
            unspaced += 1
        elif prev_cjk and curr_ch == ' ':
            # 检查空格后是否为 Latin
            if i + 1 < len(text) and text[i + 1].isascii() and text[i + 1].isalpha():
                spaced += 1
        elif prev_latin and curr_ch == ' ':
            if i + 1 < len(text) and _CJK_IDEO_RE.match(text[i + 1]):
                spaced += 1

    return spaced, unspaced


def _mark_cjk_spacing_disputes(
    paragraphs: list[ParagraphInfo],
    all_issues: list[TextIssue],
    stats: DocumentStats,
    tc_rules: dict,
) -> None:
    """标记中英文间距不一致的争议项"""
    total = stats.cjk_spaced_count + stats.cjk_unspaced_count
    if total == 0:
        return

    # 判断主流风格
    majority_spaced = stats.cjk_spaced_count >= stats.cjk_unspaced_count
    require_space = tc_rules.get('cjk_english_spacing', {}).get('require_space')

    if require_space is True:
        majority_spaced = True
    elif require_space is False:
        majority_spaced = False

    for para_info in paragraphs:
        if _cjk_ratio(para_info.text) < 0.1:
            continue
        text = _mask_urls(para_info.text)

        for i in range(1, len(text)):
            prev_ch = text[i - 1]
            curr_ch = text[i]

            prev_cjk = bool(_CJK_IDEO_RE.match(prev_ch))
            curr_cjk = bool(_CJK_IDEO_RE.match(curr_ch))
            prev_latin = prev_ch.isascii() and prev_ch.isalpha()
            curr_latin = curr_ch.isascii() and curr_ch.isalpha()

            if majority_spaced:
                # 期望有空格，但没有
                if (prev_cjk and curr_latin) or (prev_latin and curr_cjk):
                    all_issues.append(TextIssue(
                        rule="cjk_english_spacing",
                        category="通用·全半角",
                        item="中英文间距不一致",
                        status=CheckResult.WARN,
                        message=f"中英文之间缺少空格（文档中 {stats.cjk_spaced_count}/{total} 处有空格）",
                        paragraph_index=para_info.index,
                        paragraph_source=para_info.source,
                        char_offset=i,
                        context=_context_snippet(text, i),
                        fixable=True,
                        is_disputed=True,
                    ))
            else:
                # 期望无空格，但有
                if prev_cjk and curr_ch == ' ':
                    if i + 1 < len(text) and text[i + 1].isascii() and text[i + 1].isalpha():
                        all_issues.append(TextIssue(
                            rule="cjk_english_spacing",
                            category="通用·全半角",
                            item="中英文间距不一致",
                            status=CheckResult.WARN,
                            message=f"中英文之间多余的空格（文档中 {stats.cjk_unspaced_count}/{total} 处无空格）",
                            paragraph_index=para_info.index,
                            paragraph_source=para_info.source,
                            char_offset=i,
                            context=_context_snippet(text, i),
                            fixable=True,
                            is_disputed=True,
                        ))


# ============================================================
# 默认配置
# ============================================================

def _default_text_conventions() -> dict:
    """代码 fallback 默认配置"""
    return {
        'bracket_mismatch': {'enabled': True},
        'quote_mismatch': {'enabled': True},
        'duplicate_punctuation': {'enabled': True},
        'extra_spaces_in_chinese': {'enabled': True},
        'consecutive_spaces': {'enabled': True},
        'leading_trailing_spaces': {'enabled': True},
        'fullwidth_space': {'enabled': True},
        'cjk_english_spacing': {'enabled': True, 'require_space': None},
        'fullwidth_halfwidth_punctuation': {'enabled': True, 'context': 'chinese'},
        'sentence_ending_punctuation': {'enabled': True},
        'ai_review': {'enabled': True},
    }
