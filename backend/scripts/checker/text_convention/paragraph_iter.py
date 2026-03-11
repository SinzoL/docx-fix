"""
文本排版检查 — 段落遍历器与辅助函数

提供遍历文档所有段落（主体+表格+脚注+尾注）的生成器，
以及文本处理相关的辅助函数。
"""

from __future__ import annotations

from typing import Generator

from docx.document import Document

from .constants import NSMAP, W, CJK_IDEO_RE, CODE_STYLE_KEYWORDS, URL_RE
from .models import ParagraphInfo


# ============================================================
# 段落遍历器
# ============================================================

def iter_all_paragraphs(doc: Document) -> Generator[ParagraphInfo, None, None]:
    """遍历文档所有段落（主体 + 表格 + 脚注 + 尾注），附加来源标记。

    遍历顺序：
    1. doc.paragraphs（主体）
    2. doc.tables → cell.paragraphs（表格）
    3. 脚注/尾注段落

    跳过：
    - 页眉/页脚段落
    - 空段落
    """
    idx = 0

    # 1. 主体段落
    for para in doc.paragraphs:
        text = para.text
        if not text or not text.strip():
            idx += 1
            continue
        style_name = (para.style.name or "Normal") if para.style else "Normal"
        has_indent = _has_xml_first_line_indent(para)
        yield ParagraphInfo(
            paragraph=para, index=idx, source="body",
            text=text, style_name=style_name, has_xml_indent=has_indent,
        )
        idx += 1

    # 2. 表格段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    if not text or not text.strip():
                        idx += 1
                        continue
                    style_name = (para.style.name or "Normal") if para.style else "Normal"
                    yield ParagraphInfo(
                        paragraph=para, index=idx, source="table",
                        text=text, style_name=style_name,
                    )
                    idx += 1

    # 3. 脚注
    try:
        footnotes_part = doc.part.footnotes_part
        if footnotes_part is not None:
            fn_el = footnotes_part._element
            for footnote in fn_el.findall(f'{{{W}}}footnote'):
                fn_type = footnote.get(f'{{{W}}}type')
                if fn_type in ('separator', 'continuationSeparator'):
                    continue
                for p_el in footnote.findall(f'{{{W}}}p'):
                    text = _extract_text_from_element(p_el)
                    if not text or not text.strip():
                        idx += 1
                        continue
                    yield ParagraphInfo(
                        paragraph=None, index=idx, source="footnote",
                        text=text, style_name="Footnote",
                    )
                    idx += 1
    except (AttributeError, Exception):
        pass

    # 4. 尾注
    try:
        endnotes_part = doc.part.endnotes_part
        if endnotes_part is not None:
            en_el = endnotes_part._element
            for endnote in en_el.findall(f'{{{W}}}endnote'):
                en_type = endnote.get(f'{{{W}}}type')
                if en_type in ('separator', 'continuationSeparator'):
                    continue
                for p_el in endnote.findall(f'{{{W}}}p'):
                    text = _extract_text_from_element(p_el)
                    if not text or not text.strip():
                        idx += 1
                        continue
                    yield ParagraphInfo(
                        paragraph=None, index=idx, source="endnote",
                        text=text, style_name="Endnote",
                    )
                    idx += 1
    except (AttributeError, Exception):
        pass


def _extract_text_from_element(p_el) -> str:
    """从 XML 段落元素中提取文本"""
    texts = []
    for r in p_el.findall(f'{{{W}}}r'):
        for t in r.findall(f'{{{W}}}t'):
            if t.text:
                texts.append(t.text)
    return ''.join(texts)


def _has_xml_first_line_indent(para) -> bool:
    """检查段落是否有 XML 首行缩进属性"""
    try:
        pPr = para._element.find('w:pPr', NSMAP)
        if pPr is not None:
            ind = pPr.find('w:ind', NSMAP)
            if ind is not None:
                fi = ind.get(f'{{{W}}}firstLine')
                fic = ind.get(f'{{{W}}}firstLineChars')
                if (fi and fi != '0') or (fic and fic != '0'):
                    return True
    except Exception:
        pass
    return False


# ============================================================
# 辅助函数
# ============================================================

def cjk_ratio(text: str) -> float:
    """计算段落中 CJK 表意文字的占比"""
    if not text:
        return 0.0
    total = len(text.strip())
    if total == 0:
        return 0.0
    cjk_count = len(CJK_IDEO_RE.findall(text))
    return cjk_count / total


def is_code_style(style_name: str) -> bool:
    """判断是否为代码样式"""
    name_lower = style_name.lower()
    return any(kw in name_lower for kw in CODE_STYLE_KEYWORDS)


def context_snippet(text: str, pos: int, width: int = 10) -> str:
    """提取上下文片段"""
    start = max(0, pos - width)
    end = min(len(text), pos + width)
    prefix = "..." if start > 0 else ""
    suffix = "..." if end < len(text) else ""
    return f"{prefix}{text[start:end]}{suffix}"


def mask_urls(text: str) -> str:
    """将 URL 和邮箱替换为占位符，避免误判"""
    return URL_RE.sub(lambda m: ' ' * len(m.group()), text)


def source_label(source: str) -> str:
    """来源标记翻译"""
    labels = {"body": "主体", "table": "表格", "footnote": "脚注", "endnote": "尾注"}
    return labels.get(source, source)


def location_str(para_info: ParagraphInfo, char_offset: int = -1) -> str:
    """构建位置字符串"""
    loc = f"段落{para_info.index + 1} [{source_label(para_info.source)}]"
    if char_offset >= 0:
        loc += f", 第{char_offset + 1}字符"
    return loc
