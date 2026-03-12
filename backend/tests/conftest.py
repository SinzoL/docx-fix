"""
后端测试配置和通用 Fixtures

提供：
- FastAPI 测试客户端（httpx AsyncClient）
- 临时文件目录
- 测试用 .docx 文件生成器
- 规则文件路径
"""

import os
import sys
import shutil
import tempfile

import pytest
import pytest_asyncio
from httpx import AsyncClient, ASGITransport

# 设置项目路径
BACKEND_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ENGINE_DIR = os.path.join(BACKEND_DIR, "engine")
PROJECT_ROOT = os.path.dirname(BACKEND_DIR)

# 确保 backend 和 engine 目录都在 sys.path 中
if BACKEND_DIR not in sys.path:
    sys.path.insert(0, BACKEND_DIR)
if ENGINE_DIR not in sys.path:
    sys.path.insert(0, ENGINE_DIR)


@pytest.fixture(scope="session")
def project_root():
    """项目根目录"""
    return PROJECT_ROOT


@pytest.fixture(scope="session")
def rules_dir():
    """规则文件目录"""
    return os.path.join(BACKEND_DIR, "rules")


@pytest.fixture(scope="session")
def default_rule_path():
    """默认规则文件路径"""
    return os.path.join(BACKEND_DIR, "rules", "default.yaml")


@pytest.fixture
def temp_dir():
    """创建临时目录，测试结束后清理"""
    d = tempfile.mkdtemp(prefix="docx-fix-test-")
    yield d
    shutil.rmtree(d, ignore_errors=True)


@pytest.fixture
def sample_docx(temp_dir):
    """生成一个简单的测试 .docx 文件"""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    # 添加一些基本内容
    doc.add_heading("测试文档", level=1)
    para = doc.add_paragraph("这是一段测试文本。")
    # 设置字体（故意使用非标准字体以产生检查错误）
    for run in para.runs:
        run.font.size = Pt(12)
        run.font.name = "Arial"  # 非宋体，会产生检查失败

    doc.add_paragraph("第二段内容。")

    filepath = os.path.join(temp_dir, "test_sample.docx")
    doc.save(filepath)
    return filepath


@pytest.fixture
def good_docx(temp_dir):
    """生成一个格式基本正确的 .docx 文件（基于 default.yaml 规则）"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn

    doc = Document()

    # 设置页面为 A4
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    # 添加正文段落，使用正确的字体
    para = doc.add_paragraph("这是一段正文内容。")
    for run in para.runs:
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = r.makeelement(qn("w:rFonts"), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), "宋体")

    filepath = os.path.join(temp_dir, "test_good.docx")
    doc.save(filepath)
    return filepath


@pytest.fixture
def corrupted_file(temp_dir):
    """生成一个损坏的文件（不是有效的 docx）"""
    filepath = os.path.join(temp_dir, "corrupted.docx")
    with open(filepath, "wb") as f:
        f.write(b"this is not a valid docx file")
    return filepath


@pytest.fixture
def non_docx_file(temp_dir):
    """生成一个非 docx 文件"""
    filepath = os.path.join(temp_dir, "test.txt")
    with open(filepath, "w") as f:
        f.write("plain text file")
    return filepath


@pytest_asyncio.fixture
async def client(temp_dir):
    """创建 FastAPI 测试客户端，使用临时目录"""
    # 设置临时目录环境变量
    os.environ["DOCX_FIX_TEMP_DIR"] = temp_dir

    from app import app

    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        yield ac

    # 清理环境变量
    os.environ.pop("DOCX_FIX_TEMP_DIR", None)
