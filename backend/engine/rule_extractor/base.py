"""
规则提取器 — 主类

RuleExtractor 通过组合 StyleExtractorMixin 和 StructureExtractorMixin，
从 Word 模板文档中提取完整的格式规则并生成 YAML 配置文件。
"""

import os
import re
import sys
from collections import OrderedDict

import yaml
from docx import Document

from .constants import NSMAP, W, Color, OrderedDumper
from .style_extractor import StyleExtractorMixin
from .structure_extractor import StructureExtractorMixin


# ===== 分节注释的 YAML 序列化常量 =====
_YAML_SECTIONS = [
    ('meta', '元信息'),
    ('page_setup', '一、页面设置'),
    ('header_footer', '二、页眉页脚'),
    ('styles', '三、样式定义规则\n# 每个样式包含：段落格式 + 字符格式'),
    ('structure', '四、文档结构规则'),
    ('numbering', '五、编号定义规则'),
    ('special_checks', '六、特殊检查规则'),
    ('heading_style_fix', '七、标题样式自动修复规则'),
]


def rules_to_yaml(rules: dict) -> str:
    """将规则字典转为分节注释的 YAML 字符串。

    这是公共工具函数，可在 CLI（save_yaml）和 Web API（extractor_service）中复用。

    Args:
        rules: 提取的规则字典（与 YAML 规则文件结构一致）。

    Returns:
        格式化的 YAML 字符串（含分节注释头）。
    """
    lines = []
    lines.append(f'# {"=" * 60}')
    lines.append(f'# {rules.get("meta", {}).get("name", "格式规则")}')
    lines.append(f'# {rules.get("meta", {}).get("description", "")}')
    lines.append(f'# {"=" * 60}')
    lines.append('')

    for key, comment in _YAML_SECTIONS:
        if key not in rules:
            continue

        lines.append(f'# {"=" * 28}')
        lines.append(f'# {comment}')
        lines.append(f'# {"=" * 28}')

        section_data = OrderedDict([(key, rules[key])])
        yaml_str = yaml.dump(
            dict(section_data),
            Dumper=OrderedDumper,
            default_flow_style=False,
            allow_unicode=True,
            width=120,
            sort_keys=False,
        )
        lines.append(yaml_str)

    return '\n'.join(lines)


class RuleExtractor(StyleExtractorMixin, StructureExtractorMixin):
    """从 Word 模板文档提取格式规则"""

    def __init__(self, filepath):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.rules = OrderedDict()

    # ========================================
    # 元信息提取
    # ========================================
    def extract_meta(self, name=None, description=None):
        """生成规则文件的元信息"""
        basename = os.path.splitext(os.path.basename(self.filepath))[0]
        self.rules['meta'] = OrderedDict([
            ('name', name or f'从 {basename} 提取的格式规则'),
            ('version', '1.0'),
            ('description', description or f'从模板文件 "{os.path.basename(self.filepath)}" 自动提取'),
            ('source_file', os.path.basename(self.filepath)),
        ])

    # ========================================
    # 页面设置提取
    # ========================================
    def extract_page_setup(self):
        """从第一个节（section）提取页面设置"""
        if not self.doc.sections:
            return

        section = self.doc.sections[0]
        page_setup = OrderedDict()

        # 纸张大小
        if section.page_width and section.page_height:
            w_cm = round(section.page_width / 360000, 1)
            h_cm = round(section.page_height / 360000, 1)
            if abs(w_cm - 21.0) < 0.3 and abs(h_cm - 29.7) < 0.3:
                page_setup['paper_size'] = 'A4'
            else:
                page_setup['paper_size'] = 'custom'
            page_setup['width_cm'] = w_cm
            page_setup['height_cm'] = h_cm

        # 页边距
        margin_attrs = [
            ('margin_top_cm', 'top_margin'),
            ('margin_bottom_cm', 'bottom_margin'),
            ('margin_left_cm', 'left_margin'),
            ('margin_right_cm', 'right_margin'),
        ]
        for rule_key, attr in margin_attrs:
            val = getattr(section, attr, None)
            if val is not None:
                page_setup[rule_key] = round(val / 360000, 2)

        # 页眉页脚距离
        if section.header_distance is not None:
            page_setup['header_distance_cm'] = round(section.header_distance / 360000, 2)
        if section.footer_distance is not None:
            page_setup['footer_distance_cm'] = round(section.footer_distance / 360000, 2)

        # 装订线
        sect_pr = section._sectPr
        if sect_pr is not None:
            pgMar = sect_pr.find('w:pgMar', NSMAP)
            if pgMar is not None:
                gutter = pgMar.get(f'{{{W}}}gutter')
                if gutter:
                    page_setup['gutter_cm'] = round(int(gutter) / 567, 2)

        # 文档网格
        if sect_pr is not None:
            docGrid = sect_pr.find('w:docGrid', NSMAP)
            if docGrid is not None:
                grid_info = OrderedDict()
                grid_type = docGrid.get(f'{{{W}}}type')
                if grid_type:
                    grid_info['type'] = grid_type
                line_pitch = docGrid.get(f'{{{W}}}linePitch')
                if line_pitch:
                    grid_info['line_pitch'] = int(line_pitch)
                char_space = docGrid.get(f'{{{W}}}charSpace')
                if char_space:
                    grid_info['char_space'] = int(char_space)
                if grid_info:
                    page_setup['doc_grid'] = grid_info

        self.rules['page_setup'] = page_setup

    # ========================================
    # 页眉页脚提取
    # ========================================
    def extract_header_footer(self):
        """提取页眉页脚内容和格式"""
        hf = OrderedDict()

        # 页眉
        for section in self.doc.sections:
            header = section.header
            if header and header.paragraphs:
                for p in header.paragraphs:
                    text = p.text.strip()
                    if text:
                        header_info = OrderedDict()
                        header_info['text'] = text
                        style_name = p.style.name if p.style else None
                        if style_name:
                            header_info['style'] = style_name

                        # 对齐
                        pPr = p._element.find('w:pPr', NSMAP)
                        if pPr is not None:
                            jc = pPr.find('w:jc', NSMAP)
                            if jc is not None:
                                header_info['alignment'] = jc.get(f'{{{W}}}val')

                        # 字号
                        for run in p.runs:
                            if run.font.size:
                                header_info['font_size_pt'] = round(run.font.size / 12700, 1)
                                break

                        hf['header'] = header_info
                        break
                if 'header' in hf:
                    break

        # 页脚
        for section in self.doc.sections:
            footer = section.footer
            if footer and footer.paragraphs:
                for p in footer.paragraphs:
                    footer_info = OrderedDict()
                    style_name = p.style.name if p.style else None
                    if style_name:
                        footer_info['style'] = style_name

                    pPr = p._element.find('w:pPr', NSMAP)
                    if pPr is not None:
                        jc = pPr.find('w:jc', NSMAP)
                        if jc is not None:
                            footer_info['alignment'] = jc.get(f'{{{W}}}val')

                    for run in p.runs:
                        if run.font.size:
                            footer_info['font_size_pt'] = round(run.font.size / 12700, 1)
                            break

                    if footer_info:
                        hf['footer'] = footer_info
                    break
                if 'footer' in hf:
                    break

        if hf:
            self.rules['header_footer'] = hf

    # ========================================
    # 汇总提取
    # ========================================
    def extract_all(self, name=None, description=None):
        """执行所有提取步骤"""
        print(f"\n{Color.BOLD}{'=' * 60}{Color.END}")
        print(f"  {Color.CYAN}从模板提取格式规则{Color.END}")
        print(f"  源文件: {os.path.basename(self.filepath)}")
        print(f"{Color.BOLD}{'=' * 60}{Color.END}\n")

        steps = [
            ('元信息', lambda: self.extract_meta(name, description)),
            ('页面设置', self.extract_page_setup),
            ('页眉页脚', self.extract_header_footer),
            ('样式定义', self.extract_styles),
            ('文档结构', self.extract_structure),
            ('编号定义', self.extract_numbering),
            ('特殊检查规则', self.extract_special_checks),
            ('标题修复规则', self.extract_heading_style_fix),
        ]

        for step_name, func in steps:
            try:
                func()
                print(f"  {Color.GREEN}✓{Color.END} {step_name}")
            except Exception as e:
                print(f"  {Color.YELLOW}⚠{Color.END} {step_name}: {e}")

        # 收集审核上下文信息（供 LLM 审核使用）
        try:
            self._colored_text_paragraphs = self.collect_colored_text_paragraphs()
        except Exception:
            self._colored_text_paragraphs = []
        try:
            self._heading_structure = self.collect_heading_structure()
        except Exception:
            self._heading_structure = []

        return self.rules

    def save_yaml(self, output_path):
        """将提取的规则保存为 YAML 文件"""
        content = rules_to_yaml(self.rules)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)

        return output_path

    def print_summary(self):
        """打印提取结果摘要"""
        print(f"\n{Color.BOLD}{'─' * 60}{Color.END}")
        print(f"  {Color.CYAN}提取结果摘要{Color.END}")
        print(f"{Color.BOLD}{'─' * 60}{Color.END}")

        ps = self.rules.get('page_setup', {})
        if ps:
            print(f"\n  {Color.BOLD}页面设置:{Color.END}")
            print(f"    纸张: {ps.get('paper_size', '?')} ({ps.get('width_cm', '?')}×{ps.get('height_cm', '?')}cm)")
            print(f"    页边距: 上{ps.get('margin_top_cm', '?')} 下{ps.get('margin_bottom_cm', '?')} "
                  f"左{ps.get('margin_left_cm', '?')} 右{ps.get('margin_right_cm', '?')}cm")

        styles = self.rules.get('styles', {})
        if styles:
            heading_styles = []
            body_styles = []
            other_styles = []
            for name, info in styles.items():
                para = info.get('paragraph', {})
                if para.get('outline_level') is not None:
                    heading_styles.append(name)
                elif name in ('Normal', '论文正文-首行缩进', '英文正文'):
                    body_styles.append(name)
                else:
                    other_styles.append(name)

            print(f"\n  {Color.BOLD}样式定义:{Color.END} 共 {len(styles)} 个")
            if heading_styles:
                print(f"    标题样式: {', '.join(heading_styles)}")
            if body_styles:
                print(f"    正文样式: {', '.join(body_styles)}")
            if other_styles:
                print(f"    其他样式: {', '.join(other_styles[:10])}"
                      f"{'...' if len(other_styles) > 10 else ''}")

        structure = self.rules.get('structure', {})
        chapters = structure.get('required_chapters', [])
        if chapters:
            print(f"\n  {Color.BOLD}文档结构:{Color.END}")
            print(f"    章节: {len(chapters)} 个一级标题")
            for ch in chapters[:5]:
                print(f"      · {ch.get('pattern', '?')}")

        numbering = self.rules.get('numbering', {})
        if numbering:
            print(f"\n  {Color.BOLD}编号定义:{Color.END} {len(numbering)} 组")
            for nk, nv in numbering.items():
                lvl_count = len(nv.get('levels', {}))
                print(f"    {nk}: {nv.get('type', '?')} ({lvl_count} 级)")

        print()


def main():
    """CLI 入口"""
    if len(sys.argv) < 2:
        print("用法: python -m engine.rule_extractor <模板docx文件> [--output <yaml路径>] [--name <规则名称>]")
        print()
        print("示例:")
        print("  python -m engine.rule_extractor template.docx")
        print("  python -m engine.rule_extractor template.docx --output rules/my_rules.yaml")
        print("  python -m engine.rule_extractor template.docx --name '我的学校论文格式规则'")
        sys.exit(1)

    filepath = sys.argv[1]

    if not os.path.exists(filepath):
        print(f"❌ 文件不存在: {filepath}")
        sys.exit(1)

    # 解析参数
    output_path = None
    if '--output' in sys.argv:
        idx = sys.argv.index('--output')
        if idx + 1 < len(sys.argv):
            output_path = sys.argv[idx + 1]

    name = None
    if '--name' in sys.argv:
        idx = sys.argv.index('--name')
        if idx + 1 < len(sys.argv):
            name = sys.argv[idx + 1]

    # 默认输出路径
    if output_path is None:
        basename = os.path.splitext(os.path.basename(filepath))[0]
        safe_name = re.sub(r'[^\w\u4e00-\u9fff-]', '_', basename)
        output_path = os.path.join(
            os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
            'rules',
            f'{safe_name}.yaml'
        )

    # 提取规则
    extractor = RuleExtractor(filepath)
    extractor.extract_all(name=name)
    extractor.print_summary()

    # 保存
    extractor.save_yaml(output_path)
    print(f"  {Color.GREEN}✓ 规则已保存到:{Color.END} {output_path}")
    print("\n  使用方法:")
    print(f"    python docx_fixer.py check <文档.docx> --rules {output_path}")
    print(f"    python docx_fixer.py fix-format <文档.docx> --rules {output_path}")
    print()


if __name__ == "__main__":
    main()
