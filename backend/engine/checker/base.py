#!/usr/bin/env python3
"""
Word 文档格式检查引擎 — 基础模块

包含核心类 DocxChecker 的骨架、公共定义和通用检查方法。
领域检查方法分布在各子模块中（style_checker / heading_validator / numbering_checker）。

用法：
    python -m engine.checker.base <docx文件路径> [--rules <yaml规则文件>]
"""

import sys
import os
import re
import yaml
from docx import Document

from engine.shared_constants import NSMAP, W, FONT_ALIASES, fonts_match, Color  # noqa: F401


class CheckResult:
    """单条检查结果"""
    PASS = "PASS"
    WARN = "WARN"
    FAIL = "FAIL"

    def __init__(self, category, item, status, message, location=None, fixable=False, check_layer="format"):
        self.category = category
        self.item = item
        self.status = status
        self.message = message
        self.location = location
        self.fixable = fixable
        self.check_layer = check_layer

    def __str__(self):
        loc = f" [{self.location}]" if self.location else ""
        fix = " (可自动修复)" if self.fixable else ""
        if self.status == self.PASS:
            icon = f"{Color.GREEN}✓{Color.END}"
        elif self.status == self.WARN:
            icon = f"{Color.YELLOW}⚠{Color.END}"
        else:
            icon = f"{Color.RED}✗{Color.END}"
        return f"  {icon} [{self.category}] {self.item}{loc}: {self.message}{fix}"


class DocxChecker:
    """文档格式检查器"""

    def __init__(self, filepath, rules_path=None):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.results = []

        # 加载规则
        if rules_path is None:
            rules_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "rules", "default.yaml")
        with open(rules_path, 'r', encoding='utf-8') as f:
            self.rules = yaml.safe_load(f)

        # 初始化属性解析器（延迟导入避免循环引用）
        from engine.checker.property_resolver import PropertyResolver
        self.resolver = PropertyResolver(self.doc)

    def add_result(self, category, item, status, message, location=None, fixable=False, check_layer="format"):
        r = CheckResult(category, item, status, message, location, fixable, check_layer)
        self.results.append(r)
        return r

    # ========================================
    # 1. 页面设置检查
    # ========================================
    def check_page_setup(self):
        rules = self.rules.get('page_setup', {})
        if not rules:
            return

        for i, section in enumerate(self.doc.sections):
            loc = f"节{i+1}" if len(self.doc.sections) > 1 else None

            # 纸张大小
            w_cm = round((section.page_width or 0) / 360000, 1)
            h_cm = round((section.page_height or 0) / 360000, 1)
            exp_w = rules.get('width_cm', 21.0)
            exp_h = rules.get('height_cm', 29.7)
            if abs(w_cm - exp_w) > 0.2 or abs(h_cm - exp_h) > 0.2:
                self.add_result("页面设置", "纸张大小", CheckResult.FAIL,
                                f"当前 {w_cm}×{h_cm}cm, 要求 {exp_w}×{exp_h}cm (A4)", loc, fixable=True)
            else:
                self.add_result("页面设置", "纸张大小", CheckResult.PASS,
                                f"A4 ({w_cm}×{h_cm}cm)", loc)

            # 页边距
            margins = {
                '上边距': (section.top_margin, 'margin_top_cm'),
                '下边距': (section.bottom_margin, 'margin_bottom_cm'),
                '左边距': (section.left_margin, 'margin_left_cm'),
                '右边距': (section.right_margin, 'margin_right_cm'),
            }
            for name, (actual_emu, key) in margins.items():
                if actual_emu is None:
                    continue
                actual_cm = round(actual_emu / 360000, 2)
                expected_cm = rules.get(key)
                if expected_cm is not None:
                    if abs(actual_cm - expected_cm) > 0.1:
                        self.add_result("页面设置", name, CheckResult.FAIL,
                                        f"当前 {actual_cm}cm, 要求 {expected_cm}cm", loc, fixable=True)
                    else:
                        self.add_result("页面设置", name, CheckResult.PASS,
                                        f"{actual_cm}cm", loc)

            # 页眉页脚距离
            if section.header_distance is not None:
                hd_cm = round(section.header_distance / 360000, 2)
                exp_hd = rules.get('header_distance_cm')
                if exp_hd and abs(hd_cm - exp_hd) > 0.1:
                    self.add_result("页面设置", "页眉距顶端", CheckResult.FAIL,
                                    f"当前 {hd_cm}cm, 要求 {exp_hd}cm", loc, fixable=True)

            if section.footer_distance is not None:
                fd_cm = round(section.footer_distance / 360000, 2)
                exp_fd = rules.get('footer_distance_cm')
                if exp_fd and abs(fd_cm - exp_fd) > 0.1:
                    self.add_result("页面设置", "页脚距底端", CheckResult.FAIL,
                                    f"当前 {fd_cm}cm, 要求 {exp_fd}cm", loc, fixable=True)

    # ========================================
    # 2. 页眉页脚检查
    # ========================================
    def check_header_footer(self):
        hf_rules = self.rules.get('header_footer', {})
        header_rules = hf_rules.get('header', {})

        if header_rules:
            expected_text = header_rules.get('text', '')
            found = False
            for section in self.doc.sections:
                header = section.header
                if header and header.paragraphs:
                    for p in header.paragraphs:
                        if p.text.strip():
                            found = True
                            if expected_text and expected_text not in p.text.strip() and p.text.strip() not in expected_text:
                                self.add_result("页眉页脚", "页眉内容", CheckResult.WARN,
                                                f"当前: \"{p.text.strip()[:30]}\", 期望包含: \"{expected_text[:30]}\"")
                            else:
                                self.add_result("页眉页脚", "页眉内容", CheckResult.PASS,
                                                f"\"{p.text.strip()[:40]}\"")
                            break
                if found:
                    break
            if not found:
                self.add_result("页眉页脚", "页眉", CheckResult.FAIL, "未找到页眉内容")

    # ========================================
    # 6. TOC 目录域检查
    # ========================================
    def check_toc(self):
        """检查目录域代码"""
        toc_rules = self.rules.get('structure', {}).get('toc', {})
        if not toc_rules:
            return

        toc_found = False
        for para in self.doc.paragraphs:
            for run in para._element.findall('.//w:r', NSMAP):
                instrText = run.find('w:instrText', NSMAP)
                if instrText is not None and instrText.text and 'TOC' in instrText.text:
                    toc_found = True
                    code = instrText.text.strip()

                    expected_range = toc_rules.get('outline_range', '1-3')
                    m = re.search(r'\\o\s*"1-(\d+)"', code)
                    if m:
                        actual_max = m.group(1)
                        expected_max = expected_range.split('-')[-1]
                        if actual_max != expected_max:
                            self.add_result("目录", "TOC大纲范围", CheckResult.FAIL,
                                            f"当前 \\o \"1-{actual_max}\", 要求 \\o \"1-{expected_max}\"",
                                            fixable=True)
                        else:
                            self.add_result("目录", "TOC大纲范围", CheckResult.PASS,
                                            f"\\o \"1-{actual_max}\"")

                    custom_styles = toc_rules.get('custom_styles', {})
                    if custom_styles:
                        if '\\t' not in code and '\\t' not in code:
                            self.add_result("目录", "自定义样式映射", CheckResult.WARN,
                                            "TOC 中未找到 \\t 自定义样式映射")
                    break
            if toc_found:
                break

        if not toc_found:
            self.add_result("目录", "TOC域", CheckResult.WARN, "未找到目录域代码")

    # ========================================
    # 辅助方法
    # ========================================
    def _is_cover_page_paragraph(self, para):
        """判断是否为封面页段落"""
        text = para.text.strip()
        style_rules = self.rules.get('styles', {})
        cover_rules = style_rules.get('cover_title', {})
        if cover_rules.get('check_type') == 'content_match':
            patterns = cover_rules.get('content_patterns', [])
            for pattern in patterns:
                if pattern in text:
                    return True
        return False

    def _get_para_outline_level(self, para):
        """获取段落的实际大纲级别"""
        pPr = para._element.find('w:pPr', NSMAP)
        if pPr is not None:
            ol = pPr.find('w:outlineLvl', NSMAP)
            if ol is not None:
                return int(ol.get(f'{{{W}}}val'))

        if para.style and para.style.element is not None:
            style_pPr = para.style.element.find('.//w:pPr', NSMAP)
            if style_pPr is not None:
                ol = style_pPr.find('w:outlineLvl', NSMAP)
                if ol is not None:
                    return int(ol.get(f'{{{W}}}val'))

        if para.style and para.style.name:
            m = re.match(r'Heading (\d+)', para.style.name)
            if m:
                return int(m.group(1)) - 1

        return 9

    # ========================================
    # 运行所有检查
    # ========================================
    def run_all_checks(self):
        """执行所有检查项，调用各子模块的检查方法"""
        self.results = []

        # 通用检查（base.py 自身）
        self.check_page_setup()
        self.check_header_footer()

        # 样式检查（style_checker.py）
        from engine.checker.style_checker import (
            check_style_definitions,
            check_paragraph_formatting,
            check_template_instructions,
            check_font_consistency,
            check_figure_table_captions,
        )
        check_style_definitions(self)

        # 标题检查（heading_validator.py）
        from engine.checker.heading_validator import (
            check_heading_styles,
            check_document_structure,
            check_heading_hierarchy,
        )
        check_heading_styles(self)
        check_document_structure(self)
        check_heading_hierarchy(self)

        # 目录检查（base.py 自身）
        self.check_toc()

        # 编号检查（numbering_checker.py）
        from engine.checker.numbering_checker import (
            check_heading_numbering,
            check_heading_lvl_text,
            check_heading_numid_override,
            check_shared_abstract_num,
            check_heading_style_and_manual_numbering,
            check_heading_numbering_indent,
        )
        check_heading_numbering(self)
        check_heading_lvl_text(self)
        check_heading_numid_override(self)
        check_shared_abstract_num(self)
        check_heading_style_and_manual_numbering(self)

        # 段落格式检查（style_checker.py）
        check_paragraph_formatting(self)

        # 特殊检查（style_checker.py）
        check_template_instructions(self)
        check_font_consistency(self)
        check_figure_table_captions(self)

        # 编号缩进检查（numbering_checker.py）
        check_heading_numbering_indent(self)

        # 文本排版习惯检查（text_convention_checker.py）
        from engine.checker.text_convention_checker import run_text_convention_checks
        self._text_issues, self._text_stats = run_text_convention_checks(
            self, self.doc, self.rules
        )

        return self.results

    def print_report(self):
        """打印检查报告"""
        print(f"\n{'=' * 70}")
        print(f"  文档格式检查报告: {os.path.basename(self.filepath)}")
        print(f"  规则: {self.rules.get('meta', {}).get('name', 'N/A')}")
        print(f"{'=' * 70}\n")

        categories = {}
        for r in self.results:
            if r.category not in categories:
                categories[r.category] = []
            categories[r.category].append(r)

        pass_count = sum(1 for r in self.results if r.status == CheckResult.PASS)
        warn_count = sum(1 for r in self.results if r.status == CheckResult.WARN)
        fail_count = sum(1 for r in self.results if r.status == CheckResult.FAIL)
        fixable_count = sum(1 for r in self.results if r.fixable)

        for cat, items in categories.items():
            print(f"{Color.BOLD}▸ {cat}{Color.END}")
            for item in items:
                print(str(item))
            print()

        print(f"{'─' * 70}")
        print(f"  汇总: {Color.GREEN}通过 {pass_count}{Color.END} / "
              f"{Color.YELLOW}警告 {warn_count}{Color.END} / "
              f"{Color.RED}错误 {fail_count}{Color.END}")
        if fixable_count:
            print(f"  其中 {Color.BLUE}{fixable_count} 项可自动修复{Color.END}，"
                  f"运行 fix-format 命令进行修复")
        print()

        return fail_count == 0


def main():
    if len(sys.argv) < 2:
        print("用法: python -m engine.checker.base <docx文件路径> [--rules <yaml规则文件>]")
        sys.exit(1)

    filepath = sys.argv[1]
    rules_path = None
    if "--rules" in sys.argv:
        idx = sys.argv.index("--rules")
        if idx + 1 < len(sys.argv):
            rules_path = sys.argv[idx + 1]

    checker = DocxChecker(filepath, rules_path)
    checker.run_all_checks()
    checker.print_report()


if __name__ == "__main__":
    main()
