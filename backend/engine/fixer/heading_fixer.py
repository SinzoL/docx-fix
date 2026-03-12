# pyright: reportMissingImports=false, reportUnknownVariableType=false
# pyright: reportUnknownMemberType=false, reportUnknownArgumentType=false
# pyright: reportAny=false, reportExplicitAny=false
# pyright: reportUninitializedInstanceVariable=false
# pyright: reportUnusedParameter=false, reportMissingParameterType=false
"""
标题与段落修复混入类

包含：段落大纲级别、TOC域代码、标题样式替换、手动编号去除、
说明文字修复、图注样式修复、标题缩进修复等逻辑。
"""

from __future__ import annotations

import re
from typing import TYPE_CHECKING, Any, cast

from lxml import etree

from .constants import NSMAP, W

if TYPE_CHECKING:
    from docx.document import Document
    from docx.styles.style import ParagraphStyle


class HeadingFixerMixin:
    """标题与段落相关修复方法的混入类

    运行时由 DocxFixer 通过多继承组合，以下属性由 DocxFixer.__init__ 提供。
    """

    doc: Document
    rules: dict[str, Any]
    dry_run: bool
    fixes: list[tuple[str, str]]

    def log_fix(self, category: str, description: str) -> None: ...

    # ========================================
    # 3. 修复段落大纲级别
    # ========================================
    def fix_paragraph_outline_levels(self):
        """确保标题段落有正确的大纲级别"""
        style_rules = self.rules.get('styles', {})

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            style_name = (para.style.name or "None") if para.style else "None"
            if style_name.lower().startswith('toc ') or style_name == '说明文字':
                continue

            rules = style_rules.get(style_name)
            if not rules:
                continue

            para_rules = rules.get('paragraph', {})
            if 'outline_level' not in para_rules:
                continue

            expected_lvl = para_rules['outline_level']

            pPr = para._element.find('w:pPr', NSMAP)
            current_lvl = None
            ol_elem: Any = None
            if pPr is not None:
                ol_elem = pPr.find('w:outlineLvl', NSMAP)
                if ol_elem is not None:
                    current_lvl = int(ol_elem.get(f'{{{W}}}val'))

            if current_lvl is not None and current_lvl != expected_lvl:
                self.log_fix("段落大纲", f"段落{i} \"{text[:30]}\" 大纲级别 {current_lvl} → {expected_lvl}")
                if not self.dry_run and ol_elem is not None:
                    ol_elem.set(f'{{{W}}}val', str(expected_lvl))

    # ========================================
    # 4. 修复 TOC 域代码
    # ========================================
    def fix_toc(self):
        """修复目录域代码"""
        toc_rules = self.rules.get('structure', {}).get('toc', {})
        if not toc_rules:
            return

        expected_range = toc_rules.get('outline_range', '1-3')
        expected_max = expected_range.split('-')[-1]

        for para in self.doc.paragraphs:
            for run in para._element.findall('.//w:r', NSMAP):
                instrText = run.find('w:instrText', NSMAP)
                if instrText is not None and instrText.text and 'TOC' in instrText.text:
                    old_text = instrText.text
                    m = re.search(r'\\o\s*"1-(\d+)"', old_text)
                    if m and m.group(1) != expected_max:
                        new_text = re.sub(r'\\o\s*"1-\d+"', f'\\\\o "1-{expected_max}"', old_text)
                        self.log_fix("TOC", f"大纲范围 \\o \"1-{m.group(1)}\" → \\o \"1-{expected_max}\"")
                        if not self.dry_run:
                            instrText.text = new_text

    # ========================================
    # 10. 修复标题样式 & 去除手动编号
    # ========================================
    def fix_heading_style_and_manual_numbering(self):
        """修复使用错误样式的标题段落，并去除手动编号文本。"""
        fix_rules = self.rules.get('heading_style_fix', {})
        if not fix_rules.get('enabled'):
            return

        style_map = fix_rules.get('style_replacement', {})
        num_patterns = fix_rules.get('manual_numbering_patterns', {})

        if not style_map:
            return

        # 预编译目标样式对象
        target_styles = {}
        for _old_name, new_name in style_map.items():
            for style in self.doc.styles:
                if style.name == new_name:
                    target_styles[new_name] = style
                    break

        fix_count = 0

        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            style_name = (para.style.name or "None") if para.style else "None"

            if style_name.lower().startswith('toc ') or style_name == '说明文字':
                continue

            if style_name not in style_map:
                continue

            new_style_name = style_map[style_name]
            new_style = target_styles.get(new_style_name)
            if new_style is None:
                continue

            old_style = style_name
            if not self.dry_run:
                para.style = cast(ParagraphStyle, new_style)

            # 去除手动编号前缀
            pattern = num_patterns.get(new_style_name)
            removed_prefix = ""
            if pattern:
                m = re.match(pattern, text)
                if m:
                    removed_prefix = m.group(0)
                    if not self.dry_run:
                        self._remove_text_prefix(para, len(removed_prefix))

            desc = f"段落{i} \"{text[:40]}\" 样式 {old_style} → {new_style_name}"
            if removed_prefix:
                desc += f"，去除手动编号 \"{removed_prefix.strip()}\""
            self.log_fix("标题样式", desc)
            fix_count += 1

        if fix_count > 0:
            self.log_fix("标题样式", f"共修复 {fix_count} 个标题段落的样式和手动编号")

    def _remove_text_prefix(self, para: Any, char_count: int) -> None:
        """从段落开头移除指定数量的字符，保留 Run 的格式信息。"""
        remaining = char_count
        runs_to_remove = []

        for run in para.runs:
            if remaining <= 0:
                break
            rt = run.text
            if len(rt) <= remaining:
                remaining -= len(rt)
                runs_to_remove.append(run)
            else:
                run.text = rt[remaining:]
                remaining = 0

        for run in runs_to_remove:
            run._element.getparent().remove(run._element)

    # ========================================
    # 11. 修复"说明文字"样式误用
    # ========================================
    def fix_wrong_caption_style(self):
        """将被错误设置为"说明文字"样式的正文段落改为正确的正文样式。"""
        fix_rules = self.rules.get('heading_style_fix', {})
        if not fix_rules.get('enabled'):
            return

        caption_fix = fix_rules.get('caption_style_fix', {})
        if not caption_fix:
            return

        source_style = caption_fix.get('source', '说明文字')
        target_style_name = caption_fix.get('target', '论文正文-首行缩进')

        target_style = None
        for style in self.doc.styles:
            if style.name == target_style_name:
                target_style = style
                break

        if target_style is None:
            return

        fix_count = 0
        for i, para in enumerate(self.doc.paragraphs):
            if not (para.style and para.style.name == source_style and para.text.strip()):
                continue

            # 跳过红色字体的真模板说明
            is_red = False
            for run in para.runs:
                if run.font.color and run.font.color.rgb:
                    rgb = str(run.font.color.rgb)
                    if rgb.upper() in ('FF0000', 'CC0000', 'FF3333'):
                        is_red = True
                        break
            if is_red:
                continue

            text = para.text.strip()
            if not self.dry_run:
                para.style = cast(ParagraphStyle, target_style)
            self.log_fix("说明文字修复",
                         f"段落{i} \"{text[:40]}\" 样式 {source_style} → {target_style_name}")
            fix_count += 1

        if fix_count > 0:
            self.log_fix("说明文字修复", f"共修复 {fix_count} 个误用说明文字样式的段落")

    # ========================================
    # 12. 修复图注样式（Normal/正文 → 图题）
    # ========================================
    def fix_figure_caption_style(self):
        """将以"图X-Y"开头的正文段落改为"图题"样式。"""
        fig_pattern = re.compile(r'^图\s*\d+-\d+')

        fig_style = None
        for style in self.doc.styles:
            if style.name == '图题':
                fig_style = style
                break
        if fig_style is None:
            return

        fixable_sources = {'Normal', '论文正文-首行缩进'}

        fix_count = 0
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else 'None'
            if style_name not in fixable_sources:
                continue
            if not fig_pattern.match(text):
                continue

            self.log_fix("图注样式",
                         f"段落{i} \"{text[:40]}\" 样式 {style_name} → 图题")
            if not self.dry_run:
                para.style = cast(ParagraphStyle, fig_style)
                # 段落级设置 numId=0 禁用自动编号
                pPr = para._element.find('w:pPr', NSMAP)
                if pPr is None:
                    pPr = etree.SubElement(para._element, f'{{{W}}}pPr')
                numPr = pPr.find('w:numPr', NSMAP)
                if numPr is None:
                    numPr = etree.SubElement(pPr, f'{{{W}}}numPr')
                numId_el = numPr.find('w:numId', NSMAP)
                if numId_el is None:
                    numId_el = etree.SubElement(numPr, f'{{{W}}}numId')
                numId_el.set(f'{{{W}}}val', '0')
            fix_count += 1

        if fix_count > 0:
            self.log_fix("图注样式", f"共修复 {fix_count} 个图注段落的样式")

    # ========================================
    # 14. 修复标题段落级多余缩进
    # ========================================
    def fix_heading_paragraph_indent(self):
        """移除标题段落级别多余的缩进属性。"""
        style_rules = self.rules.get('styles', {})
        heading_styles = set()
        for name, rules in style_rules.items():
            if rules.get('paragraph', {}).get('outline_level') is not None:
                heading_styles.add(name)

        fix_count = 0
        for i, para in enumerate(self.doc.paragraphs):
            style_name = para.style.name if para.style else 'None'
            if style_name not in heading_styles:
                continue

            pPr = para._element.find('w:pPr', NSMAP)
            if pPr is None:
                continue
            ind = pPr.find('w:ind', NSMAP)
            if ind is None:
                continue

            left_val = ind.get(f'{{{W}}}left')
            if left_val and left_val != '0':
                text = para.text.strip()
                self.log_fix("标题缩进",
                             f"段落{i} \"{text[:30]}\" 移除段落级 ind left={left_val}")
                if not self.dry_run:
                    if f'{{{W}}}left' in ind.attrib:
                        del ind.attrib[f'{{{W}}}left']
                    if len(ind.attrib) == 0:
                        pPr.remove(ind)
                fix_count += 1

        if fix_count > 0:
            self.log_fix("标题缩进", f"共修复 {fix_count} 个标题段落的多余缩进")
