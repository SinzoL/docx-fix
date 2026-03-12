"""
TextExtractor — 文档段落提取器

从 Word 文档中提取可润色段落，记录完整的 Run 结构信息（文本 + 格式 + 偏移量），
用于后续 LLM 润色和精确回写。
"""

from __future__ import annotations

import re
import logging
from dataclasses import dataclass, field
from typing import Optional

from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# 不可润色段落的识别正则
_RE_FIGURE_CAPTION = re.compile(r"^图\s*\d+[-.]?\d*")
_RE_TABLE_CAPTION = re.compile(r"^表\s*\d+[-.]?\d*")
_RE_REFERENCE_ITEM = re.compile(r"^\[\d+\]")

# 段落最小长度（低于此值跳过润色）
MIN_POLISHABLE_LENGTH = 5


@dataclass
class RunInfo:
    """Run 级别的信息快照"""
    index: int                              # Run 在段落中的索引
    text: str                               # Run 文本
    start_offset: int                       # 在段落纯文本中的起始字符偏移
    end_offset: int                         # 结束字符偏移
    # 格式属性
    font_name: Optional[str] = None         # 字体名（ascii）
    font_name_east_asia: Optional[str] = None  # 中文字体名
    font_size_pt: Optional[float] = None    # 字号（磅）
    bold: Optional[bool] = None             # 加粗
    italic: Optional[bool] = None           # 斜体
    underline: Optional[bool] = None        # 下划线
    color_rgb: Optional[str] = None         # 颜色（RGB 十六进制）
    superscript: Optional[bool] = None      # 上标
    subscript: Optional[bool] = None        # 下标


@dataclass
class ParagraphSnapshot:
    """段落快照 — 记录段落的完整信息，用于提取和回写"""
    index: int                              # 段落在 doc.paragraphs 中的位置
    text: str                               # 纯文本内容（所有 Run 拼接）
    style_name: str                         # 段落样式名
    element_type: str                       # "title" | "narrative" | "list" | "toc" | "caption" | "reference" | "formula" | "empty"
    runs: list[RunInfo] = field(default_factory=list)
    is_polishable: bool = True              # 是否适合润色
    skip_reason: Optional[str] = None       # 跳过原因（当 is_polishable=False 时）


class TextExtractor:
    """文档段落提取器 — 提取可润色段落并记录 Run 结构"""

    def __init__(self, doc: Document) -> None:
        """
        Args:
            doc: python-docx Document 对象
        """
        self._doc = doc
        self._snapshots: list[ParagraphSnapshot] = []
        self._extracted = False

    def extract_paragraphs(self) -> list[ParagraphSnapshot]:
        """提取文档所有段落的快照

        遍历 doc.paragraphs，对每个段落：
        1. 分类为可润色/不可润色
        2. 记录 Run 结构信息（文本 + 格式 + 偏移量）

        Returns:
            段落快照列表（保持文档中的原始顺序）
        """
        self._snapshots = []
        for idx, paragraph in enumerate(self._doc.paragraphs):
            snapshot = self._build_snapshot(idx, paragraph)
            self._snapshots.append(snapshot)

        self._extracted = True
        polishable_count = sum(1 for s in self._snapshots if s.is_polishable)
        logger.info(
            f"段落提取完成：总计 {len(self._snapshots)} 段，"
            f"可润色 {polishable_count} 段，"
            f"跳过 {len(self._snapshots) - polishable_count} 段"
        )
        return self._snapshots

    def get_polishable_paragraphs(self) -> list[ParagraphSnapshot]:
        """获取所有可润色段落

        Returns:
            is_polishable=True 的段落快照列表
        """
        if not self._extracted:
            self.extract_paragraphs()
        return [s for s in self._snapshots if s.is_polishable]

    @staticmethod
    def batch_paragraphs(
        snapshots: list[ParagraphSnapshot],
        batch_size: int = 5,
    ) -> list[list[ParagraphSnapshot]]:
        """将可润色段落分批

        Args:
            snapshots: 段落快照列表（应为可润色段落）
            batch_size: 每批段落数（默认 5）

        Returns:
            分批后的段落快照列表
        """
        polishable = [s for s in snapshots if s.is_polishable]
        return [
            polishable[i:i + batch_size]
            for i in range(0, len(polishable), batch_size)
        ]

    # ----------------------------------------------------------------
    # 内部方法
    # ----------------------------------------------------------------

    def _build_snapshot(self, idx: int, paragraph) -> ParagraphSnapshot:
        """为单个段落构建快照"""
        text = paragraph.text
        style_name = paragraph.style.name if paragraph.style else ""

        # 记录 Run 结构
        runs_info = self._extract_runs(paragraph)

        # 分类段落
        element_type, is_polishable, skip_reason = self._classify_paragraph(
            text, style_name, paragraph
        )

        return ParagraphSnapshot(
            index=idx,
            text=text,
            style_name=style_name,
            element_type=element_type,
            runs=runs_info,
            is_polishable=is_polishable,
            skip_reason=skip_reason,
        )

    def _extract_runs(self, paragraph) -> list[RunInfo]:
        """提取段落中所有 Run 的信息快照"""
        runs_info: list[RunInfo] = []
        offset = 0

        for i, run in enumerate(paragraph.runs):
            run_text = run.text
            start = offset
            end = offset + len(run_text)

            # 提取格式属性
            font = run.font
            info = RunInfo(
                index=i,
                text=run_text,
                start_offset=start,
                end_offset=end,
                font_name=font.name,
                font_name_east_asia=self._get_east_asia_font(run),
                font_size_pt=font.size.pt if font.size else None,
                bold=font.bold,
                italic=font.italic,
                underline=font.underline is not None and font.underline is not False,
                color_rgb=str(font.color.rgb) if font.color and font.color.rgb else None,
                superscript=font.superscript,
                subscript=font.subscript,
            )
            runs_info.append(info)
            offset = end

        return runs_info

    @staticmethod
    def _get_east_asia_font(run) -> Optional[str]:
        """获取 Run 的中文字体名

        python-docx 的 font.name 只返回 ascii 字体，
        东亚字体需要从 XML 中提取。
        """
        try:
            rpr = run._element.find(qn("w:rPr"))
            if rpr is not None:
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is not None:
                    return rfonts.get(qn("w:eastAsia"))
        except Exception:
            pass
        return None

    def _classify_paragraph(
        self, text: str, style_name: str, paragraph
    ) -> tuple[str, bool, Optional[str]]:
        """分类段落类型，判断是否可润色

        Returns:
            (element_type, is_polishable, skip_reason)
        """
        stripped = text.strip()

        # 1. 空段落
        if not stripped:
            return "empty", False, "空段落"

        # 2. 短文本
        if len(stripped) < MIN_POLISHABLE_LENGTH:
            return "narrative", False, f"文本过短（< {MIN_POLISHABLE_LENGTH} 字符）"

        # 3. TOC 段落
        style_lower = style_name.lower()
        if "toc" in style_lower:
            return "toc", False, "目录（TOC）段落"

        # 4. 标题段落（大纲级别 < 9）
        if style_lower.startswith("heading") or style_lower.startswith("标题"):
            return "title", False, "标题段落"

        # 5. 图注
        if _RE_FIGURE_CAPTION.match(stripped):
            return "caption", False, "图注段落"

        # 6. 表注
        if _RE_TABLE_CAPTION.match(stripped):
            return "caption", False, "表注段落"

        # 7. 参考文献列表项
        if _RE_REFERENCE_ITEM.match(stripped):
            return "reference", False, "参考文献列表项"

        # 8. 公式段落（包含 OMath XML 元素）
        if self._has_omath(paragraph):
            return "formula", False, "公式段落"

        # 9. 默认 — 可润色的正文段落
        return "narrative", True, None

    @staticmethod
    def _has_omath(paragraph) -> bool:
        """检查段落是否包含 OMath（公式）元素"""
        try:
            elem = paragraph._element
            # 检查 <m:oMath> 和 <m:oMathPara>
            omath_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
            if elem.findall(f"{{{omath_ns}}}oMath") or elem.findall(f"{{{omath_ns}}}oMathPara"):
                return True
        except Exception:
            pass
        return False
